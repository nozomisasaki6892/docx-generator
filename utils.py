# utils.py
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
# Import các hằng số cần thiết từ config mới
from config import (
    FONT_NAME, FONT_SIZE_DEFAULT, # Lấy cỡ mặc định
    MARGIN_TOP, MARGIN_BOTTOM, MARGIN_LEFT_DEFAULT, MARGIN_RIGHT_DEFAULT,
    MARGIN_LEFT_CONTRACT, MARGIN_RIGHT_CONTRACT, # Giữ lại nếu cần
    FIRST_LINE_INDENT, LINE_SPACING_DEFAULT, LINE_SPACING_BODY
)

def set_paragraph_format(
    paragraph,
    alignment=None,
    left_indent=None,
    right_indent=None,
    first_line_indent=None,
    space_before=None,
    space_after=None,
    line_spacing=None, # Sẽ dùng LINE_SPACING_BODY làm mặc định nếu cần
    line_spacing_rule=WD_LINE_SPACING.MULTIPLE,
    keep_together=None,
    keep_with_next=None,
    page_break_before=None,
    widow_control=True
):
    """Thiết lập định dạng cơ bản cho đối tượng Paragraph."""
    paragraph_format = paragraph.paragraph_format
    if alignment is not None:
        paragraph_format.alignment = alignment
    if left_indent is not None:
        paragraph_format.left_indent = left_indent
    if right_indent is not None:
        paragraph_format.right_indent = right_indent
    if first_line_indent is not None:
        paragraph_format.first_line_indent = first_line_indent
    if space_before is not None:
        paragraph_format.space_before = space_before
    if space_after is not None:
        paragraph_format.space_after = space_after
    # Sử dụng LINE_SPACING_BODY nếu không có giá trị cụ thể truyền vào
    effective_line_spacing = line_spacing if line_spacing is not None else LINE_SPACING_BODY
    if effective_line_spacing is not None:
        paragraph_format.line_spacing = effective_line_spacing
        paragraph_format.line_spacing_rule = line_spacing_rule
    if keep_together is not None:
        paragraph_format.keep_together = keep_together
    if keep_with_next is not None:
        paragraph_format.keep_with_next = keep_with_next
    if page_break_before is not None:
        paragraph_format.page_break_before = page_break_before
    if widow_control is not None:
        paragraph_format.widow_control = widow_control

def set_run_format(
    run,
    font_name=FONT_NAME,
    size=None, # Ưu tiên dùng cỡ chữ từ config được truyền vào
    bold=None,
    italic=None,
    underline=None,
    uppercase=None
):
    """Thiết lập định dạng cơ bản cho đối tượng Run."""
    font = run.font
    font.name = font_name
    try:
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    except AttributeError:
        pass # Bỏ qua nếu không set được
    # Sử dụng FONT_SIZE_DEFAULT nếu không có size cụ thể
    effective_size = size if size is not None else FONT_SIZE_DEFAULT
    if effective_size is not None:
        font.size = effective_size
    if bold is not None:
        font.bold = bold
    if italic is not None:
        font.italic = italic
    if underline is not None:
        font.underline = underline
    if uppercase:
         if run.text != run.text.upper():
              run.text = run.text.upper()

def add_run_with_format(
    paragraph,
    text,
    font_name=FONT_NAME,
    size=None,
    bold=None,
    italic=None,
    underline=None,
    uppercase=None
):
    """Thêm một Run vào Paragraph và áp dụng định dạng cơ bản."""
    run = paragraph.add_run(text)
    set_run_format(run, font_name, size, bold, italic, underline, uppercase)
    return run

def add_paragraph_with_text(
    document,
    text,
    alignment=None,
    left_indent=None,
    right_indent=None,
    first_line_indent=None,
    space_before=None,
    space_after=Pt(6), # Mặc định cách sau 6pt
    line_spacing=LINE_SPACING_BODY, # Mặc định giãn dòng 1.5
    line_spacing_rule=WD_LINE_SPACING.MULTIPLE,
    keep_together=None,
    keep_with_next=None,
    page_break_before=None,
    widow_control=True,
    font_name=FONT_NAME,
    size=FONT_SIZE_DEFAULT, # Mặc định dùng cỡ chữ default
    bold=None,
    italic=None,
    underline=None,
    uppercase=None
):
    """Thêm một Paragraph với text và định dạng đầy đủ, có giá trị mặc định."""
    paragraph = document.add_paragraph()
    set_paragraph_format(
        paragraph, alignment=alignment, left_indent=left_indent,
        right_indent=right_indent, first_line_indent=first_line_indent,
        space_before=space_before, space_after=space_after,
        line_spacing=line_spacing, line_spacing_rule=line_spacing_rule,
        keep_together=keep_together, keep_with_next=keep_with_next,
        page_break_before=page_break_before, widow_control=widow_control
    )
    # Thêm text vào run cuối cùng hoặc tạo run mới nếu chưa có
    if not paragraph.runs:
        run = paragraph.add_run(text)
    else:
        run = paragraph.runs[-1]
        # Nếu run cuối cùng không có text thì dùng nó, ngược lại thêm run mới
        if run.text:
             run = paragraph.add_run(text)
        else:
             run.text = text

    set_run_format(
        run, font_name=font_name, size=size, bold=bold,
        italic=italic, underline=underline, uppercase=uppercase
    )
    return paragraph


def apply_standard_margins(document):
    """Áp dụng lề trang chuẩn NĐ30."""
    try:
        section = document.sections[0]
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = MARGIN_TOP
        section.bottom_margin = MARGIN_BOTTOM
        section.left_margin = MARGIN_LEFT_DEFAULT
        section.right_margin = MARGIN_RIGHT_DEFAULT
        section.orientation = 0 # PORTRAIT
        print("  UTILS: Applied standard margins.", flush=True)
    except Exception as e:
        print(f"  ERROR applying standard margins: {e}", flush=True)

# Các hàm apply margin khác có thể giữ lại nếu cần