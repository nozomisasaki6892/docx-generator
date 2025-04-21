# -*- coding: utf-8 -*-
from flask import Flask, request, send_file, jsonify
from flask_cors import CORS
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn # Để set font Châu Á (Times New Roman)
import io, uuid, os, time, re
import requests # Thư viện để gọi Gemini API

app = Flask(__name__)
CORS(app) # Bật CORS cho phép gọi từ frontend

# --- Cấu hình API Gemini ---
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY", "YOUR_API_KEY_HERE") # Lấy key từ biến môi trường
GEMINI_API_URL = "https://generativelanguage.googleapis.com/v1beta/models/gemini-pro:generateContent"

# --- Prompt cho AI (Có thể tùy chỉnh thêm) ---
AI_PROMPT_TEMPLATE = """
Bạn là một trợ lý soạn thảo văn bản hành chính Việt Nam chuyên nghiệp. Nhiệm vụ của bạn là đọc kỹ nội dung dưới đây, xác định loại văn bản (ví dụ: Nghị định, Tờ trình, Báo cáo, Công văn, Quy chế, Giấy mời, Phát biểu, Tiểu luận), sau đó làm sạch và chuẩn hóa nội dung theo đúng thể thức quy định tại Nghị định 30/2020/NĐ-CP và các quy tắc sau:
1.  **Làm sạch:** Sửa lỗi chính tả, ngữ pháp, loại bỏ từ ngữ thừa, câu lặp, định dạng markdown thừa (nếu có). Đảm bảo văn phong hành chính trang trọng, rõ ràng, mạch lạc.
2.  **Chuẩn hóa cấu trúc:** Dựa vào loại văn bản đã xác định, sắp xếp lại các phần (nếu cần) theo đúng bố cục chuẩn (Vd: Quốc hiệu/Tiêu ngữ -> Số/Ký hiệu -> Địa danh/Ngày tháng -> Tên loại/Trích yếu -> Căn cứ -> Nội dung -> Chữ ký -> Nơi nhận).
3.  **Quốc hiệu/Tiêu ngữ:**
    * Nếu nội dung đầu vào **đã có** "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM" và "Độc lập - Tự do - Hạnh phúc", **không** thêm lại. Chỉ chuẩn hóa viết hoa, in đậm nếu cần.
    * Nếu nội dung đầu vào **thiếu**, hãy **thêm** vào đầu văn bản theo chuẩn:
        ```
        [TÊN CƠ QUAN BAN HÀNH - NẾU CÓ] (In hoa, đậm, căn giữa/trái tùy loại)
        ------------------------------------ (Gạch ngang nếu là cơ quan, tổ chức)
        CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM (In hoa, đậm, căn giữa)
        Độc lập - Tự do - Hạnh phúc (Thường, đậm, căn giữa)
        ------------------------------------ (Gạch ngang dưới tiêu ngữ)
        ```
4.  **Đánh dấu các phần đặc biệt:** Sử dụng các thẻ đánh dấu đơn giản để tôi có thể xử lý định dạng Word sau này:
    * `[QUOCHIEU]`...`[/QUOCHIEU]`
    * `[TIEUNGU]`...`[/TIEUNGU]`
    * `[SOKYHIEU]`Số: .../TTr-ABC`[/SOKYHIEU]`
    * `[DIADANHNGAYTHANG]`Hà Nội, ngày ... tháng ... năm ...`[/DIADANHNGAYTHANG]`
    * `[TENLOAI]`TỜ TRÌNH`[/TENLOAI]`
    * `[TRICHYEU]`Về việc ABC`[/TRICHYEU]`
    * `[CANCU]`Căn cứ Luật...`[/CANCU]`
    * `[KINHGUI]`Kính gửi: ...`[/KINHGUI]`
    * `[CHUONG]`Chương I`[/CHUONG]`
    * `[DIEU]`Điều 1.`[/DIEU]`
    * `[NOIDUNG]`Nội dung chính...`[/NOIDUNG]`
    * `[CHUCVU_QTVK]`(Vd: TM. CHÍNH PHỦ)`[/CHUCVU_QTVK]`
    * `[CHUCVU_NK]`(Vd: THỦ TƯỚNG)`[/CHUCVU_NK]`
    * `[TEN_NK]`Nguyễn Văn A`[/TEN_NK]`
    * `[NOINHAN]`Nơi nhận:`[/NOINHAN]`
    * `[DS_NOINHAN]`- Như trên;`[/DS_NOINHAN]`
    * `[TEN_CQBH]`(Vd: BỘ TÀI CHÍNH)`[/TEN_CQBH]`
5.  **Không thay đổi ý nghĩa gốc** của văn bản.
Hãy xử lý nội dung sau và trả về VĂN BẢN HOÀN CHỈNH đã được làm sạch và đánh dấu các phần tử:

Nội dung cần xử lý:
{text_input}
"""

# --- Hàm tiện ích ---
def set_paragraph_format(paragraph, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, left_indent=Cm(0), first_line_indent=Cm(0), space_before=Pt(0), space_after=Pt(6), line_spacing=1.5, line_spacing_rule=WD_LINE_SPACING.MULTIPLE):
    p_format = paragraph.paragraph_format
    p_format.alignment = alignment
    p_format.left_indent = left_indent
    p_format.first_line_indent = first_line_indent
    p_format.space_before = space_before
    p_format.space_after = space_after
    p_format.line_spacing = line_spacing
    p_format.line_spacing_rule = line_spacing_rule

def set_run_format(run, font_name='Times New Roman', size=Pt(13), bold=False, italic=False, underline=False, uppercase=False):
    font = run.font
    font.name = font_name
    # Set font cho ký tự Châu Á (quan trọng cho tiếng Việt)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    font.size = size
    font.bold = bold
    font.italic = italic
    font.underline = underline
    run.text = run.text.upper() if uppercase else run.text

# --- Hàm nhận diện loại văn bản (Sơ bộ - cần cải thiện) ---
def identify_doc_type(title, body):
    body_upper = body.upper()
    title_upper = title.upper()

    if "NGHỊ ĐỊNH" in title_upper or "NGHỊ ĐỊNH" in body_upper[:500]:
        return "NghiDinh"
    if "QUY CHẾ" in title_upper or "QUY CHẾ" in body_upper[:500]:
        return "QuyChe"
    if "TỜ TRÌNH" in title_upper:
        return "ToTrinh"
    if "PHIẾU TRÌNH" in title_upper:
         return "PhieuTrinh" # Cần xử lý bảng riêng
    if "BÁO CÁO" in title_upper:
        return "BaoCao"
    if "GIẤY MỜI" in title_upper or "TRÂN TRỌNG KÍNH MỜI" in body_upper:
        return "GiayMoi"
    if "PHÁT BIỂU" in title_upper or "KÍNH THƯA" in body_upper[:200]:
        # Kiểm tra thêm nếu không có các dấu hiệu văn bản hành chính khác
        if not ("CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM" in body_upper or "Số:" in body[:100]):
             return "PhatBieu"
    if "TIỂU LUẬN" in title_upper or "LỜI MỞ ĐẦU" in body_upper[:200] or "MỤC LỤC" in body_upper[:200]:
         return "TieuLuan"
    # Mặc định là Công văn hoặc loại khác nếu không xác định được
    return "CongVan" # Hoặc một loại mặc định khác

# --- Hàm gọi Gemini API (Cần xử lý lỗi và nội dung dài) ---
def call_gemini_api_for_cleanup(text_to_clean):
    if not GEMINI_API_KEY or GEMINI_API_KEY == "YOUR_API_KEY_HERE":
        print("WARNING: GEMINI_API_KEY is not set. Skipping AI cleanup.")
        return text_to_clean # Trả về text gốc nếu không có key

    headers = {
        "Content-Type": "application/json",
    }
    # Chia nhỏ text nếu quá dài (ví dụ > 15000 ký tự)
    max_len = 15000
    parts = [text_to_clean[i:i+max_len] for i in range(0, len(text_to_clean), max_len)]
    cleaned_parts = []
    print(f"Chia văn bản thành {len(parts)} phần để xử lý AI.")

    for i, part in enumerate(parts):
        print(f"Đang xử lý phần {i+1}/{len(parts)} bằng Gemini...")
        prompt = AI_PROMPT_TEMPLATE.format(text_input=part)
        payload = {
            "contents": [{"parts": [{"text": prompt}]}],
            "generationConfig": {
                "temperature": 0.5, # Giảm nhiệt độ để kết quả nhất quán hơn
                "maxOutputTokens": 8192 # Tăng giới hạn token ra
            }
        }
        try:
            response = requests.post(f"{GEMINI_API_URL}?key={GEMINI_API_KEY}", headers=headers, json=payload, timeout=120) # Tăng timeout
            response.raise_for_status() # Check lỗi HTTP
            result = response.json()

            # Kiểm tra cấu trúc response của Gemini
            if 'candidates' in result and len(result['candidates']) > 0 and 'content' in result['candidates'][0] and 'parts' in result['candidates'][0]['content'] and len(result['candidates'][0]['content']['parts']) > 0:
                 cleaned_text = result['candidates'][0]['content']['parts'][0]['text']
                 cleaned_parts.append(cleaned_text)
                 print(f"Xử lý xong phần {i+1}.")
            else:
                 print(f"Lỗi: Không tìm thấy nội dung trả về từ Gemini cho phần {i+1}. Payload: {result}")
                 # Có thể trả về phần gốc hoặc báo lỗi
                 cleaned_parts.append(part) # Trả về phần gốc nếu lỗi
            # Nghỉ giữa các lần gọi để tránh rate limit
            if i < len(parts) - 1:
                print("Nghỉ 5 giây...")
                time.sleep(5)

        except requests.exceptions.RequestException as e:
            print(f"Lỗi gọi Gemini API cho phần {i+1}: {e}")
            # Có thể thêm xử lý retry hoặc trả về phần gốc
            cleaned_parts.append(part) # Trả về phần gốc nếu lỗi API
        except Exception as e:
            print(f"Lỗi không xác định khi xử lý Gemini response phần {i+1}: {e}")
            cleaned_parts.append(part)

    print("Ghép các phần đã xử lý...")
    full_cleaned_text = "\n".join(cleaned_parts)
    # print("--- Nội dung sau khi qua AI ---")
    # print(full_cleaned_text[:1000] + "...") # In thử 1000 ký tự đầu
    # print("--- Kết thúc nội dung AI ---")
    return full_cleaned_text

# --- Hàm định dạng Word bằng Python-Docx ---
def format_document_with_python_docx(title, body, doc_type):
    document = Document()
    # --- 1. Thiết lập trang và Style cơ bản ---
    section = document.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    # Đặt lề theo loại văn bản (Nghị định 30 là chuẩn chung)
    if doc_type == "TieuLuan":
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.0)
    else: # Mặc định theo NĐ 30
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(3.5)
        section.right_margin = Cm(2.0)

    # Style chữ chung (Times New Roman, 13pt hoặc 14pt)
    style = document.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    run = style.element.rPr.rFonts
    run.set(qn('w:eastAsia'), 'Times New Roman') # Font Châu Á
    font_size = Pt(14) if doc_type in ["QuyChe", "ToTrinh", "PhieuTrinh", "BaoCao", "GiayMoi", "PhatBieu"] else Pt(13)
    if doc_type == "TieuLuan": font_size = Pt(13) # Tiểu luận thường 13pt
    font.size = font_size
    # --- 2. Xử lý nội dung đã được AI đánh dấu hoặc cấu trúc gốc ---
    # Tách nội dung thành các dòng để xử lý
    lines = body.split('\n')
    # Cờ để kiểm soát việc thêm các phần tử chuẩn (tránh lặp nếu AI đã thêm)
    added_qh_tn = False
    added_noi_nhan_label = False

    # Xử lý từng dòng
    current_paragraph = None
    is_noi_nhan_section = False

    for line in lines:
        stripped_line = line.strip()
        if not stripped_line: # Bỏ qua dòng trống
             # Thêm khoảng cách nếu đang trong nội dung
             if current_paragraph and not is_noi_nhan_section:
                 # document.add_paragraph() # Tạo khoảng cách bằng dòng trống
                 pass # Hoặc dùng space_after
             continue

        # ---- Nhận diện và định dạng các phần tử đặc biệt ----
        # (Phần này cần logic regex hoặc phân tích cấu trúc phức tạp hơn dựa trên thẻ AI)
        # Ví dụ sơ bộ:
        if "[TEN_CQBH]" in line:
            text = re.search(r'\[TEN_CQBH\](.*?)\[/TEN_CQBH\]', line).group(1).strip()
            p = document.add_paragraph()
            set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER)
            run = p.add_run(text)
            set_run_format(run, size=Pt(12), bold=True, uppercase=True)
            # Thêm gạch ngang nếu là Quy chế hoặc Tờ trình,...
            if doc_type in ["QuyChe", "ToTrinh", "BaoCao", "GiayMoi"]:
                 p_line = document.add_paragraph("_______")
                 set_paragraph_format(p_line, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(6))
            added_qh_tn = True # Giả sử AI đã thêm luôn QH TN nếu có CQBH
        elif "[QUOCHIEU]" in line:
             text = re.search(r'\[QUOCHIEU\](.*?)\[/QUOCHIEU\]', line).group(1).strip()
             p = document.add_paragraph()
             set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER)
             run = p.add_run(text)
             set_run_format(run, size=Pt(12), bold=True, uppercase=True)
             added_qh_tn = True
        elif "[TIEUNGU]" in line:
             text = re.search(r'\[TIEUNGU\](.*?)\[/TIEUNGU\]', line).group(1).strip()
             p = document.add_paragraph()
             set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER)
             run = p.add_run(text)
             set_run_format(run, size=Pt(13), bold=True) # Tiêu ngữ cỡ lớn hơn QH
             # Thêm gạch ngang dưới tiêu ngữ
             p_line = document.add_paragraph("------------------------------------")
             set_paragraph_format(p_line, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(12)) # Khoảng cách lớn hơn
             added_qh_tn = True
        elif "[SOKYHIEU]" in line:
            text = re.search(r'\[SOKYHIEU\](.*?)\[/SOKYHIEU\]', line).group(1).strip()
            p = document.add_paragraph()
            set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.LEFT)
            run = p.add_run(text)
            set_run_format(run, size=Pt(13))
        elif "[DIADANHNGAYTHANG]" in line:
             text = re.search(r'\[DIADANHNGAYTHANG\](.*?)\[/DIADANHNGAYTHANG\]', line).group(1).strip()
             # Kiểm tra xem dòng Số ký hiệu đã được thêm chưa, nếu rồi thì căn phải cùng dòng
             # Logic này phức tạp, tạm thời thêm dòng mới căn phải
             p = document.add_paragraph()
             set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
             run = p.add_run(text)
             set_run_format(run, size=Pt(13), italic=True) # Địa danh ngày tháng nghiêng
        elif "[TENLOAI]" in line:
            text = re.search(r'\[TENLOAI\](.*?)\[/TENLOAI\]', line).group(1).strip()
            p = document.add_paragraph()
            # Tăng khoảng cách trước tên loại
            set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=Pt(12), space_after=Pt(6))
            run = p.add_run(text)
            # Cỡ chữ và kiểu chữ tùy loại văn bản
            size = Pt(14)
            is_bold = True
            is_upper = True
            if doc_type == "PhatBieu":
                 size = Pt(16)
            elif doc_type == "TieuLuan":
                 size = Pt(22) # Ví dụ
            set_run_format(run, size=size, bold=is_bold, uppercase=is_upper)
        elif "[TRICHYEU]" in line:
            text = re.search(r'\[TRICHYEU\](.*?)\[/TRICHYEU\]', line).group(1).strip()
            # Bỏ "Về việc" nếu có
            text = text.replace("Về việc", "").replace("V/v", "").strip()
            p = document.add_paragraph()
            set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(12))
            run = p.add_run(f"Về việc {text}") # Thêm lại chuẩn
            size = Pt(14) if doc_type != "TieuLuan" else Pt(16)
            set_run_format(run, size=size, bold=True) # Trích yếu đậm
            # Thêm dòng kẻ dưới trích yếu (NĐ30)
            if doc_type not in ["PhatBieu", "TieuLuan", "PhieuTrinh"]:
                p_line = document.add_paragraph("--------------")
                set_paragraph_format(p_line, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(12))

        elif "[CANCU]" in line:
             text = re.search(r'\[CANCU\](.*?)\[/CANCU\]', line).group(1).strip()
             p = document.add_paragraph()
             set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line_indent=Cm(1.0), line_spacing=1.5)
             run = p.add_run(text)
             set_run_format(run, size=font_size, italic=True) # Căn cứ thường nghiêng
        elif "[KINHGUI]" in line:
             text = re.search(r'\[KINHGUI\](.*?)\[/KINHGUI\]', line).group(1).strip()
             p = document.add_paragraph()
             set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=Pt(6), space_after=Pt(6))
             run = p.add_run(text)
             set_run_format(run, size=font_size, bold=True) # Kính gửi đậm
        elif "[CHUONG]" in line:
             text = re.search(r'\[CHUONG\](.*?)\[/CHUONG\]', line).group(1).strip()
             p = document.add_paragraph()
             set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=Pt(12), space_after=Pt(6))
             run = p.add_run(text.upper()) # Chương in hoa
             set_run_format(run, size=font_size, bold=True)
        elif "[DIEU]" in line:
             text = re.search(r'\[DIEU\](.*?)\[/DIEU\]', line).group(1).strip()
             # Tách số Điều và nội dung điều
             match = re.match(r'(Điều\s+\d+\.?)(.*)', text, re.IGNORECASE)
             if match:
                 dieu_part = match.group(1)
                 title_part = match.group(2).strip()
                 p = document.add_paragraph()
                 set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=Pt(6), space_after=Pt(6))
                 run_dieu = p.add_run(dieu_part + " ")
                 set_run_format(run_dieu, size=font_size, bold=True)
                 if title_part: # Nếu có tên điều
                      run_title = p.add_run(title_part)
                      set_run_format(run_title, size=font_size, bold=True) # Tên điều cũng đậm
             else: # Nếu không tách được thì in đậm cả dòng
                 p = document.add_paragraph()
                 set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=Pt(6), space_after=Pt(6))
                 run = p.add_run(text)
                 set_run_format(run, size=font_size, bold=True)
        elif "[CHUCVU_QTVK]" in line or "[CHUCVU_NK]" in line or "[TEN_NK]" in line:
            # Xử lý khối chữ ký - cần canh phải và có khoảng trống
            is_qtvk = "[CHUCVU_QTVK]" in line
            is_chucvu_nk = "[CHUCVU_NK]" in line
            is_ten_nk = "[TEN_NK]" in line

            if is_qtvk: text = re.search(r'\[CHUCVU_QTVK\](.*?)\[/CHUCVU_QTVK\]', line).group(1).strip()
            elif is_chucvu_nk: text = re.search(r'\[CHUCVU_NK\](.*?)\[/CHUCVU_NK\]', line).group(1).strip()
            else: text = re.search(r'\[TEN_NK\](.*?)\[/TEN_NK\]', line).group(1).strip()

            p = document.add_paragraph()
            # Căn phải cho khối chữ ký
            set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.RIGHT, space_before=Pt(0), space_after=Pt(0))
            run = p.add_run(text)
            size = Pt(13) if doc_type != "TieuLuan" else Pt(14) # Chữ ký thường 13-14
            is_bold = True # Quyền hạn, chức vụ, tên đều đậm
            is_upper = is_qtvk or is_chucvu_nk # Quyền hạn, chức vụ in hoa
            set_run_format(run, size=size, bold=is_bold, uppercase=is_upper)
            # Thêm khoảng trống sau chức vụ, trước tên người ký
            if is_chucvu_nk:
                 for _ in range(4): # Thêm 4 dòng trống cho chữ ký
                      document.add_paragraph()

        elif "[NOINHAN]" in line:
             text = re.search(r'\[NOINHAN\](.*?)\[/NOINHAN\]', line).group(1).strip()
             # Thêm khoảng cách lớn trước nơi nhận
             p = document.add_paragraph()
             set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=Pt(24), space_after=Pt(0))
             run = p.add_run(text) # Ví dụ: "Nơi nhận:"
             set_run_format(run, size=Pt(12), bold=True, italic=True) # Nơi nhận: 12pt, đậm, nghiêng
             is_noi_nhan_section = True
             added_noi_nhan_label = True
        elif "[DS_NOINHAN]" in line or (is_noi_nhan_section and stripped_line.startswith('-')):
             # Xử lý dòng trong danh sách nơi nhận
             if "[DS_NOINHAN]" in line:
                 text = re.search(r'\[DS_NOINHAN\](.*?)\[/DS_NOINHAN\]', line).group(1).strip()
             else:
                 text = stripped_line
             p = document.add_paragraph()
             set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=Pt(0), space_after=Pt(0), line_spacing=1.0)
             run = p.add_run(text)
             set_run_format(run, size=Pt(11)) # Danh sách nơi nhận 11pt
             is_noi_nhan_section = True # Vẫn đang trong mục nơi nhận

        # ---- Nội dung thông thường ----
        else:
            is_noi_nhan_section = False # Thoát khỏi mục nơi nhận
            # Nếu chưa thêm Quốc hiệu/Tiêu ngữ và đây là dòng đầu tiên của nội dung chính
            if not added_qh_tn and doc_type not in ["PhatBieu", "TieuLuan", "PhieuTrinh"]:
                 # Thêm QH TN chuẩn nếu chưa có
                 p_qh = document.add_paragraph("CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM")
                 set_paragraph_format(p_qh, alignment=WD_ALIGN_PARAGRAPH.CENTER)
                 set_run_format(p_qh.runs[0], size=Pt(12), bold=True, uppercase=True)
                 p_tn = document.add_paragraph("Độc lập - Tự do - Hạnh phúc")
                 set_paragraph_format(p_tn, alignment=WD_ALIGN_PARAGRAPH.CENTER)
                 set_run_format(p_tn.runs[0], size=Pt(13), bold=True)
                 p_line = document.add_paragraph("------------------------------------")
                 set_paragraph_format(p_line, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(12))
                 added_qh_tn = True

            # Xử lý đoạn văn bản thường
            # Kiểm tra xem có phải đoạn mới không hay là tiếp tục đoạn cũ
            # Logic này cần cải thiện để xử lý xuống dòng trong cùng 1 đoạn
            current_paragraph = document.add_paragraph()
            first_line_indent = Cm(1.0) if doc_type not in ["PhatBieu", "GiayMoi"] else Cm(0) # Thụt lề đầu dòng cho văn bản thường
            align = WD_ALIGN_PARAGRAPH.JUSTIFY if doc_type not in ["PhatBieu"] else WD_ALIGN_PARAGRAPH.LEFT # Phát biểu canh trái
            line_spacing_val = 1.5 if doc_type != "PhatBieu" else 1.5 # Phát biểu có thể giãn dòng hơn
            set_paragraph_format(current_paragraph, alignment=align, first_line_indent=first_line_indent, line_spacing=line_spacing_val)
            run = current_paragraph.add_run(stripped_line)
            set_run_format(run, size=font_size)

    # --- 3. Thêm nơi nhận mặc định nếu AI không đánh dấu và không phải PhatBieu/TieuLuan ---
    if not added_noi_nhan_label and doc_type not in ["PhatBieu", "TieuLuan", "PhieuTrinh"]:
         p = document.add_paragraph()
         set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=Pt(24), space_after=Pt(0))
         run = p.add_run("Nơi nhận:")
         set_run_format(run, size=Pt(12), bold=True, italic=True)
         p_ds = document.add_paragraph()
         set_paragraph_format(p_ds, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=Pt(0), space_after=Pt(0), line_spacing=1.0)
         run_ds = p_ds.add_run("- Như trên;")
         set_run_format(run_ds, size=Pt(11))
         p_luu = document.add_paragraph()
         set_paragraph_format(p_luu, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=Pt(0), space_after=Pt(0), line_spacing=1.0)
         run_luu = p_luu.add_run("- Lưu: VT, ...") # Thêm đơn vị lưu trữ
         set_run_format(run_luu, size=Pt(11))

    # --- 4. Đánh số trang (Ví dụ: cuối trang, căn giữa - theo NĐ30) ---
    # Logic đánh số trang phức tạp hơn, cần truy cập header/footer
    # Tạm thời bỏ qua trong ví dụ này để đơn giản

    return document

# --- Route chính ---
@app.route("/generate", methods=["POST"])
def generate_docx_route():
    try:
        data = request.get_json(force=True)
        title = data.get("title", "VĂN BẢN CHUNG")
        body = data.get("body", "")

        if not body:
            return jsonify({"error": "Nội dung văn bản không được để trống"}), 400

        print(f"Nhận yêu cầu tạo văn bản: '{title}'")

        # 1. Nhận diện loại văn bản
        doc_type = identify_doc_type(title, body)
        print(f"Loại văn bản nhận diện: {doc_type}")

        # 2. Làm sạch và chuẩn hóa nội dung bằng AI (Nếu có API Key)
        # cleaned_body = call_gemini_api_for_cleanup(body)
        # Tạm thời dùng body gốc để test định dạng python-docx
        cleaned_body = body
        print("Đã bỏ qua bước gọi AI (hoặc AI đã xử lý xong - nếu có key). Bắt đầu định dạng Word...")


        # 3. Tạo file Word với định dạng tương ứng
        document = format_document_with_python_docx(title, cleaned_body, doc_type)
        print("Định dạng Word hoàn tất.")

        # 4. Xuất file Word ra stream để gửi về client
        output_stream = io.BytesIO()
        document.save(output_stream)
        output_stream.seek(0)

        # Tạo tên file động dựa trên title
        safe_title = re.sub(r'[\\/*?:"<>|]', "", title) # Loại bỏ ký tự không hợp lệ
        safe_title = safe_title.replace(" ", "_")
        filename = f"{safe_title}_{uuid.uuid4().hex[:6]}.docx"
        print(f"Tạo file: {filename}")


        return send_file(
            output_stream,
            as_attachment=True,
            download_name=filename,
            # mimetype chuẩn cho .docx
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    except Exception as e:
        print(f"Lỗi nghiêm trọng xảy ra: {e}")
        import traceback
        traceback.print_exc() # In chi tiết lỗi ra log server
        return jsonify({"error": f"Đã xảy ra lỗi trong quá trình xử lý: {str(e)}"}), 500

@app.route("/", methods=["GET"])
def home():
    return jsonify({"message": "API tạo văn bản Word chuẩn hành chính sẵn sàng tại /generate"})

if __name__ == "__main__":
    # Chạy trên cổng 10000 và lắng nghe từ mọi IP cho Render
    app.run(host="0.0.0.0", port=int(os.environ.get("PORT", 10000)), debug=False) # Tắt debug khi deploy thật