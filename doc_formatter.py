# doc_formatter.py
import re
from docx import Document
from docx.shared import Cm, Inches # Đảm bảo có Inches
import traceback # Thêm import

# Import config và các formatter (giữ nguyên)
from config import (MARGIN_TOP, MARGIN_BOTTOM, MARGIN_LEFT_DEFAULT, MARGIN_RIGHT_DEFAULT,
                    MARGIN_LEFT_CONTRACT, FONT_SIZE_DEFAULT, FONT_SIZE_TITLE) # Thêm các hằng số có thể dùng
try:
    # ... (toàn bộ khối import formatters giữ nguyên) ...
    from formatters import (ban_ghi_nho, ban_thoa_thuan, bang_tot_nghiep, bao_cao, bien_ban,
                           chi_thi, chuong_trinh, cong_dien, cong_van, de_an,
                           de_cuong_mh, don_nhap_hoc, du_an, giao_trinh, giay_bao_trung_tuyen,
                           giay_gioi_thieu, giay_moi, giay_nghi_phep, giay_uy_quyen, giay_xac_nhan_sv,
                           hop_dong, huong_dan, huong_dan_hs, ke_hoach, luat,
                           nghi_dinh, nghi_dinh_qppl, nghi_quyet, nghi_quyet_qh, phap_lenh,
                           phat_bieu, phieu, phieu_trinh, phuong_an, quy_che,
                           quy_che_ts, quy_dinh, quy_dinh_nt, quyet_dinh, quyet_dinh_ts,
                           quyet_dinh_ttg, thong_bao, thong_bao_nt, thong_bao_ts, thong_cao,
                           thong_tu, thu_cong, tieu_luan, to_trinh)
    print("Tất cả formatters đã được import thành công (doc_formatter).", flush=True)
except ImportError as e:
    print(f"LỖI IMPORT FORMATTERS (doc_formatter): {e}. Sẽ sử dụng Placeholder.", flush=True)
    # ... (khối fallback giữ nguyên) ...
    ban_ghi_nho = ban_thoa_thuan = bang_tot_nghiep = bao_cao = bien_ban = \
    chi_thi = chuong_trinh = cong_dien = cong_van = de_an = \
    de_cuong_mh = don_nhap_hoc = du_an = giao_trinh = giay_bao_trung_tuyen = \
    giay_gioi_thieu = giay_moi = giay_nghi_phep = giay_uy_quyen = giay_xac_nhan_sv = \
    hop_dong = huong_dan = huong_dan_hs = ke_hoach = luat = \
    nghi_dinh = nghi_dinh_qppl = nghi_quyet = nghi_quyet_qh = phap_lenh = \
    phat_bieu = phieu = phieu_trinh = phuong_an = quy_che = \
    quy_che_ts = quy_dinh = quy_dinh_nt = quyet_dinh = quyet_dinh_ts = \
    quyet_dinh_ttg = thong_bao = thong_bao_nt = thong_bao_ts = thong_cao = \
    thong_tu = thu_cong = tieu_luan = to_trinh = PlaceholderFormatter()


# PlaceholderFormatter class (giữ nguyên)
class PlaceholderFormatter:
    def format(self, document, data):
        print(f"Warning: PlaceholderFormatter activated. Using basic paragraph.", flush=True)
        try:
            document.add_paragraph(data.get('body', '[Placeholder Body - No Formatter Found or Error]'))
        except Exception as e:
            print(f"Error in PlaceholderFormatter: {e}", flush=True)

# DOC_TYPE_FORMATTERS dictionary (giữ nguyên)
DOC_TYPE_FORMATTERS = {
    # ... (giữ nguyên toàn bộ ánh xạ) ...
    "Luat": luat, "NghiQuyetQH": nghi_quyet_qh, "PhapLenh": phap_lenh,
    "NghiDinhQPPL": nghi_dinh_qppl, "QuyetDinhTTg": quyet_dinh_ttg, "ThongTu": thong_tu,
    "CongVan": cong_van, "QuyetDinh": quyet_dinh, "ChiThi": chi_thi, "ThongBao": thong_bao,
    "KeHoach": ke_hoach, "NghiQuyet": nghi_quyet, "QuyDinh": quy_dinh, "HuongDan": huong_dan,
    "ChuongTrinh": chuong_trinh, "BienBan": bien_ban, "DeAn": de_an, "ThongCao": thong_cao,
    "PhuongAn": phuong_an, "DuAn": du_an, "CongDien": cong_dien, "BanGhiNho": ban_ghi_nho,
    "BanThoaThuan": ban_thoa_thuan, "GiayUyQuyen": giay_uy_quyen, "GiayGioiThieu": giay_gioi_thieu,
    "GiayNghiPhep": giay_nghi_phep, "Phieu": phieu, "ThuCong": thu_cong, "HopDong": hop_dong,
    "NghiDinh": nghi_dinh, "QuyChe": quy_che, "ToTrinh": to_trinh, "PhieuTrinh": phieu_trinh,
    "BaoCao": bao_cao, "GiayMoi": giay_moi, "PhatBieu": phat_bieu, "TieuLuan": tieu_luan,
    "ThongBaoTS": thong_bao_ts, "QuyCheTS": quy_che_ts, "HuongDanHS": huong_dan_hs,
    "GiayBaoTrungTuyen": giay_bao_trung_tuyen, "QuyetDinhTS": quyet_dinh_ts,
    "GiayXacNhanSV": giay_xac_nhan_sv, "DonNhapHoc": don_nhap_hoc,
    "DeCuongMH": de_cuong_mh, "QuyDinhNT": quy_dinh_nt,
    "ThongBaoNT": thong_bao_nt, "BangTotNghiep": bang_tot_nghiep, "GiaoTrinh": giao_trinh,
}

# identify_doc_type function (giữ nguyên)
def identify_doc_type(title, body):
    # ... (giữ nguyên logic nhận diện) ...
    # Thêm print ở cuối để biết kết quả nhận diện
    identified_type = None # Biến tạm
    # ... (logic if/else của bạn) ...
    # Ví dụ:
    if "BẰNG TỐT NGHIỆP" in title.upper(): identified_type = "BangTotNghiep"
    # ... (các điều kiện khác) ...
    elif "CÔNG VĂN" in title.upper(): identified_type = "CongVan" # Ví dụ

    print(f"--- DOC_FORMATTER identify_doc_type: Title='{title[:50]}...', Identified='{identified_type}' ---", flush=True)
    return identified_type


# apply_docx_formatting function (thêm nhiều print)
def apply_docx_formatting(data, recognized_doc_type, intended_doc_type):
    print(f"--- apply_docx_formatting START: recognized='{recognized_doc_type}', intended='{intended_doc_type}' ---", flush=True)
    try:
        document = Document()
        print(f"  DEBUG: Initial Document object created.", flush=True)
        section = document.sections[0]
        # Apply margins FIRST
        # Determine margin type based on intended or recognized type BEFORE selecting formatter
        temp_doc_type = intended_doc_type or recognized_doc_type or "CongVan" # Ưu tiên intended
        print(f"  DEBUG: Determining margins based on type '{temp_doc_type}'...", flush=True)
        if temp_doc_type in ["BanGhiNho", "BanThoaThuan", "HopDong"]:
            section.left_margin = MARGIN_LEFT_CONTRACT
            section.right_margin = MARGIN_RIGHT_DEFAULT # Giả sử dùng MARGIN_RIGHT_DEFAULT
            print(f"  DEBUG: Applied CONTRACT margins.", flush=True)
        elif temp_doc_type == "BangTotNghiep":
            section.orientation = 1
            section.page_width = Cm(29.7)
            section.page_height = Cm(21.0)
            section.left_margin = Cm(1.5)
            section.right_margin = Cm(1.5)
            section.top_margin = Cm(1.5)
            section.bottom_margin = Cm(1.5)
            print(f"  DEBUG: Applied LANDSCAPE margins for BangTotNghiep.", flush=True)
        else: # Lề mặc định
            section.left_margin = MARGIN_LEFT_DEFAULT
            section.right_margin = MARGIN_RIGHT_DEFAULT
            print(f"  DEBUG: Applied DEFAULT margins.", flush=True)
        section.top_margin = MARGIN_TOP
        section.bottom_margin = MARGIN_BOTTOM


        formatter_module = None
        doc_type_for_filename = None
        is_placeholder_used = False

        # Logic chọn formatter (Thêm print)
        print(f"  DEBUG: Selecting formatter...", flush=True)
        if recognized_doc_type and recognized_doc_type in DOC_TYPE_FORMATTERS:
            formatter_module = DOC_TYPE_FORMATTERS[recognized_doc_type]
            if not isinstance(formatter_module, PlaceholderFormatter):
                doc_type_for_filename = recognized_doc_type
                print(f"  DEBUG: Selected formatter by RECOGNIZED type: '{recognized_doc_type}' -> {getattr(formatter_module, '__name__', 'N/A')}", flush=True)
            else:
                print(f"  DEBUG: Recognized type '{recognized_doc_type}' maps to PlaceholderFormatter. Will try intended type.", flush=True)
                is_placeholder_used = True
                formatter_module = None # Reset
        else:
             print(f"  DEBUG: Recognized type '{recognized_doc_type}' not found or invalid in DOC_TYPE_FORMATTERS.", flush=True)


        if formatter_module is None and intended_doc_type and intended_doc_type in DOC_TYPE_FORMATTERS:
            formatter_module = DOC_TYPE_FORMATTERS[intended_doc_type]
            if not isinstance(formatter_module, PlaceholderFormatter):
                doc_type_for_filename = intended_doc_type
                print(f"  DEBUG: Selected formatter by INTENDED type: '{intended_doc_type}' -> {getattr(formatter_module, '__name__', 'N/A')}", flush=True)
                is_placeholder_used = False # Reset flag
            else:
                print(f"  DEBUG: Intended type '{intended_doc_type}' maps to PlaceholderFormatter. Will use fallback.", flush=True)
                is_placeholder_used = True
                formatter_module = None # Reset
        elif formatter_module is None:
             print(f"  DEBUG: Intended type '{intended_doc_type}' not found or invalid in DOC_TYPE_FORMATTERS.", flush=True)


        if formatter_module is None:
            if is_placeholder_used:
                 print("  DEBUG: Using PlaceholderFormatter due to previous mapping.", flush=True)
                 formatter_module = PlaceholderFormatter()
                 doc_type_for_filename = "Placeholder"
            else:
                 print("  DEBUG: Using fallback formatter: 'cong_van'", flush=True)
                 formatter_module = DOC_TYPE_FORMATTERS.get("CongVan", PlaceholderFormatter()) # An toàn hơn
                 doc_type_for_filename = "CongVanFallback"
                 if isinstance(formatter_module, PlaceholderFormatter):
                      print("  WARNING: Fallback formatter 'CongVan' also seems to be a Placeholder!", flush=True)


        final_formatter_name = getattr(formatter_module, '__name__', 'PlaceholderFormatter')
        print(f"  DEBUG: Final formatter selected: '{final_formatter_name}'", flush=True)
        print(f"  DEBUG: Data keys being passed: {list(data.keys())}", flush=True) # In các key của data


        if hasattr(formatter_module, 'format'):
            print(f"--- DOC_FORMATTER DEBUG: Calling {final_formatter_name}.format() ---", flush=True)
            formatter_module.format(document, data) # Gọi hàm format
            print(f"--- DOC_FORMATTER DEBUG: Finished calling {final_formatter_name}.format() ---", flush=True)
        else:
            error_msg = f"Lỗi: Module {final_formatter_name} không có hàm format."
            print(error_msg, flush=True)
            document.add_paragraph(error_msg)

        print("--- apply_docx_formatting END ---", flush=True)
        # Sử dụng tên của formatter thực sự được gọi cho tên file
        final_doc_type_name = doc_type_for_filename if doc_type_for_filename else "UnknownType"
        return document, final_doc_type_name

    except Exception as e:
        print(f"!!!!!!!! ERROR in apply_docx_formatting !!!!!!!!!!", flush=True)
        error_trace = traceback.format_exc()
        print(error_trace, flush=True)
        # Tạo document mới chứa lỗi nếu có lỗi nghiêm trọng xảy ra ở đây
        error_doc = Document()
        error_doc.add_paragraph(f"--- LỖI NGHIÊM TRỌNG TRONG apply_docx_formatting ---")
        error_doc.add_paragraph(f"Error: {e}")
        error_doc.add_paragraph(error_trace)
        return error_doc, "ErrorDoc"