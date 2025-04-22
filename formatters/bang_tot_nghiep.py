# formatters/bang_tot_nghiep.py
import time
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_ORIENTATION
from utils import set_paragraph_format, set_run_format, add_run_with_format, add_centered_text
from config import FONT_SIZE_DEFAULT

# Lưu ý: Định dạng bằng tốt nghiệp rất phức tạp và đa dạng, phụ thuộc mẫu cụ thể.
# Code này chỉ tạo cấu trúc cơ bản, cần tùy chỉnh nhiều theo mẫu thực tế.

def format(document, data):
    print("Bắt đầu định dạng Bằng tốt nghiệp (cơ bản)...")

    # Lấy dữ liệu - cần chuẩn hóa key từ data
    university_name = data.get("university_name", "TRƯỜNG ĐẠI HỌC ABC").upper()
    degree_title = data.get("degree_title", "BẰNG CỬ NHÂN").upper()
    student_name = data.get("student_name", "Nguyễn Văn A").upper()
    dob = data.get("dob", "01/01/2000")
    major = data.get("major", "Công nghệ Thông tin")
    graduation_rank = data.get("graduation_rank", "Xuất sắc")
    degree_mod = data.get("degree_mod", "Chính quy") # Mode of Degree
    decision_number = data.get("decision_number", "Số 123/QĐ-ĐHABC")
    decision_date = data.get("decision_date", "ngày 15 tháng 6 năm 2025")
    diploma_number = data.get("diploma_number", "Số vào sổ: 12345")
    issuing_location = data.get("issuing_location", "Hà Nội")
    issuing_date_str = data.get("issuing_date", time.strftime(f"ngày %d tháng %m năm %Y"))
    rector_title = data.get("rector_title", "HIỆU TRƯỞNG").upper()
    rector_name = data.get("rector_name", "GS.TS. Trần Văn B")


    # Định dạng trang ngang (Landscape) - Di chuyển ra doc_formatter.py
    # section = document.sections[0]
    # section.orientation = WD_ORIENTATION.LANDSCAPE
    # section.page_width = Cm(29.7)
    # section.page_height = Cm(21.0)
    # section.left_margin = Cm(1.5)
    # section.right_margin = Cm(1.5)
    # section.top_margin = Cm(1.5)
    # section.bottom_margin = Cm(1.5)

    # Cấu trúc thường dùng Table để định vị
    # Ví dụ đơn giản dùng paragraph căn giữa

    add_centered_text(document, "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM", size=Pt(12), bold=True, space_after=0)
    add_centered_text(document, "Độc lập - Tự do - Hạnh phúc", size=Pt(13), bold=True, space_after=12)

    add_centered_text(document, university_name, size=Pt(14), bold=True, space_after=18)

    add_centered_text(document, degree_title, size=Pt(24), bold=True, space_after=18)

    add_centered_text(document, f"Chứng nhận Sinh viên: {student_name}", size=Pt(14), bold=True, space_after=6)
    add_centered_text(document, f"Ngày sinh: {dob}", size=Pt(14), space_after=12)

    add_centered_text(document, f"Đã tốt nghiệp ngành: {major}", size=Pt(14), bold=True, space_after=6)
    add_centered_text(document, f"Xếp loại tốt nghiệp: {graduation_rank}", size=Pt(14), space_after=6)
    add_centered_text(document, f"Hình thức đào tạo: {degree_mod}", size=Pt(14), space_after=12)

    add_centered_text(document, f"Theo Quyết định số {decision_number} {decision_date}", size=Pt(12), space_after=18)


    # Số vào sổ và Ngày cấp
    table = document.add_table(rows=1, cols=2)
    table.autofit = False
    table.allow_autofit = False
    # Điều chỉnh độ rộng cột phù hợp với trang ngang
    table.columns[0].width = Inches(4.5)
    table.columns[1].width = Inches(4.5)

    cell_left = table.cell(0, 0)
    cell_right = table.cell(0, 1)
    cell_left._element.clear_content()
    cell_right._element.clear_content()

    p_diploma_num = cell_left.add_paragraph(diploma_number)
    set_paragraph_format(p_diploma_num, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=Pt(12))
    set_run_format(p_diploma_num.runs[0], size=Pt(12))

    p_issue_date = cell_right.add_paragraph(f"{issuing_location}, {issuing_date_str}")
    set_paragraph_format(p_issue_date, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=Pt(12))
    set_run_format(p_issue_date.runs[0], size=Pt(12), italic=True)


    # Chữ ký Hiệu trưởng
    p_rector_title = cell_right.add_paragraph(rector_title)
    set_paragraph_format(p_rector_title, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=Pt(12), space_after=Pt(6))
    set_run_format(p_rector_title.runs[0], size=Pt(14), bold=True)
    cell_right.add_paragraph("\n\n\n\n")
    p_rector_name = cell_right.add_paragraph(rector_name)
    set_paragraph_format(p_rector_name, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_run_format(p_rector_name.runs[0], size=Pt(14), bold=True)


    print("Định dạng Bằng tốt nghiệp (cơ bản) hoàn tất.")