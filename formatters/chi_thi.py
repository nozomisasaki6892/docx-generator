# formatters/chi_thi.py
import re
import time
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from utils import set_paragraph_format, set_run_format, add_run_with_format
try:
    from .common_elements import format_basic_header, format_signature_block, format_recipient_list
except ImportError:
    # Fallback nếu chạy độc lập
    def format_basic_header(document, data, doc_type): pass
    def format_signature_block(document, data): pass
    def format_recipient_list(document, data): pass

from config import FONT_SIZE_DEFAULT, FONT_SIZE_TITLE, FIRST_LINE_INDENT

def format(document, data):
    print("Bắt đầu định dạng Chỉ thị...")
    title = data.get("title", "Chỉ thị về việc Tăng cường ABC")
    body = data.get("body", "Để thực hiện..., Thủ tướng Chính phủ yêu cầu:\n1. Các Bộ, ngành...\n2. UBND các tỉnh...")
    doc_type_label = "CHỈ THỊ"

    # 1. Header (Sử dụng header cơ bản)
    # Chỉ thị của Thủ tướng/Chủ tịch UBND... dùng header cơ bản
    format_basic_header(document, data, "ChiThi")

    # 2. Tên loại và Trích yếu
    p_title = document.add_paragraph(doc_type_label)
    set_paragraph_format(p_title, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=Pt(12), space_after=Pt(6))
    set_run_format(p_title.runs[0], size=FONT_SIZE_TITLE, bold=True)

    subject = title.replace("Chỉ thị", "").strip()
    p_subject = document.add_paragraph(f"Về việc {subject}")
    set_paragraph_format(p_subject, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(12))
    set_run_format(p_subject.runs[0], size=Pt(14), bold=True)


    # 3. Nội dung Chỉ thị
    body_lines = body.split('\n')
    for line in body_lines:
        stripped_line = line.strip()
        if not stripped_line: continue
        p = document.add_paragraph()

        # Logic định dạng cơ bản cho các yêu cầu trong chỉ thị
        is_numbered_item = re.match(r'^(\d+)\.\s+', stripped_line) # 1., 2.
        is_alpha_item = re.match(r'^[a-z]\)\s+', stripped_line) # a), b)
        is_bullet = stripped_line.startswith("-") or stripped_line.startswith("+") or stripped_line.startswith("*")

        align = WD_ALIGN_PARAGRAPH.JUSTIFY
        left_indent = Cm(0)
        first_indent = FIRST_LINE_INDENT if not (is_numbered_item or is_alpha_item or is_bullet) else Cm(0)
        is_bold = False # Nội dung chỉ thị thường không đậm
        is_italic = False
        size = FONT_SIZE_DEFAULT
        space_before = Pt(0)
        space_after = Pt(6)
        line_spacing = 1.5

        if is_numbered_item:
            align = WD_ALIGN_PARAGRAPH.LEFT # Căn trái mục lớn
            left_indent = Cm(0.5) # Thụt lề mục
            # is_bold = True # Đầu mục có thể đậm
        elif is_alpha_item:
             align = WD_ALIGN_PARAGRAPH.LEFT
             left_indent = Cm(1.0)
        elif is_bullet:
             align = WD_ALIGN_PARAGRAPH.LEFT
             left_indent = Cm(1.5)
             first_indent = Cm(-0.5) # Hanging indent

        set_paragraph_format(p, alignment=align, left_indent=left_indent, first_line_indent=first_indent, line_spacing=line_spacing, space_before=space_before, space_after=space_after)
        add_run_with_format(p, stripped_line, size=size, bold=is_bold, italic=is_italic)


    # 4. Yêu cầu tổ chức thực hiện (Thường là đoạn cuối)
    # Xử lý như nội dung thông thường


    # 5. Chữ ký (Sử dụng khối chữ ký cơ bản)
    # Cần xác định rõ người ký và chức vụ từ data (VD: THỦ TƯỚNG, CHỦ TỊCH UBND...)
    if 'signer_title' not in data: data['signer_title'] = "CHỨC VỤ NGƯỜI KÝ"
    if 'signer_name' not in data: data['signer_name'] = "[Họ và tên]"
    document.add_paragraph() # Khoảng cách trước chữ ký
    format_signature_block(document, data)

    # 6. Nơi nhận (Sử dụng nơi nhận cơ bản)
    if 'recipients' not in data: data['recipients'] = ["- Như trên;", "- Lưu: VT, ...;"]
    format_recipient_list(document, data)

    print("Định dạng Chỉ thị hoàn tất.")