# formatters/nghi_dinh_qppl.py
import re
import time
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
# Header của Chính phủ giống Luật, Chữ ký khác
from .luat import format_qppl_header
from utils import set_paragraph_format, set_run_format, add_run_with_format
from config import FONT_SIZE_DEFAULT, FONT_SIZE_TITLE, FIRST_LINE_INDENT, FONT_SIZE_SIGNATURE, FONT_SIZE_SIGNER_NAME

def format_cp_signature(document, signer_title, signer_name):
     # Chữ ký NĐ CP thường căn phải
     sig_paragraph = document.add_paragraph()
     set_paragraph_format(sig_paragraph, alignment=WD_ALIGN_PARAGRAPH.RIGHT, space_before=Pt(6), space_after=Pt(0), line_spacing=1.15)
     # TM. CHÍNH PHỦ
     add_run_with_format(sig_paragraph, "TM. CHÍNH PHỦ\n", size=FONT_SIZE_SIGNATURE, bold=True)
     # THỦ TƯỚNG
     add_run_with_format(sig_paragraph, signer_title.upper() + "\n\n\n\n\n", size=FONT_SIZE_SIGNATURE, bold=True)
     # Tên Thủ tướng
     add_run_with_format(sig_paragraph, signer_name, size=FONT_SIZE_SIGNER_NAME, bold=True)


def format(document, data):
    print("Bắt đầu định dạng Nghị định QPPL...")
    title = data.get("title", "Nghị định quy định chi tiết ABC")
    body = data.get("body", "Nội dung nghị định...")
    decree_number = data.get("decree_number", "Nghị định số: .../20.../NĐ-CP")
    issuing_date_str = data.get("issuing_date", time.strftime("ngày %d tháng %m năm %Y"))
    issuing_location = data.get("issuing_location", "Hà Nội")

    # 1. Header (CHÍNH PHỦ)
    format_qppl_header(document, "CHÍNH PHỦ")

    # 2. Số hiệu và Ngày tháng ban hành (căn giữa)
    p_num_date = document.add_paragraph(f"{decree_number}\n{issuing_location}, {issuing_date_str}")
    set_paragraph_format(p_num_date, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(12))
    add_run_with_format(p_num_date, p_num_date.text, size=FONT_SIZE_DEFAULT)


    # 3. Tên Nghị định
    nd_title = title.upper() # Thường đã có chữ NGHỊ ĐỊNH
    p_tenloai = document.add_paragraph(nd_title)
    set_paragraph_format(p_tenloai, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=Pt(6), space_after=Pt(12))
    add_run_with_format(p_tenloai, p_tenloai.text, size=FONT_SIZE_TITLE, bold=True, uppercase=True)

    # 4. Cơ quan ban hành (CHÍNH PHỦ)
    p_issuer = document.add_paragraph("CHÍNH PHỦ")
    set_paragraph_format(p_issuer, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(6))
    add_run_with_format(p_issuer, p_issuer.text, size=FONT_SIZE_DEFAULT, bold=True, uppercase=True)

    # 5. Căn cứ ban hành
    preamble = data.get("preamble", "Căn cứ Luật Tổ chức Chính phủ ngày ... tháng ... năm ...;\nCăn cứ [Luật/Pháp lệnh được hướng dẫn];\nTheo đề nghị của [Bộ trưởng/Thủ trưởng cơ quan];\nChính phủ ban hành Nghị định ...")
    preamble_lines = preamble.split('\n')
    for line in preamble_lines:
         p_pre = document.add_paragraph(line)
         set_paragraph_format(p_pre, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line_indent=FIRST_LINE_INDENT, space_after=Pt(0), line_spacing=1.5)
         add_run_with_format(p_pre, line, size=FONT_SIZE_DEFAULT, italic=True)
    document.add_paragraph()

    # 6. Nội dung (Chương, Mục, Điều, Khoản, Điểm) - Tương tự Luật
    body_lines = body.split('\n')
    for line in body_lines:
        stripped_line = line.strip()
        if not stripped_line: continue
        p = document.add_paragraph()
        # (Copy logic xử lý Chương, Mục, Điều, Khoản, Điểm từ formatters/luat.py)
        is_chuong = stripped_line.upper().startswith("CHƯƠNG")
        is_muc = re.match(r'^(MỤC\s+\d+)\.?\s+', stripped_line.upper())
        is_dieu = stripped_line.upper().startswith("ĐIỀU")
        is_khoan = re.match(r'^\d+\.\s+', stripped_line)
        is_diem = re.match(r'^[a-z]\)\s+', stripped_line)

        align = WD_ALIGN_PARAGRAPH.JUSTIFY
        left_indent = Cm(0)
        first_indent = FIRST_LINE_INDENT
        is_bold = False
        size = FONT_SIZE_DEFAULT
        space_before = Pt(0)
        space_after = Pt(6)

        if is_chuong:
            align = WD_ALIGN_PARAGRAPH.CENTER
            first_indent = Cm(0)
            is_bold = True
            space_before = Pt(12)
        elif is_muc:
            align = WD_ALIGN_PARAGRAPH.CENTER
            first_indent = Cm(0)
            is_bold = True
            space_before = Pt(6)
        elif is_dieu:
            align = WD_ALIGN_PARAGRAPH.LEFT
            first_indent = Cm(0)
            is_bold = True
            space_before = Pt(6)
        elif is_khoan:
            left_indent = Cm(0.5)
            first_indent = Cm(0)
        elif is_diem:
            left_indent = Cm(1.0)
            first_indent = Cm(0)

        set_paragraph_format(p, alignment=align, left_indent=left_indent, first_line_indent=first_indent, line_spacing=1.5, space_before=space_before, space_after=space_after)
        add_run_with_format(p, stripped_line, size=size, bold=is_bold)

    # 7. Chữ ký (TM. CHÍNH PHỦ, THỦ TƯỚNG)
    format_cp_signature(document, "THỦ TƯỚNG", data.get("signer_name", "[Tên Thủ tướng]"))

    # Nghị định thường không có nơi nhận ở cuối

    print("Định dạng Nghị định QPPL hoàn tất.")