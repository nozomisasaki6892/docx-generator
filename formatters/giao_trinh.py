# formatters/giao_trinh.py
import re
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from utils import set_paragraph_format, set_run_format, add_run_with_format, add_centered_text
from config import FONT_SIZE_DEFAULT, FONT_SIZE_TITLE, FIRST_LINE_INDENT

# Lưu ý: Giáo trình có cấu trúc rất phức tạp (Mục lục, Lời nói đầu, Chương, Mục, Hình ảnh, Bảng biểu, Tài liệu tham khảo...)
# Code này chỉ cung cấp định dạng rất cơ bản cho tiêu đề và các đề mục chính.

def format(document, data):
    print("Bắt đầu định dạng Giáo trình (cơ bản)...")
    title = data.get("title", "Giáo trình ABC")
    body = data.get("body", "Lời nói đầu...\nChương 1:...\n   1.1 Mục...\n   1.2 Mục...\nChương 2:...")
    authors = data.get("authors", ["[Tên tác giả 1]", "[Tên tác giả 2]"]) # List of authors
    publisher = data.get("publisher", "Nhà xuất bản XYZ")
    publish_year = data.get("publish_year", "2025")


    # --- Trang bìa (Giả định, cần thiết kế chi tiết) ---
    # Thông tin cơ quan chủ quản, trường (Nếu có)
    add_centered_text(document, data.get("university_name", "TRƯỜNG ĐẠI HỌC").upper(), size=Pt(14), bold=True, space_after=12)
    # Tên tác giả
    for author in authors:
        add_centered_text(document, author, size=Pt(14), bold=True, space_after=0)
    document.add_paragraph("\n") # Khoảng cách

    # Tên giáo trình
    p_title = document.add_paragraph(title.upper())
    set_paragraph_format(p_title, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=Pt(36), space_after=Pt(36))
    set_run_format(p_title.runs[0], size=Pt(24), bold=True) # Cỡ chữ lớn cho tên GT

    # Thông tin Nhà xuất bản, Năm XB
    add_centered_text(document, publisher.upper(), size=Pt(14), bold=True, space_before=Pt(72), space_after=0)
    add_centered_text(document, publish_year, size=Pt(14), bold=True, space_after=0)

    document.add_page_break()

    # --- Nội dung Giáo trình ---
    # Cần có logic phức tạp để xử lý Mục lục, Lời nói đầu, Tham khảo...
    # Code dưới đây chỉ định dạng cơ bản Chương, Mục

    body_lines = body.split('\n')
    for line in body_lines:
        stripped_line = line.strip()
        if not stripped_line: continue
        p = document.add_paragraph()

        # Logic định dạng cơ bản
        is_chapter = stripped_line.upper().startswith("CHƯƠNG")
        is_section_roman = re.match(r'^([IVXLCDM]+)\.\s+', stripped_line) # Mục La Mã I, II
        is_section_digit = re.match(r'^(\d+)\.\s+', stripped_line) # Mục 1, 2, 3
        is_subsection_digit = re.match(r'^(\d+\.\d+\.?)\s+', stripped_line) # Mục 1.1, 1.2
        is_subsubsection_alpha = re.match(r'^[a-z]\)\s+', stripped_line) # Mục a, b, c
        is_bullet = stripped_line.startswith("-") or stripped_line.startswith("+") or stripped_line.startswith("*") or stripped_line.startswith("•")

        align = WD_ALIGN_PARAGRAPH.JUSTIFY
        left_indent = Cm(0)
        first_indent = FIRST_LINE_INDENT if not (is_chapter or is_section_roman or is_section_digit or is_subsection_digit or is_subsubsection_alpha or is_bullet) else Cm(0)
        is_bold = bool(is_chapter or is_section_roman or is_section_digit or is_subsection_digit) # Mục lớn đậm
        is_italic = False
        size = FONT_SIZE_DEFAULT # Cỡ chữ giáo trình thường là 13pt
        space_before = Pt(0)
        space_after = Pt(6)
        line_spacing = 1.5

        if is_chapter:
            align = WD_ALIGN_PARAGRAPH.CENTER
            space_before = Pt(18)
            space_after = Pt(12)
            size = Pt(14) # Chương to hơn
        elif is_section_roman or is_section_digit:
            align = WD_ALIGN_PARAGRAPH.LEFT
            space_before = Pt(12)
            size = Pt(13)
        elif is_subsection_digit:
            align = WD_ALIGN_PARAGRAPH.LEFT
            left_indent = Cm(0.5)
            space_before = Pt(6)
        elif is_subsubsection_alpha:
             align = WD_ALIGN_PARAGRAPH.LEFT
             left_indent = Cm(1.0)
        elif is_bullet:
             align = WD_ALIGN_PARAGRAPH.LEFT
             left_indent = Cm(1.5)
             first_indent = Cm(-0.5) # Hanging indent

        set_paragraph_format(p, alignment=align, left_indent=left_indent, first_line_indent=first_indent, line_spacing=line_spacing, space_before=space_before, space_after=space_after)
        add_run_with_format(p, stripped_line, size=size, bold=is_bold, italic=is_italic)


    print("Định dạng Giáo trình (cơ bản) hoàn tất.")