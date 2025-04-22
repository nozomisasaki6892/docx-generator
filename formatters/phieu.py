# formatters/phieu.py
import time
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from utils import set_paragraph_format, set_run_format, add_run_with_format
try:
    # Phiếu thường không có header chuẩn, có chữ ký và nơi nhận
    from .common_elements import format_signature_block, format_recipient_list
except ImportError:
    def format_signature_block(document, data): pass
    def format_recipient_list(document, data): pass

from config import FONT_SIZE_DEFAULT, FONT_SIZE_TITLE, FIRST_LINE_INDENT, FONT_SIZE_HEADER

def format(document, data):
    print("Bắt đầu định dạng Phiếu...")
    title = data.get("title", "PHIẾU GỬI").upper() # VD: PHIẾU GỬI, PHIẾU CHUYỂN...
    body = data.get("body", "Kính gửi:...\nNội dung:...\nYêu cầu xử lý:...")
    issuing_org = data.get("issuing_org", "TÊN CƠ QUAN/ĐƠN VỊ").upper()
    doc_number = data.get("doc_number", "Số: ...... /PG-...")
    issuing_date_str = data.get("issuing_date", time.strftime(f"ngày %d tháng %m năm %Y"))


    # 1. Thông tin cơ quan và Số hiệu (Có thể căn góc trái hoặc giữa)
    p_org = document.add_paragraph(issuing_org)
    set_paragraph_format(p_org, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(0))
    set_run_format(p_org.runs[0], size=FONT_SIZE_HEADER, bold=True)

    p_num = document.add_paragraph(doc_number)
    set_paragraph_format(p_num, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(6))
    set_run_format(p_num.runs[0], size=Pt(13))

    p_line = document.add_paragraph("-----------")
    set_paragraph_format(p_line, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(12))


    # 2. Tên Phiếu
    p_title = document.add_paragraph(title)
    set_paragraph_format(p_title, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=Pt(6), space_after=Pt(12))
    set_run_format(p_title.runs[0], size=FONT_SIZE_TITLE, bold=True)


    # 3. Kính gửi / Nơi nhận chính
    recipient = data.get("recipient_main", "Kính gửi: [Tên đơn vị/cá nhân nhận]")
    p_kg = document.add_paragraph(recipient)
    set_paragraph_format(p_kg, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=Pt(12))
    set_run_format(p_kg.runs[0], size=FONT_SIZE_DEFAULT, bold=True)


    # 4. Nội dung Phiếu
    body_lines = body.split('\n')
    for line in body_lines:
        stripped_line = line.strip()
        if not stripped_line: continue

        p = document.add_paragraph()
        # Nội dung phiếu thường căn trái
        is_info_line = ":" in stripped_line and len(stripped_line.split(":")[0]) < 30

        align = WD_ALIGN_PARAGRAPH.LEFT
        first_indent = FIRST_LINE_INDENT if not is_info_line else Cm(0)
        is_bold = False

        set_paragraph_format(p, alignment=align, space_after=Pt(6), first_line_indent=first_indent, line_spacing=1.5)

        if is_info_line:
            parts = stripped_line.split(":", 1)
            add_run_with_format(p, parts[0] + ":", size=FONT_SIZE_DEFAULT, bold=True) # Nhãn đậm
            add_run_with_format(p, parts[1], size=FONT_SIZE_DEFAULT)
        else:
             add_run_with_format(p, stripped_line, size=FONT_SIZE_DEFAULT)


    # 5. Ngày tháng lập phiếu (Nếu khác ngày ban hành)
    # Có thể thêm ở góc phải dưới


    # 6. Chữ ký (Người lập phiếu / Thủ trưởng đơn vị)
    if 'signer_title' not in data: data['signer_title'] = "NGƯỜI LẬP PHIẾU" # Hoặc chức vụ khác
    if 'signer_name' not in data: data['signer_name'] = "[Họ và tên]"
    document.add_paragraph()
    format_signature_block(document, data) # Dùng signature block chuẩn


    # 7. Nơi nhận (Nếu cần gửi nhiều nơi)
    if data.get('recipients'):
         format_recipient_list(document, data)


    print("Định dạng Phiếu hoàn tất.")