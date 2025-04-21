# formatters/bang_diem.py
import re
import time
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_ALIGN_VERTICAL
from utils import set_paragraph_format, set_run_format, add_run_with_format
# Giả sử dùng cell helper từ thoi_khoa_bieu
try:
    from .thoi_khoa_bieu import set_cell_properties
    from .common_elements import format_basic_header # Header trường
except ImportError:
    print("Warning: Cannot import helper from thoi_khoa_bieu.py")
    # Định nghĩa lại set_cell_properties nếu cần
    def set_cell_properties(cell, text, bold=False, italic=False, size=Pt(10), align=WD_ALIGN_PARAGRAPH.CENTER, valign=WD_ALIGN_VERTICAL.CENTER):
        if len(cell.paragraphs) > 0: p = cell.paragraphs[0]; p.clear()
        else: p = cell.add_paragraph()
        run = p.add_run(text); set_run_format(run, size=size, bold=bold, italic=italic)
        p.alignment = align; cell.vertical_alignment = valign
    from common_elements import format_basic_header


from config import FONT_SIZE_DEFAULT, FONT_SIZE_TITLE, FONT_NAME, FONT_SIZE_HEADER

def format(document, data):
    print("Bắt đầu định dạng Bảng điểm...")
    # Thông tin sinh viên
    student = data.get("student", {"name": "[HỌ TÊN]", "id": "[Mã SV]", "class": "[Lớp]", "major": "[Ngành]", "faculty": "[Khoa]", "level": "Đại học", "mode": "Chính quy"})
    # Dữ liệu điểm: list of dictionaries
    grades_data = data.get("grades_data", [
        {"stt": 1, "ma_hp": "MATH101", "ten_hp": "Toán cao cấp A1", "tin_chi": 3, "diem_qt": 8.0, "diem_thi": 7.5, "diem_tk": 7.7, "diem_chu": "B", "ghi_chu": ""},
        {"stt": 2, "ma_hp": "PHYS101", "ten_hp": "Vật lý đại cương", "tin_chi": 3, "diem_qt": 7.0, "diem_thi": 6.5, "diem_tk": 6.7, "diem_chu": "C+", "ghi_chu": ""},
    ])
    # Thông tin tổng kết (có thể tính toán hoặc nhận từ data)
    summary = data.get("summary", {"total_credits_registered": 6, "total_credits_earned": 6, "gpa_term": 7.2, "gpa_cumulative": 7.2})
    issuing_org = data.get("issuing_org", "TÊN TRƯỜNG").upper()
    issuing_dept = data.get("issuing_dept", "PHÒNG ĐÀO TẠO").upper()

    # 1. Header (Trường, Phòng Đào tạo)
    p_org = document.add_paragraph(issuing_org)
    set_paragraph_format(p_org, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(0))
    add_run_with_format(p_org, issuing_org, size=FONT_SIZE_HEADER, bold=True)
    p_dept = document.add_paragraph(issuing_dept)
    set_paragraph_format(p_dept, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(12))
    add_run_with_format(p_dept, issuing_dept, size=Pt(11), bold=True)

    # 2. Tên Bảng điểm
    term = data.get("term", "Học kỳ ... Năm học ...") # VD: Học kỳ I Năm học 2024-2025
    p_title = document.add_paragraph(f"BẢNG ĐIỂM HỌC TẬP\n{term.upper()}")
    set_paragraph_format(p_title, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(12))
    add_run_with_format(p_title, p_title.text, size=FONT_SIZE_TITLE, bold=True, uppercase=True)

    # 3. Thông tin sinh viên
    # Có thể dùng table 2 cột cho đẹp
    info_table = document.add_table(rows=3, cols=4) # Ví dụ chia 4 cột
    info_table.autofit = False
    col_widths = [Inches(1.0), Inches(2.0), Inches(1.0), Inches(2.0)]
    for i, w in enumerate(col_widths): info_table.columns[i].width = w

    # Hàng 1
    set_cell_properties(info_table.cell(0, 0), "Họ và tên:", bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)
    set_cell_properties(info_table.cell(0, 1), student['name'], bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)
    set_cell_properties(info_table.cell(0, 2), "Mã số SV:", bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)
    set_cell_properties(info_table.cell(0, 3), student['id'], align=WD_ALIGN_PARAGRAPH.LEFT)
    # Hàng 2
    set_cell_properties(info_table.cell(1, 0), "Lớp:", bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)
    set_cell_properties(info_table.cell(1, 1), student['class'], align=WD_ALIGN_PARAGRAPH.LEFT)
    set_cell_properties(info_table.cell(1, 2), "Ngành:", bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)
    set_cell_properties(info_table.cell(1, 3), student['major'], align=WD_ALIGN_PARAGRAPH.LEFT)
    # Hàng 3
    set_cell_properties(info_table.cell(2, 0), "Khoa:", bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)
    set_cell_properties(info_table.cell(2, 1), student['faculty'], align=WD_ALIGN_PARAGRAPH.LEFT)
    set_cell_properties(info_table.cell(2, 2), "Hệ/Mode:", bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)
    set_cell_properties(info_table.cell(2, 3), f"{student['level']}/{student['mode']}", align=WD_ALIGN_PARAGRAPH.LEFT)

    document.add_paragraph() # Khoảng cách

    # 4. Bảng điểm chi tiết
    num_grades = len(grades_data)
    # Xác định các cột cần hiển thị (tùy theo data)
    grade_headers = ["STT", "Mã HP", "Tên học phần", "Số TC", "Điểm QT", "Điểm Thi", "Điểm TK (10)", "Điểm Chữ", "Ghi chú"]
    num_grade_cols = len(grade_headers)

    grade_table = document.add_table(rows=num_grades + 1, cols=num_grade_cols)
    grade_table.style = 'Table Grid'
    grade_table.autofit = False
    # Thiết lập độ rộng cột điểm
    # Ước lượng: STT(0.4), Mã HP(0.8), Tên HP(2.5), TC(0.5), Điểm QT(0.6), Thi(0.6), TK10(0.7), Chữ(0.6), Ghi chú(1.0) -> Total ~8.3 inches (Cần >6.3)
    # Điều chỉnh lại:
    grade_col_widths = [Inches(0.4), Inches(0.8), Inches(2.0), Inches(0.5), Inches(0.6), Inches(0.6), Inches(0.7), Inches(0.6), Inches(0.8)] # ~7 inches
    if len(grade_col_widths) == num_grade_cols:
         for i, w in enumerate(grade_col_widths): grade_table.columns[i].width = w

    # Header bảng điểm
    header_cells = grade_table.rows[0].cells
    for i, header_text in enumerate(grade_headers):
        set_cell_properties(header_cells[i], header_text, bold=True, size=Pt(10))

    # Dữ liệu điểm
    for r, grade_row in enumerate(grades_data):
        row_cells = grade_table.rows[r + 1].cells
        set_cell_properties(row_cells[0], str(grade_row.get("stt", r+1)), size=Pt(10))
        set_cell_properties(row_cells[1], str(grade_row.get("ma_hp", "")), size=Pt(10), align=WD_ALIGN_PARAGRAPH.LEFT)
        set_cell_properties(row_cells[2], str(grade_row.get("ten_hp", "")), size=Pt(10), align=WD_ALIGN_PARAGRAPH.LEFT)
        set_cell_properties(row_cells[3], str(grade_row.get("tin_chi", "")), size=Pt(10))
        set_cell_properties(row_cells[4], str(grade_row.get("diem_qt", "")), size=Pt(10))
        set_cell_properties(row_cells[5], str(grade_row.get("diem_thi", "")), size=Pt(10))
        set_cell_properties(row_cells[6], str(grade_row.get("diem_tk", "")), size=Pt(10))
        set_cell_properties(row_cells[7], str(grade_row.get("diem_chu", "")), size=Pt(10))
        set_cell_properties(row_cells[8], str(grade_row.get("ghi_chu", "")), size=Pt(10), align=WD_ALIGN_PARAGRAPH.LEFT)

    # 5. Thông tin tổng kết
    document.add_paragraph() # Khoảng cách
    p_summary = document.add_paragraph()
    set_paragraph_format(p_summary, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=Pt(2), line_spacing=1.15)
    add_run_with_format(p_summary, "Tổng số tín chỉ đăng ký:", bold=True)
    add_run_with_format(p_summary, f" {summary['total_credits_registered']}\n")
    add_run_with_format(p_summary, "Tổng số tín chỉ tích lũy:", bold=True)
    add_run_with_format(p_summary, f" {summary['total_credits_earned']}\n")
    add_run_with_format(p_summary, f"Điểm trung bình học kỳ (hệ 4):", bold=True) # Cần quy đổi nếu cần
    add_run_with_format(p_summary, f" {summary['gpa_term']:.2f}\n") # Làm tròn 2 chữ số
    add_run_with_format(p_summary, f"Điểm trung bình tích lũy (hệ 4):", bold=True)
    add_run_with_format(p_summary, f" {summary['gpa_cumulative']:.2f}")


    # 6. Ngày tháng và chữ ký
    p_date_create = document.add_paragraph(f"{data.get('issuing_location', '........')}, ngày {time.strftime('%d')} tháng {time.strftime('%m')} năm {time.strftime('%Y')}")
    set_paragraph_format(p_date_create, alignment=WD_ALIGN_PARAGRAPH.RIGHT, space_before=Pt(12), space_after=Pt(0))
    add_run_with_format(p_date_create, p_date_create.text, size=FONT_SIZE_DEFAULT, italic=True)

    p_signer_title = document.add_paragraph()
    set_paragraph_format(p_signer_title, alignment=WD_ALIGN_PARAGRAPH.RIGHT, space_before=Pt(0), space_after=Pt(60))
    add_run_with_format(p_signer_title, data.get("signer_title", "TL. HIỆU TRƯỞNG\nTRƯỞNG PHÒNG ĐÀO TẠO").upper(), bold=True)

    p_signer_name = document.add_paragraph()
    set_paragraph_format(p_signer_name, alignment=WD_ALIGN_PARAGRAPH.RIGHT, space_before=Pt(0))
    add_run_with_format(p_signer_name, data.get("signer_name", "[Ký, ghi rõ họ tên]"), bold=True)


    print("Định dạng Bảng điểm hoàn tất.")