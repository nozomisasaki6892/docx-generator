# formatters/quyet_dinh.py
import re
import time
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from utils import set_paragraph_format, set_run_format, add_run_with_format
try:
    from .common_elements import format_basic_header, format_signature_block, format_recipient_list
except ImportError:
    def format_basic_header(document, data, doc_type): pass
    def format_signature_block(document, data): pass
    def format_recipient_list(document, data): pass

from config import FONT_SIZE_DEFAULT, FONT_SIZE_TITLE, FIRST_LINE_INDENT

def format(document, data):
    print("Bắt đầu định dạng Quyết định (cá biệt)...")
    title = data.get("title", "Quyết định về việc Thành lập/Phê duyệt/...")
    body = data.get("body", "Căn cứ...\nXét đề nghị của...\nQUYẾT ĐỊNH:\nĐiều 1...\nĐiều 2...")
    issuing_authority = data.get("issuing_org", "THỦ TRƯỞNG CƠ QUAN").upper() # VD: GIÁM ĐỐC, HIỆU TRƯỞNG
    doc_type_label = "QUYẾT ĐỊNH"

    # 1. Header (Cơ quan ban hành Quyết định)
    data['issuing_org'] = issuing_authority # Cập nhật data
    format_basic_header(document, data, "QuyetDinh")

    # 2. Tên loại và Trích yếu
    p_tenloai = document.add_paragraph(doc_type_label)
    set_paragraph_format(p_tenloai, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=Pt(12), space_after=Pt(6))
    add_run_with_format(p_tenloai, doc_type_label, size=FONT_SIZE_TITLE, bold=True, uppercase=True)

    trich_yeu_text = title.replace("Quyết định", "").strip()
    if trich_yeu_text.lower().startswith("về việc"):
        trich_yeu_text = trich_yeu_text.split(" ", 2)[-1]

    p_trichyeu = document.add_paragraph(f"Về việc {trich_yeu_text}")
    set_paragraph_format(p_trichyeu, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(12))
    add_run_with_format(p_trichyeu, f"Về việc {trich_yeu_text}", size=Pt(14), bold=True)


    # 3. Thẩm quyền ban hành (Tên chức vụ người ký)
    signer_title_upper = data.get("signer_title", "THỦ TRƯỞNG CƠ QUAN").upper()
    p_authority = document.add_paragraph(signer_title_upper)
    set_paragraph_format(p_authority, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(6))
    add_run_with_format(p_authority, signer_title_upper, size=FONT_SIZE_DEFAULT, bold=True, uppercase=True)


    # 4. Căn cứ ban hành
    body_lines = body.split('\n')
    processed_indices = set()

    for i, line in enumerate(body_lines):
        stripped_line = line.strip()
        if not stripped_line: continue
        if stripped_line.lower().startswith("căn cứ") or stripped_line.lower().startswith("xét đề nghị"):
            p = document.add_paragraph(stripped_line)
            set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line_indent=FIRST_LINE_INDENT, line_spacing=1.5, space_after=Pt(0))
            add_run_with_format(p.runs[0], size=FONT_SIZE_DEFAULT, italic=True)
            processed_indices.add(i)
        elif any(l.strip().lower().startswith("căn cứ") or l.strip().lower().startswith("xét đề nghị") for l in body_lines[:i]):
            break # Dừng khi hết căn cứ

    # 5. Phần QUYẾT ĐỊNH:
    added_qd_label = False
    quyet_dinh_intro_index = -1
    for i, line in enumerate(body_lines):
         if i in processed_indices: continue
         stripped_line = line.strip()
         if "QUYẾT ĐỊNH:" in stripped_line.upper():
             # Nếu dòng đó chỉ có QUYẾT ĐỊNH:
             if stripped_line.upper() == "QUYẾT ĐỊNH:":
                 p_qd_label = document.add_paragraph("QUYẾT ĐỊNH:")
                 set_paragraph_format(p_qd_label, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=Pt(12), space_after=Pt(12))
                 add_run_with_format(p_qd_label, "QUYẾT ĐỊNH:", size=FONT_SIZE_DEFAULT, bold=True, uppercase=True)
                 processed_indices.add(i)
                 added_qd_label = True
             else: # Nếu nằm trong câu khác (ít gặp)
                 p_qd_intro = document.add_paragraph(stripped_line)
                 set_paragraph_format(p_qd_intro, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=Pt(12), space_after=Pt(12))
                 add_run_with_format(p_qd_intro, stripped_line, size=FONT_SIZE_DEFAULT, bold=True)
                 processed_indices.add(i)
                 added_qd_label = True
             quyet_dinh_intro_index = i
             break

    if processed_indices and quyet_dinh_intro_index == -1:
         document.add_paragraph() # Khoảng trắng sau căn cứ

    # Thêm label nếu chưa có
    if not added_qd_label:
        p_qd_label = document.add_paragraph("QUYẾT ĐỊNH:")
        set_paragraph_format(p_qd_label, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=Pt(12), space_after=Pt(12))
        add_run_with_format(p_qd_label, "QUYẾT ĐỊNH:", size=FONT_SIZE_DEFAULT, bold=True, uppercase=True)


    # 6. Nội dung Quyết định (Điều, Khoản, Điểm)
    for i, line in enumerate(body_lines):
        if i in processed_indices: continue
        stripped_line = line.strip()
        if not stripped_line: continue

        # Bỏ qua label QUYẾT ĐỊNH: nếu đã thêm
        if "QUYẾT ĐỊNH:" in stripped_line.upper() and not added_qd_label:
            continue

        p = document.add_paragraph()
        is_dieu = stripped_line.upper().startswith("ĐIỀU")
        is_khoan = re.match(r'^\d+\.\s+', stripped_line)
        is_diem = re.match(r'^[a-z]\)\s+', stripped_line)

        align = WD_ALIGN_PARAGRAPH.JUSTIFY
        left_indent = Cm(0)
        first_indent = FIRST_LINE_INDENT if not (is_dieu or is_khoan or is_diem) else Cm(0)
        is_bold = False
        size = FONT_SIZE_DEFAULT

        if is_dieu:
            align = WD_ALIGN_PARAGRAPH.LEFT
            is_bold = True
            space_before = Pt(6)
        elif is_khoan:
            left_indent = Cm(0.5)
        elif is_diem:
            left_indent = Cm(1.0)

        set_paragraph_format(p, alignment=align, left_indent=left_indent, first_line_indent=first_indent, line_spacing=1.5, space_after=Pt(6))
        add_run_with_format(p, stripped_line, size=size, bold=is_bold)

    # 7. Chữ ký
    # Chức vụ đã có ở mục 3, giờ chỉ cần tên
    if 'signer_title' not in data: data['signer_title'] = signer_title_upper # Dùng lại chức vụ ở mục 3
    if 'signer_name' not in data: data['signer_name'] = "[Họ và tên]"
    format_signature_block(document, data) # Dùng signature block chuẩn

    # 8. Nơi nhận
    format_recipient_list(document, data)

    print("Định dạng Quyết định (cá biệt) hoàn tất.")