# formatters/hop_dong.py (Đã sửa SyntaxError)
import re
import time
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from utils import set_paragraph_format, set_run_format, add_run_with_format, add_centered_text
# Hợp đồng dùng lại format parties và signature của Bản ghi nhớ/Thỏa thuận
try:
    # Cần đảm bảo import đúng hàm đã sửa lỗi indentation trước đó nếu có
    from .ban_ghi_nho import format_parties_block as format_parties_block_mou, format_signature_mou
except ImportError:
    # Fallback với định nghĩa tạm thời để tránh lỗi import nếu file kia có vấn đề
    print("Warning: Could not import from ban_ghi_nho. Using placeholder functions.")
    def format_parties_block_mou(document, parties): pass
    def format_signature_mou(document, signer_a_title, signer_a_name, signer_b_title, signer_b_name): pass

from config import FONT_SIZE_DEFAULT, FONT_SIZE_TITLE, FONT_SIZE_HEADER, FIRST_LINE_INDENT, FONT_SIZE_PLACE_DATE, FONT_SIZE_SIGNATURE, FONT_SIZE_SIGNER_NAME # Thêm các import cần thiết

# Định nghĩa hàm format_party_info ngay trong file này để tránh phụ thuộc vòng
def format_party_info(document, party_label, party_data):
    p_label = document.add_paragraph(party_label)
    # --- Sửa lỗi thiếu dấu ) ở dòng dưới ---
    set_paragraph_format(p_label, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=Pt(6), first_line_indent=Cm(0)) # Thêm dấu )
    set_run_format(p_label.runs[0], size=FONT_SIZE_DEFAULT, bold=True)

    default_info = {
        "name": "[Tên đầy đủ]", "address": "[Địa chỉ]", "phone": "[Điện thoại]",
        "fax": "[Fax]", "tax_code": "[Mã số thuế]", "account_number": "[Số tài khoản]",
        "bank_name": "[Ngân hàng]", "representative": "[Người đại diện]", "position": "[Chức vụ]"
    }
    # Các trường thông tin phổ biến của Doanh nghiệp/Tổ chức
    for key, label in [("name", "Tên đơn vị"), ("address", "Địa chỉ trụ sở"), ("phone", "Điện thoại"),
                       ("fax", "Fax"), ("tax_code", "Mã số thuế"), ("account_number", "Số tài khoản"),
                       ("bank_name", "Tại Ngân hàng"), ("representative", "Người đại diện"), ("position", "Chức vụ")]:
         value = party_data.get(key)
         if value: # Chỉ thêm nếu có giá trị
             p_info = document.add_paragraph()
             set_paragraph_format(p_info, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=Pt(3), left_indent=Cm(0.5), first_line_indent=Cm(0))
             add_run_with_format(p_info, f"- {label}: ", size=FONT_SIZE_DEFAULT)
             add_run_with_format(p_info, str(value), size=FONT_SIZE_DEFAULT) # Đảm bảo value là string


def format(document, data):
    print("Bắt đầu định dạng Hợp đồng...")
    title = data.get("title", "HỢP ĐỒNG KINH TẾ").upper() # Hoặc Hợp đồng Dịch vụ, Mua bán...
    contract_number = data.get("contract_number", "Số: ...... /HĐKT")
    issuing_location = data.get("issuing_location", "Hà Nội")
    current_date_str = time.strftime(f"ngày %d tháng %m năm %Y")
    preamble = data.get("preamble", ["Căn cứ Bộ luật Dân sự số 91/2015/QH13 ngày 24 tháng 11 năm 2015;", "Căn cứ Luật Thương mại số 36/2005/QH11 ngày 14 tháng 6 năm 2005;", "Căn cứ vào khả năng và nhu cầu của hai Bên."])
    party_a_data = data.get("party_a", {}) # Dữ liệu Bên A (dict)
    party_b_data = data.get("party_b", {}) # Dữ liệu Bên B (dict)
    body = data.get("body", "Điều 1: Đối tượng hợp đồng\nĐiều 2: Giá trị hợp đồng và phương thức thanh toán\nĐiều 3: Quyền và nghĩa vụ của Bên A...")


    # 1. Quốc hiệu, Tiêu ngữ
    add_centered_text(document, "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM", size=FONT_SIZE_HEADER, bold=True, space_after=0)
    add_centered_text(document, "Độc lập - Tự do - Hạnh phúc", size=Pt(13), bold=True, space_after=18)

    # 2. Tên Hợp đồng
    add_centered_text(document, title, size=FONT_SIZE_TITLE, bold=True, space_before=12, space_after=6)
    # Số hợp đồng
    add_centered_text(document, contract_number, size=Pt(13), bold=False, space_after=18)


    # 3. Căn cứ ký kết
    if isinstance(preamble, str): preamble = preamble.split('\n')
    if preamble:
        for line in preamble:
            p_pre = document.add_paragraph(line)
            set_paragraph_format(p_pre, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=Pt(0), first_line_indent=FIRST_LINE_INDENT, line_spacing=1.5)
            set_run_format(p_pre.runs[0], size=FONT_SIZE_DEFAULT, italic=True)
        document.add_paragraph()

    # 4. Ngày tháng, địa điểm ký
    p_intro = document.add_paragraph(f"Hôm nay, {current_date_str}, tại {issuing_location}, chúng tôi gồm có:")
    set_paragraph_format(p_intro, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=Pt(12), line_spacing=1.5, first_line_indent=Cm(0))
    set_run_format(p_intro.runs[0], size=FONT_SIZE_DEFAULT)


    # 5. Thông tin các bên (Dùng hàm helper nội bộ)
    format_party_info(document, "BÊN A:", party_a_data)
    document.add_paragraph()
    format_party_info(document, "BÊN B:", party_b_data)
    document.add_paragraph()

    p_agree = document.add_paragraph("Sau khi bàn bạc, hai bên thống nhất ký kết Hợp đồng này với các điều khoản sau đây:")
    set_paragraph_format(p_agree, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=Pt(6), first_line_indent=FIRST_LINE_INDENT, line_spacing=1.5)
    set_run_format(p_agree.runs[0], size=FONT_SIZE_DEFAULT)


    # 6. Nội dung Hợp đồng (Điều khoản)
    body_lines = body.split('\n')
    for line in body_lines:
        stripped_line = line.strip()
        if not stripped_line: continue
        p = document.add_paragraph()
        is_dieu = stripped_line.upper().startswith("ĐIỀU")
        is_khoan = re.match(r'^\d+\.\s+', stripped_line) # 1. 2.
        is_sub_khoan = re.match(r'^\d+\.\d+\.?\s+', stripped_line) # 1.1, 1.2
        is_diem = re.match(r'^[a-z]\)\s+', stripped_line) # a) b)

        align = WD_ALIGN_PARAGRAPH.JUSTIFY
        left_indent = Cm(0)
        first_indent = FIRST_LINE_INDENT if not (is_dieu or is_khoan or is_sub_khoan or is_diem) else Cm(0)
        is_bold = bool(is_dieu or is_khoan) # Điều, khoản cấp 1 đậm
        size = FONT_SIZE_DEFAULT
        space_before = Pt(0)
        space_after = Pt(6)
        line_spacing = 1.5

        if is_dieu:
            align = WD_ALIGN_PARAGRAPH.CENTER # Điều khoản thường căn giữa
            is_bold = True
            space_before = Pt(12)
        elif is_khoan:
            align = WD_ALIGN_PARAGRAPH.LEFT
            left_indent = Cm(0.5)
            is_bold = True # Khoản 1. 2. đậm
            space_before = Pt(6)
        elif is_sub_khoan:
             align = WD_ALIGN_PARAGRAPH.LEFT
             left_indent = Cm(1.0)
             is_bold = False # Mục con nhỏ không đậm
        elif is_diem:
            align = WD_ALIGN_PARAGRAPH.LEFT
            left_indent = Cm(1.5)
            is_bold = False


        set_paragraph_format(p, alignment=align, left_indent=left_indent, first_line_indent=first_indent, line_spacing=line_spacing, space_before=space_before, space_after=space_after)
        add_run_with_format(p, stripped_line, size=size, bold=is_bold)


    # 7. Điều khoản chung (Hiệu lực, giải quyết tranh chấp...)
    # Thường là các Điều cuối cùng, xử lý như trên

    p_validity = document.add_paragraph("Hợp đồng này được lập thành 02 (hai) bản có giá trị pháp lý như nhau, mỗi bên giữ 01 (một) bản và có hiệu lực kể từ ngày ký.")
    set_paragraph_format(p_validity, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_before=Pt(18), space_after=Pt(18), first_line_indent=FIRST_LINE_INDENT, line_spacing=1.5)
    set_run_format(p_validity.runs[0], size=FONT_SIZE_DEFAULT)


    # 8. Chữ ký hai bên (Dùng hàm helper từ Bản ghi nhớ)
    signer_a_title = party_a_data.get("signer_title", "ĐẠI DIỆN BÊN A").upper()
    signer_a_name = party_a_data.get("signer_name", "[Tên người ký Bên A]")
    signer_b_title = party_b_data.get("signer_title", "ĐẠI DIỆN BÊN B").upper()
    signer_b_name = party_b_data.get("signer_name", "[Tên người ký Bên B]")
    format_signature_mou(document, signer_a_title, signer_a_name, signer_b_title, signer_b_name)


    print("Định dạng Hợp đồng hoàn tất.")