# formatters/quyet_dinh_ttg.py
import re
import time
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
# Header QPPL, Chữ ký Thủ tướng
try:
    from .luat import format_qppl_header
except ImportError:
    def format_qppl_header(document, issuer_name): pass
# Signature TTg giống NĐ Chính phủ
try:
    from .nghi_dinh_qppl import format_cp_signature
except ImportError:
    def format_cp_signature(document, signer_title, signer_name): pass

from utils import set_paragraph_format, set_run_format, add_run_with_format
from config import FONT_SIZE_DEFAULT, FONT_SIZE_TITLE, FIRST_LINE_INDENT

def format(document, data):
    print("Bắt đầu định dạng Quyết định Thủ tướng (QPPL)...")
    title = data.get("title", "Quyết định Phê duyệt Chiến lược/Đề án Quốc gia...").upper()
    body = data.get("body", "Căn cứ Luật...\nXét đề nghị của Bộ trưởng...\nQUYẾT ĐỊNH:\nĐiều 1...\nĐiều 2...")
    decision_number = data.get("decision_number", "Số: .../QĐ-TTg")
    issuing_date_str = data.get("issuing_date", time.strftime("ngày %d tháng %m năm %Y"))
    issuing_location = data.get("issuing_location", "Hà Nội")
    issuer_name = "THỦ TƯỚNG CHÍNH PHỦ"


    # 1. Header (THỦ TƯỚNG CHÍNH PHỦ - Kiểu QPPL)
    format_qppl_header(document, issuer_name)


    # 2. Số hiệu và Ngày tháng ban hành (Căn giữa)
    p_num_date = document.add_paragraph(f"{decision_number}       {issuing_location}, {issuing_date_str}")
    set_paragraph_format(p_num_date, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(12))
    run_num = add_run_with_format(p_num_date, f"{decision_number}       ", size=FONT_SIZE_DEFAULT)
    run_date = add_run_with_format(p_num_date, f"{issuing_location}, {issuing_date_str}", size=FONT_SIZE_DEFAULT, italic=True)


    # 3. Tên Quyết định (IN HOA, đậm, căn giữa)
    p_tenloai = document.add_paragraph(title)
    set_paragraph_format(p_tenloai, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=Pt(6), space_after=Pt(12))
    add_run_with_format(p_tenloai, p_tenloai.text, size=FONT_SIZE_TITLE, bold=True, uppercase=True)


    # 4. Cơ quan ban hành (THỦ TƯỚNG CHÍNH PHỦ - Lặp lại)
    p_issuer_body = document.add_paragraph(issuer_name.upper())
    set_paragraph_format(p_issuer_body, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(6))
    add_run_with_format(p_issuer_body, p_issuer_body.text, size=FONT_SIZE_DEFAULT, bold=True, uppercase=True)


    # 5. Căn cứ ban hành
    body_lines = body.split('\n')
    processed_indices = set()

    for i, line in enumerate(body_lines):
        stripped_line = line.strip()
        if not stripped_line: continue
        if stripped_line.lower().startswith("căn cứ") or "xét đề nghị của" in stripped_line.lower():
            p = document.add_paragraph(stripped_line)
            set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line_indent=FIRST_LINE_INDENT, line_spacing=1.5, space_after=Pt(0))
            set_run_format(p.runs[0], size=FONT_SIZE_DEFAULT, italic=True)
            processed_indices.add(i)
        # Dừng khi hết căn cứ hoặc gặp QUYẾT ĐỊNH:
        elif any(l.strip().lower().startswith("căn cứ") or "xét đề nghị của" in l.strip().lower() for l in body_lines[:i]):
             if "QUYẾT ĐỊNH:" in stripped_line.upper() or stripped_line.upper().startswith("ĐIỀU"):
                 break

    # 6. Phần QUYẾT ĐỊNH:
    added_qd_label = False
    for i, line in enumerate(body_lines):
        if i in processed_indices: continue
        stripped_line = line.strip()
        if "QUYẾT ĐỊNH:" in stripped_line.upper():
             if stripped_line.upper() == "QUYẾT ĐỊNH:":
                 p_qd_label = document.add_paragraph("QUYẾT ĐỊNH:")
                 set_paragraph_format(p_qd_label, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=Pt(12), space_after=Pt(12))
                 add_run_with_format(p_qd_label, "QUYẾT ĐỊNH:", size=FONT_SIZE_DEFAULT, bold=True, uppercase=True)
                 processed_indices.add(i)
                 added_qd_label = True
             else: # Ít gặp
                 p_qd_intro = document.add_paragraph(stripped_line)
                 set_paragraph_format(p_qd_intro, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=Pt(12), space_after=Pt(12))
                 add_run_with_format(p_qd_intro, stripped_line, size=FONT_SIZE_DEFAULT, bold=True)
                 processed_indices.add(i)
                 added_qd_label = True
             break # Chỉ xử lý 1 lần

    if processed_indices and not added_qd_label: document.add_paragraph() # Khoảng cách

    if not added_qd_label:
        p_qd_label = document.add_paragraph("QUYẾT ĐỊNH:")
        set_paragraph_format(p_qd_label, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=Pt(12), space_after=Pt(12))
        add_run_with_format(p_qd_label, "QUYẾT ĐỊNH:", size=FONT_SIZE_DEFAULT, bold=True, uppercase=True)

    # 7. Nội dung (Điều, Khoản, Điểm)
    for i, line in enumerate(body_lines):
        if i in processed_indices: continue
        stripped_line = line.strip()
        if not stripped_line: continue
        # Bỏ qua label QUYẾT ĐỊNH: nếu đã thêm
        if "QUYẾT ĐỊNH:" in stripped_line.upper() and not added_qd_label:
            continue

        p = document.add_paragraph()
        is_dieu = stripped_line.upper().startswith("ĐIỀU")
        is_khoan = re.match(r'^\d+\.\s+', stripped_line)
        is_diem = re.match(r'^[a-z]\)\s+', stripped_line)

        align = WD_ALIGN_PARAGRAPH.JUSTIFY
        left_indent = Cm(0)
        first_indent = FIRST_LINE_INDENT if not (is_dieu or is_khoan or is_diem) else Cm(0)
        is_bold = False
        size = FONT_SIZE_DEFAULT
        space_before = Pt(0)
        space_after = Pt(6)
        line_spacing = 1.5

        if is_dieu:
            align = WD_ALIGN_PARAGRAPH.LEFT
            is_bold = True
            space_before = Pt(6)
            size = Pt(13)
        elif is_khoan:
            left_indent = Cm(0.5)
        elif is_diem:
            left_indent = Cm(1.0)

        set_paragraph_format(p, alignment=align, left_indent=left_indent, first_line_indent=first_indent, line_spacing=line_spacing, space_before=space_before, space_after=space_after)
        add_run_with_format(p, stripped_line, size=size, bold=is_bold)


    # 8. Chữ ký (THỦ TƯỚNG) - Không có TM.
    signer_title = data.get("signer_title", "THỦ TƯỚNG")
    signer_name = data.get("signer_name", "[Tên Thủ tướng]")
    sig_paragraph = document.add_paragraph()
    set_paragraph_format(sig_paragraph, alignment=WD_ALIGN_PARAGRAPH.RIGHT, space_before=Pt(12), space_after=Pt(0), line_spacing=1.15)
    add_run_with_format(sig_paragraph, signer_title.upper() + "\n\n\n\n\n", size=FONT_SIZE_SIGNATURE, bold=True)
    add_run_with_format(sig_paragraph, signer_name, size=FONT_SIZE_SIGNER_NAME, bold=True)


    # Quyết định TTg thường không có nơi nhận ở cuối

    print("Định dạng Quyết định Thủ tướng (QPPL) hoàn tất.")