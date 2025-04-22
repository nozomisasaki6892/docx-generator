# formatters/bao_cao.py (Đã sửa IndentationError)
import re # Thêm import re nếu chưa có
import time
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from utils import set_paragraph_format, set_run_format, add_run_with_format
try:
    from .common_elements import format_basic_header, format_signature_block, format_recipient_list
except ImportError:
    # Fallback nếu chạy độc lập
    def format_basic_header(document, data, doc_type): pass
    def format_signature_block(document, data): pass
    def format_recipient_list(document, data): pass

from config import FONT_SIZE_DEFAULT, FONT_SIZE_TITLE, FIRST_LINE_INDENT, FONT_SIZE_SMALL # Thêm FONT_SIZE_SMALL

def format(document, data):
    print("Bắt đầu định dạng Báo cáo...")
    title = data.get("title", "Báo cáo Công tác Tháng/Quý/Năm")
    body = data.get("body", "Phần I: Tình hình thực hiện...\nPhần II: Phương hướng...")
    doc_type_label = "BÁO CÁO"

    # 1. Header (Sử dụng header cơ bản)
    format_basic_header(document, data, "BaoCao")

    # 2. Tiêu đề
    p_title = document.add_paragraph(doc_type_label)
    set_paragraph_format(p_title, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=Pt(12), space_after=Pt(6))
    set_run_format(p_title.runs[0], size=FONT_SIZE_TITLE, bold=True)

    # Trích yếu nội dung báo cáo
    report_subject = title.replace("Báo cáo", "").strip()
    # Bỏ V/v:, về việc nếu có trong title
    if report_subject.lower().startswith("v/v:"):
        report_subject = report_subject[4:].strip()
    elif report_subject.lower().startswith("về việc"):
        report_subject = report_subject.split(" ", 2)[-1]

    p_subject = document.add_paragraph(f"V/v: {report_subject}")
    set_paragraph_format(p_subject, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(6))
    set_run_format(p_subject.runs[0], size=Pt(14), bold=True) # Có thể không đậm tùy yêu cầu
    p_line_sub = document.add_paragraph("-" * 15)
    set_paragraph_format(p_line_sub, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(12))


    # 3. Kính gửi (Nếu có)
    # --- Sửa lỗi thụt lề ở dòng dưới ---
    kinh_gui = data.get("kinh_gui") # Sửa lại key và bỏ khoảng trắng thừa đầu dòng
    if kinh_gui:
        p_kg = document.add_paragraph(f"Kính gửi: {kinh_gui}")
        set_paragraph_format(p_kg, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=Pt(12), first_line_indent=Cm(0)) # Kính gửi không thụt đầu dòng
        set_run_format(p_kg.runs[0], size=FONT_SIZE_DEFAULT, bold=True)


    # 4. Nội dung báo cáo
    body_lines = body.split('\n')
    for line in body_lines:
        stripped_line = line.strip()
        if not stripped_line: continue
        p = document.add_paragraph()

        # Logic định dạng cơ bản cho các đề mục phổ biến trong báo cáo
        is_part_roman = re.match(r'^(PHẦN\s+[IVXLCDM]+)\.?\s+', stripped_line.upper())
        is_part_digit = re.match(r'^([A-Z]\.)\s+', stripped_line) # Mục A, B, C
        is_section_digit = re.match(r'^(\d+\.)\s+', stripped_line) # Mục 1, 2, 3
        is_subsection_alpha = re.match(r'^[a-z]\)\s+', stripped_line) # Mục a, b, c
        is_bullet = stripped_line.startswith("-") or stripped_line.startswith("+") or stripped_line.startswith("*")

        align = WD_ALIGN_PARAGRAPH.JUSTIFY
        left_indent = Cm(0)
        first_indent = FIRST_LINE_INDENT if not (is_part_roman or is_part_digit or is_section_digit or is_subsection_alpha or is_bullet) else Cm(0)
        is_bold = bool(is_part_roman or is_part_digit or is_section_digit)
        is_italic = False
        size = FONT_SIZE_DEFAULT
        space_before = Pt(0)
        space_after = Pt(6)
        line_spacing = 1.5

        if is_part_roman or is_part_digit:
            align = WD_ALIGN_PARAGRAPH.CENTER
            space_before = Pt(12)
            size = Pt(13) # Cỡ chữ to hơn một chút
        elif is_section_digit:
            align = WD_ALIGN_PARAGRAPH.LEFT
            left_indent = Cm(0)
            space_before = Pt(6)
        elif is_subsection_alpha:
            align = WD_ALIGN_PARAGRAPH.LEFT
            left_indent = Cm(0.5)
        elif is_bullet:
            align = WD_ALIGN_PARAGRAPH.LEFT
            left_indent = Cm(1.0) # Điều chỉnh thụt lề bullet nếu cần
            first_indent = Cm(-0.5) # Hanging indent for bullets


        set_paragraph_format(p, alignment=align, left_indent=left_indent, first_line_indent=first_indent, line_spacing=line_spacing, space_before=space_before, space_after=space_after)
        add_run_with_format(p, stripped_line, size=size, bold=is_bold, italic=is_italic)


    # 5. Chữ ký (Sử dụng khối chữ ký cơ bản)
    # Cần xác định rõ người ký và chức vụ từ data
    if 'signer_title' not in data: data['signer_title'] = "NGƯỜI LẬP BÁO CÁO" # Hoặc chức vụ cụ thể
    if 'signer_name' not in data: data['signer_name'] = "[Họ và tên]"
    document.add_paragraph() # Khoảng cách trước chữ ký
    format_signature_block(document, data)

    # 6. Nơi nhận (Sử dụng nơi nhận cơ bản)
    if 'recipients' not in data: data['recipients'] = ["- Như kính gửi (nếu có);", "- Lưu: VT, ...;"]
    format_recipient_list(document, data)

    print("Định dạng Báo cáo hoàn tất.")