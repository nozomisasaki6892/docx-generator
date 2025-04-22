# formatters/quy_dinh_nt.py
import re
import time
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from utils import set_paragraph_format, set_run_format, add_run_with_format
try:
    # Dùng header, signature như Quy định, Quy chế
    from .common_elements import format_basic_header, format_signature_block, format_recipient_list
except ImportError:
    def format_basic_header(document, data, doc_type): pass
    def format_signature_block(document, data): pass
    def format_recipient_list(document, data): pass

from config import FONT_SIZE_DEFAULT, FONT_SIZE_TITLE, FIRST_LINE_INDENT

def format(document, data):
    print("Bắt đầu định dạng Quy định Nhà trường...")
    title = data.get("title", "Quy định về Tác phong Sinh viên/Nội quy KTX...")
    # Thường ban hành kèm Quyết định của Hiệu trưởng
    enactment_info = data.get("enactment_info", "Ban hành kèm theo Quyết định số .../QĐ-ĐH... ngày ... tháng ... năm ... của Hiệu trưởng")
    body = data.get("body", "Chương I: QUY ĐỊNH CHUNG\nĐiều 1. Phạm vi và đối tượng...\nChương II: QUY ĐỊNH CỤ THỂ\nĐiều 2...")
    doc_type_label = "QUY ĐỊNH"


    # 1. Header (Tên Trường)
    if 'issuing_org' not in data: data['issuing_org'] = "TRƯỜNG ĐẠI HỌC XYZ"
    format_basic_header(document, data, "QuyDinhNT")


    # 2. Tên Quy định
    p_title = document.add_paragraph(doc_type_label)
    set_paragraph_format(p_title, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=Pt(12), space_after=Pt(6))
    set_run_format(p_title.runs[0], size=FONT_SIZE_TITLE, bold=True)

    # Tên Quy định cụ thể
    rule_name = title.replace("Quy định", "").strip()
    p_name = document.add_paragraph(rule_name)
    set_paragraph_format(p_name, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(6))
    set_run_format(p_name.runs[0], size=Pt(14), bold=True) # Tên QĐ đậm

    # Thông tin ban hành kèm theo
    p_enactment = document.add_paragraph(f"({enactment_info})")
    set_paragraph_format(p_enactment, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(12))
    set_run_format(p_enactment.runs[0], size=FONT_SIZE_DEFAULT, italic=True) # Nghiêng


    # 3. Nội dung Quy định (Chương, Điều, Khoản, Điểm)
    body_lines = body.split('\n')
    for line in body_lines:
        stripped_line = line.strip()
        if not stripped_line: continue
        p = document.add_paragraph()

        # Logic định dạng cơ bản (Tương tự Quy chế)
        is_chuong = stripped_line.upper().startswith("CHƯƠNG")
        is_dieu = stripped_line.upper().startswith("ĐIỀU")
        is_khoan = re.match(r'^\d+\.\s+', stripped_line)
        is_diem = re.match(r'^[a-z]\)\s+', stripped_line)

        align = WD_ALIGN_PARAGRAPH.JUSTIFY
        left_indent = Cm(0)
        first_indent = FIRST_LINE_INDENT if not (is_chuong or is_dieu or is_khoan or is_diem) else Cm(0)
        is_bold = False
        size = FONT_SIZE_DEFAULT
        space_before = Pt(0)
        space_after = Pt(6)
        line_spacing = 1.5

        if is_chuong:
            align = WD_ALIGN_PARAGRAPH.CENTER
            is_bold = True
            space_before = Pt(12)
            size = Pt(13)
        elif is_dieu:
            align = WD_ALIGN_PARAGRAPH.LEFT
            is_bold = True
            space_before = Pt(6)
            size = Pt(13)
        elif is_khoan:
            left_indent = Cm(0.5)
        elif is_diem:
            left_indent = Cm(1.0)

        set_paragraph_format(p, alignment=align, left_indent=left_indent, first_line_indent=first_indent, line_spacing=line_spacing, space_before=space_before, space_after=space_after)
        add_run_with_format(p, stripped_line, size=size, bold=is_bold)


    # 4. Chữ ký (Người ký Quyết định ban hành Quy định)
    # Thường không có trên Quy định kèm theo
    if data.get('signer_title') and data.get('signer_name'):
         document.add_paragraph()
         format_signature_block(document, data)
         # format_recipient_list(document, data)

    print("Định dạng Quy định Nhà trường hoàn tất.")