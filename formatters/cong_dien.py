# formatters/cong_dien.py
import time
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from utils import set_paragraph_format, set_run_format, add_run_with_format
try:
    # Công điện dùng header và signature riêng
    from .common_elements import format_recipient_list # Chỉ dùng nơi nhận
except ImportError:
    def format_recipient_list(document, data): pass

from config import FONT_SIZE_DEFAULT, FONT_SIZE_TITLE, FIRST_LINE_INDENT, FONT_SIZE_SIGNATURE, FONT_SIZE_SIGNER_NAME

def format_cong_dien_header(document, data):
    # Header công điện thường đơn giản hơn
    issuing_org = data.get("issuing_org", "TÊN CƠ QUAN GỬI ĐIỆN").upper()
    doc_number = data.get("doc_number", "Số:       /CĐ-...")

    p_org = document.add_paragraph(issuing_org)
    set_paragraph_format(p_org, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(0))
    set_run_format(p_org.runs[0], size=Pt(13), bold=True)

    p_line_org = document.add_paragraph("*******")
    set_paragraph_format(p_line_org, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(6))

    p_num = document.add_paragraph(doc_number)
    set_paragraph_format(p_num, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(12))
    set_run_format(p_num.runs[0], size=Pt(13))

def format(document, data):
    print("Bắt đầu định dạng Công điện...")
    title = data.get("title", "Công điện về việc Khẩn trương ứng phó XYZ")
    body = data.get("body", "Nội dung công điện yêu cầu...")
    issuing_location = data.get("issuing_location", "Hà Nội")
    current_date_str = time.strftime(f"%H giờ %M, ngày %d tháng %m năm %Y") # Công điện có giờ, phút

    # 1. Header Công điện
    format_cong_dien_header(document, data)

    # 2. Tiêu đề CÔNG ĐIỆN
    p_title = document.add_paragraph("CÔNG ĐIỆN")
    set_paragraph_format(p_title, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=Pt(12), space_after=Pt(12))
    set_run_format(p_title.runs[0], size=FONT_SIZE_TITLE, bold=True)

    # 3. Nơi gửi điện (Kính gửi)
    recipients_direct = data.get("recipients_direct", ["Ban Chỉ huy PCTT và TKCN các tỉnh, thành phố...", "Các Bộ, ngành liên quan..."])
    if recipients_direct:
        p_kg_label = document.add_paragraph("Kính gửi:")
        set_paragraph_format(p_kg_label, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(6))
        set_run_format(p_kg_label.runs[0], size=FONT_SIZE_DEFAULT, bold=True)
        for recipient in recipients_direct:
             p_rec = document.add_paragraph(recipient)
             set_paragraph_format(p_rec, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(0))
             set_run_format(p_rec.runs[0], size=FONT_SIZE_DEFAULT, bold=True) # Tên nơi nhận đậm
        document.add_paragraph()

    # 4. Nội dung công điện
    # Thường viết liền mạch, không chia điều khoản phức tạp
    body_lines = body.split('\n')
    for line in body_lines:
        stripped_line = line.strip()
        if not stripped_line: continue
        p = document.add_paragraph(stripped_line)
        set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=Pt(6), first_line_indent=FIRST_LINE_INDENT, line_spacing=1.5)
        set_run_format(p.runs[0], size=FONT_SIZE_DEFAULT)


    # 5. Thời gian ban hành (có giờ, phút) - Căn phải
    p_datetime = document.add_paragraph(f"{issuing_location}, {current_date_str}")
    set_paragraph_format(p_datetime, alignment=WD_ALIGN_PARAGRAPH.RIGHT, space_before=Pt(12), space_after=Pt(6))
    set_run_format(p_datetime.runs[0], size=Pt(13), italic=True)


    # 6. Chữ ký (Căn phải, giống common_elements nhưng không có thẩm quyền)
    signer_title = data.get("signer_title", "CHỨC VỤ NGƯỜI KÝ").upper()
    signer_name = data.get("signer_name", "Người Ký")

    sig_paragraph = document.add_paragraph()
    set_paragraph_format(sig_paragraph, alignment=WD_ALIGN_PARAGRAPH.RIGHT, space_before=Pt(0), space_after=Pt(0), line_spacing=1.0)
    add_run_with_format(sig_paragraph, signer_title + "\n\n\n\n\n", size=FONT_SIZE_SIGNATURE, bold=True)
    add_run_with_format(sig_paragraph, signer_name, size=FONT_SIZE_SIGNER_NAME, bold=True)

    # 7. Nơi nhận (Dùng format_recipient_list)
    # Thường không có "Như trên" mà liệt kê cụ thể
    if 'recipients' not in data: data['recipients'] = ["- Như kính gửi;", "- Thủ tướng Chính phủ (để b/c);", "- Lưu: VT, ...;"]
    # Thêm ngắt trang trước nơi nhận nếu cần
    # document.add_page_break()
    format_recipient_list(document, data)


    print("Định dạng Công điện hoàn tất.")