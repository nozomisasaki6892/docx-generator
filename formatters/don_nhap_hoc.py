# formatters/don_nhap_hoc.py
import time
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from utils import set_paragraph_format, set_run_format, add_run_with_format, add_centered_text
from config import FONT_SIZE_DEFAULT, FONT_SIZE_TITLE, FIRST_LINE_INDENT, FONT_SIZE_HEADER

def format(document, data):
    print("Bắt đầu định dạng Đơn xin nhập học...")
    title = data.get("title", "ĐƠN XIN NHẬP HỌC").upper()
    body = data.get("body", "Kính gửi:...\nTên em là:...\nNgày sinh:...\nTrúng tuyển ngành:...\nĐiểm thi:...\nEm xin trình bày nguyện vọng...")
    student_name = data.get("student_name", "[Họ và tên sinh viên]")
    issuing_location = data.get("issuing_location", "Hà Nội") # Nơi viết đơn
    submission_date_str = data.get("submission_date", time.strftime(f"ngày %d tháng %m năm %Y")) # Ngày viết đơn


    # 1. Quốc hiệu, Tiêu ngữ (Căn giữa)
    add_centered_text(document, "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM", size=FONT_SIZE_HEADER, bold=True, space_after=0)
    add_centered_text(document, "Độc lập - Tự do - Hạnh phúc", size=Pt(13), bold=True, space_after=18)

    # 2. Tên Đơn
    add_centered_text(document, title, size=FONT_SIZE_TITLE, bold=True, space_before=12, space_after=12)

    # 3. Kính gửi
    recipient = data.get("recipient", "Ban Giám hiệu Trường [Tên trường]\nvà Phòng Công tác Sinh viên") # VD
    p_kg = document.add_paragraph(f"Kính gửi: {recipient}")
    set_paragraph_format(p_kg, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=Pt(12))
    set_run_format(p_kg.runs[0], size=FONT_SIZE_DEFAULT, bold=True)

    # 4. Nội dung Đơn
    body_lines = body.split('\n')
    for line in body_lines:
        stripped_line = line.strip()
        if not stripped_line: continue

        p = document.add_paragraph()
        # Nội dung đơn thường căn trái hoặc đều, thụt lề dòng đầu
        # Xử lý riêng các dòng khai thông tin
        is_info_line = ":" in stripped_line and len(stripped_line.split(":")[0]) < 30 # Heuristic for info lines

        align = WD_ALIGN_PARAGRAPH.JUSTIFY
        first_indent = FIRST_LINE_INDENT
        is_bold = False
        if is_info_line:
            align = WD_ALIGN_PARAGRAPH.LEFT
            first_indent = Cm(0)

        set_paragraph_format(p, alignment=align, space_after=Pt(6), first_line_indent=first_indent, line_spacing=1.5)

        if is_info_line:
            parts = stripped_line.split(":", 1)
            add_run_with_format(p, parts[0] + ":", size=FONT_SIZE_DEFAULT, bold=False) # Nhãn thông tin thường
            add_run_with_format(p, parts[1], size=FONT_SIZE_DEFAULT, bold=False) # Nội dung thông tin thường
        else:
             add_run_with_format(p, stripped_line, size=FONT_SIZE_DEFAULT)


    # 5. Lời cam đoan / Cảm ơn
    p_thanks = document.add_paragraph("Em xin chân thành cảm ơn!")
    set_paragraph_format(p_thanks, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=Pt(12), space_after=Pt(12), first_line_indent=FIRST_LINE_INDENT, line_spacing=1.5)
    set_run_format(p_thanks.runs[0], size=FONT_SIZE_DEFAULT, italic=False) # Thường không nghiêng

    # 6. Ngày tháng và Chữ ký người làm đơn (Căn phải)
    table = document.add_table(rows=1, cols=2)
    table.autofit = False
    table.allow_autofit = False
    table.columns[0].width = Cm(8.0) # Cột trống bên trái
    table.columns[1].width = Cm(8.5) # Cột ngày tháng, chữ ký bên phải

    cell_right = table.cell(0, 1)
    cell_right._element.clear_content()

    p_date = cell_right.add_paragraph(f"{issuing_location}, {submission_date_str}")
    set_paragraph_format(p_date, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(6))
    set_run_format(p_date.runs[0], size=FONT_SIZE_DEFAULT, italic=True)

    p_signer_label = cell_right.add_paragraph("NGƯỜI LÀM ĐƠN")
    set_paragraph_format(p_signer_label, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(0))
    set_run_format(p_signer_label.runs[0], size=FONT_SIZE_DEFAULT, bold=True)

    p_signer_note = cell_right.add_paragraph("(Ký và ghi rõ họ tên)")
    set_paragraph_format(p_signer_note, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(12))
    set_run_format(p_signer_note.runs[0], size=Pt(11), italic=True)

    cell_right.add_paragraph("\n\n\n") # Khoảng trống ký

    p_signer_name = cell_right.add_paragraph(student_name)
    set_paragraph_format(p_signer_name, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_run_format(p_signer_name.runs[0], size=FONT_SIZE_DEFAULT, bold=True)


    print("Định dạng Đơn xin nhập học hoàn tất.")