# formatters/phap_lenh.py
import re
import time
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
# Sử dụng header và signature của Luật/NQ QH (QPPL)
try:
    from .luat import format_qppl_header, format_qppl_signature
except ImportError:
    def format_qppl_header(document, issuer_name): pass
    def format_qppl_signature(document, signer_title, signer_name): pass

from utils import set_paragraph_format, set_run_format, add_run_with_format
from config import FONT_SIZE_DEFAULT, FONT_SIZE_TITLE, FIRST_LINE_INDENT

def format(document, data):
    print("Bắt đầu định dạng Pháp lệnh...")
    title = data.get("title", "PHÁP LỆNH Ưu đãi người có công với cách mạng").upper()
    body = data.get("body", "Căn cứ Hiến pháp...\nỦy ban thường vụ Quốc hội ban hành Pháp lệnh...\nChương I...\nĐiều 1...")
    ordinance_number = data.get("ordinance_number", "Pháp lệnh số: .../.../UBTVQH...")
    issuing_date_str = data.get("issuing_date", time.strftime("ngày %d tháng %m năm %Y")) # Ngày thông qua
    issuing_location = data.get("issuing_location", "Hà Nội")
    # Cơ quan ban hành Pháp lệnh là UBTVQH
    issuer_name = "ỦY BAN THƯỜNG VỤ QUỐC HỘI"


    # 1. Header (ỦY BAN THƯỜNG VỤ QUỐC HỘI - Kiểu QPPL)
    format_qppl_header(document, issuer_name)

    # 2. Số hiệu và Ngày tháng ban hành (Căn giữa)
    p_num_date = document.add_paragraph(f"{ordinance_number}       {issuing_location}, {issuing_date_str}")
    set_paragraph_format(p_num_date, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(12))
    run_num = add_run_with_format(p_num_date, f"{ordinance_number}       ", size=FONT_SIZE_DEFAULT)
    run_date = add_run_with_format(p_num_date, f"{issuing_location}, {issuing_date_str}", size=FONT_SIZE_DEFAULT, italic=True)


    # 3. Tên Pháp lệnh (IN HOA, đậm, căn giữa)
    p_tenloai = document.add_paragraph(title)
    set_paragraph_format(p_tenloai, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=Pt(6), space_after=Pt(12))
    add_run_with_format(p_tenloai, p_tenloai.text, size=FONT_SIZE_TITLE, bold=True, uppercase=True)


    # 4. Cơ quan ban hành (ỦY BAN THƯỜNG VỤ QUỐC HỘI - Lặp lại)
    p_issuer_body = document.add_paragraph(issuer_name.upper())
    set_paragraph_format(p_issuer_body, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(6))
    add_run_with_format(p_issuer_body, p_issuer_body.text, size=FONT_SIZE_DEFAULT, bold=True, uppercase=True)


    # 5. Căn cứ ban hành
    body_lines = body.split('\n')
    processed_indices = set()

    for i, line in enumerate(body_lines):
        stripped_line = line.strip()
        if not stripped_line: continue
        if stripped_line.lower().startswith("căn cứ") or "ủy ban thường vụ quốc hội ban hành" in stripped_line.lower():
            p = document.add_paragraph(stripped_line)
            set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line_indent=FIRST_LINE_INDENT, line_spacing=1.5, space_after=Pt(0))
            set_run_format(p.runs[0], size=FONT_SIZE_DEFAULT, italic=True)
            processed_indices.add(i)
            if "ủy ban thường vụ quốc hội ban hành" in stripped_line.lower():
                 break # Dừng sau câu này
        # Dừng khi hết căn cứ hoặc gặp Chương/Điều
        elif any(l.strip().lower().startswith("căn cứ") for l in body_lines[:i]):
            if stripped_line.upper().startswith("CHƯƠNG") or stripped_line.upper().startswith("ĐIỀU"):
                 break

    if processed_indices: document.add_paragraph()


    # 6. Nội dung (Chương, Mục, Điều, Khoản, Điểm) - Tương tự Luật
    for i, line in enumerate(body_lines):
        if i in processed_indices: continue
        stripped_line = line.strip()
        if not stripped_line: continue
        p = document.add_paragraph()

        is_chuong = stripped_line.upper().startswith("CHƯƠNG")
        is_muc = re.match(r'^(MỤC\s+\d+)\.?\s+', stripped_line.upper())
        is_dieu = stripped_line.upper().startswith("ĐIỀU")
        is_khoan = re.match(r'^\d+\.\s+', stripped_line)
        is_diem = re.match(r'^[a-z]\)\s+', stripped_line)

        align = WD_ALIGN_PARAGRAPH.JUSTIFY
        left_indent = Cm(0)
        first_indent = FIRST_LINE_INDENT if not (is_chuong or is_muc or is_dieu or is_khoan or is_diem) else Cm(0)
        is_bold = False
        size = FONT_SIZE_DEFAULT
        space_before = Pt(0)
        space_after = Pt(6)
        line_spacing = 1.5

        if is_chuong:
            align = WD_ALIGN_PARAGRAPH.CENTER
            is_bold = True
            space_before = Pt(12)
            size = Pt(13)
        elif is_muc:
            align = WD_ALIGN_PARAGRAPH.CENTER
            is_bold = True
            space_before = Pt(6)
            size = Pt(13)
        elif is_dieu:
            align = WD_ALIGN_PARAGRAPH.LEFT
            is_bold = True
            space_before = Pt(6)
            size = Pt(13)
        elif is_khoan:
            left_indent = Cm(0.5)
        elif is_diem:
            left_indent = Cm(1.0)

        set_paragraph_format(p, alignment=align, left_indent=left_indent, first_line_indent=first_indent, line_spacing=line_spacing, space_before=space_before, space_after=space_after)
        add_run_with_format(p, stripped_line, size=size, bold=is_bold)


    # 7. Thông tin thông qua Pháp lệnh (nếu có)
    adoption_info = data.get("adoption_info", f"Pháp lệnh này đã được Ủy ban thường vụ Quốc hội nước Cộng hòa xã hội chủ nghĩa Việt Nam khóa ... thông qua ngày {issuing_date_str}.")
    p_adoption = document.add_paragraph(adoption_info)
    set_paragraph_format(p_adoption, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=Pt(12), space_after=Pt(12))
    set_run_format(p_adoption.runs[0], size=FONT_SIZE_DEFAULT, italic=True)


    # 8. Chữ ký TM. UBTVQH, Chủ tịch Quốc hội
    signer_title = data.get("signer_title", "TM. ỦY BAN THƯỜNG VỤ QUỐC HỘI\nCHỦ TỊCH")
    signer_name = data.get("signer_name", "[Tên Chủ tịch QH]")
    # Chữ ký Pháp lệnh căn giữa
    sig_paragraph = document.add_paragraph()
    set_paragraph_format(sig_paragraph, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=Pt(0), space_after=Pt(0), line_spacing=1.15)
    # Tách dòng cho thẩm quyền và chức vụ
    title_lines = signer_title.upper().split('\n')
    for line in title_lines:
         add_run_with_format(sig_paragraph, line + "\n", size=Pt(14), bold=True)

    sig_paragraph.add_run("\n\n\n\n\n") # Khoảng trống ký
    add_run_with_format(sig_paragraph, signer_name, size=Pt(14), bold=True) # Tên đậm, to

    # Pháp lệnh thường không có nơi nhận ở cuối

    print("Định dạng Pháp lệnh hoàn tất.")