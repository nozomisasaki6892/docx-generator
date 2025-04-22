# formatters/giay_moi.py
import time
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from utils import set_paragraph_format, set_run_format, add_run_with_format
try:
    from .common_elements import format_basic_header, format_signature_block, format_recipient_list
except ImportError:
    # Fallback nếu chạy độc lập
    def format_basic_header(document, data, doc_type): pass
    def format_signature_block(document, data): pass
    def format_recipient_list(document, data): pass

from config import FONT_SIZE_DEFAULT, FONT_SIZE_TITLE, FIRST_LINE_INDENT

def format(document, data):
    print("Bắt đầu định dạng Giấy mời...")
    title = "GIẤY MỜI"
    body = data.get("body", "Trân trọng kính mời Ông/Bà:...\nTới dự:...\nThời gian:...\nĐịa điểm:...\nRất mong Ông/Bà thu xếp thời gian đến dự.")
    recipient_name = data.get("recipient_name", "Kính gửi: Ông/Bà [Tên người được mời]") # Tên người mời cụ thể


    # 1. Header (Sử dụng header cơ bản)
    format_basic_header(document, data, "GiayMoi")


    # 2. Tên loại văn bản
    p_title = document.add_paragraph(title)
    set_paragraph_format(p_title, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=Pt(12), space_after=Pt(18))
    set_run_format(p_title.runs[0], size=FONT_SIZE_TITLE, bold=True)


    # 3. Kính gửi (Tên người được mời)
    p_kg = document.add_paragraph(recipient_name)
    set_paragraph_format(p_kg, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(12))
    set_run_format(p_kg.runs[0], size=FONT_SIZE_DEFAULT, bold=True)

    # 4. Nội dung mời
    body_lines = body.split('\n')
    inviting_org = data.get("issuing_org", "[Tên cơ quan mời]")
    p_intro = document.add_paragraph(f"{inviting_org} trân trọng kính mời:") # Có thể bỏ nếu Kính gửi đã rõ
    set_paragraph_format(p_intro, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=Pt(6), first_line_indent=FIRST_LINE_INDENT, line_spacing=1.5)
    set_run_format(p_intro.runs[0], size=FONT_SIZE_DEFAULT)

    for line in body_lines:
        stripped_line = line.strip()
        if not stripped_line or "trân trọng kính mời" in stripped_line.lower(): continue # Bỏ dòng đầu nếu có trong body

        p = document.add_paragraph()
        is_info_line = ":" in stripped_line and len(stripped_line.split(":")[0]) < 30

        align = WD_ALIGN_PARAGRAPH.LEFT
        first_indent = Cm(0) # Thông tin mời thường căn trái, không thụt lề

        set_paragraph_format(p, alignment=align, space_after=Pt(6), first_line_indent=first_indent, left_indent=Cm(1.0), line_spacing=1.5) # Thụt lề nội dung mời

        if is_info_line:
            parts = stripped_line.split(":", 1)
            # In đậm nhãn thông tin quan trọng
            if "tới dự" in parts[0].lower() or "thời gian" in parts[0].lower() or "địa điểm" in parts[0].lower():
                 add_run_with_format(p, parts[0] + ":", size=FONT_SIZE_DEFAULT, bold=True)
                 add_run_with_format(p, parts[1], size=FONT_SIZE_DEFAULT, bold=True)
            else:
                 add_run_with_format(p, parts[0] + ":", size=FONT_SIZE_DEFAULT)
                 add_run_with_format(p, parts[1], size=FONT_SIZE_DEFAULT)
        elif "rất mong" in stripped_line.lower():
            # Đoạn kết mời không thụt lề trái
            set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=Pt(12), space_after=Pt(6), first_line_indent=Cm(0), left_indent=Cm(0), line_spacing=1.5)
            add_run_with_format(p, stripped_line, size=FONT_SIZE_DEFAULT, italic=False, bold=True) # Câu kết đậm
        else:
             add_run_with_format(p, stripped_line, size=FONT_SIZE_DEFAULT)


    # 5. Chữ ký (Sử dụng khối chữ ký cơ bản)
    # Chức vụ người ký giấy mời
    if 'signer_title' not in data: data['signer_title'] = "THỦ TRƯỞNG CƠ QUAN"
    if 'signer_name' not in data: data['signer_name'] = "[Họ và tên]"
    document.add_paragraph()
    format_signature_block(document, data)

    # 6. Nơi nhận (Thường chỉ có Lưu VT)
    if 'recipients' not in data: data['recipients'] = ["- Như trên;", "- Lưu: VT."]
    format_recipient_list(document, data)

    print("Định dạng Giấy mời hoàn tất.")