# formatters/huong_dan.py (Đã sửa SyntaxError)
import re
import time
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from utils import set_paragraph_format, set_run_format, add_run_with_format
try:
    from .common_elements import format_basic_header, format_signature_block, format_recipient_list
except ImportError:
    # Fallback nếu chạy độc lập
    def format_basic_header(document, data, doc_type): pass
    def format_signature_block(document, data): pass
    def format_recipient_list(document, data): pass

from config import FONT_SIZE_DEFAULT, FONT_SIZE_TITLE, FIRST_LINE_INDENT

def format(document, data):
    print("Bắt đầu định dạng Hướng dẫn...")
    title = data.get("title", "Hướng dẫn Thực hiện Nghị định/Thông tư...")
    body = data.get("body", "Căn cứ...\nĐể triển khai thực hiện..., [Cơ quan ban hành] hướng dẫn như sau:\nI. QUY ĐỊNH CHUNG\nII. NỘI DUNG HƯỚNG DẪN CỤ THỂ\n   1. Về đối tượng...\n   2. Về trình tự...\nIII. TỔ CHỨC THỰC HIỆN...")
    doc_type_label = "HƯỚNG DẪN"

    # 1. Header (Sử dụng header cơ bản)
    format_basic_header(document, data, "HuongDan")

    # 2. Tiêu đề
    p_title = document.add_paragraph(doc_type_label)
    set_paragraph_format(p_title, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=Pt(12), space_after=Pt(6))
    set_run_format(p_title.runs[0], size=FONT_SIZE_TITLE, bold=True)

    # Trích yếu nội dung hướng dẫn
    subject = title.replace("Hướng dẫn", "").strip()
    if subject.lower().startswith("về việc"):
        subject = subject.split(" ", 2)[-1]
    elif subject.lower().startswith("thực hiện"):
         subject = subject.split(" ", 1)[-1]

    p_subject = document.add_paragraph(subject)
    set_paragraph_format(p_subject, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(12))
    set_run_format(p_subject.runs[0], size=Pt(14), bold=True) # Trích yếu đậm


    # 3. Căn cứ ban hành (nếu có trong body)
    body_lines = body.split('\n')
    processed_indices = set()

    for i, line in enumerate(body_lines):
        stripped_line = line.strip()
        if not stripped_line: continue
        if stripped_line.lower().startswith("căn cứ"):
            p = document.add_paragraph(stripped_line)
            set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line_indent=FIRST_LINE_INDENT, line_spacing=1.5, space_after=Pt(0))
            set_run_format(p.runs[0], size=FONT_SIZE_DEFAULT, italic=True)
            processed_indices.add(i)
        # Dừng khi hết căn cứ
        elif any(l.strip().lower().startswith("căn cứ") for l in body_lines[:i]):
             break
    if processed_indices: document.add_paragraph() # Thêm khoảng trống sau căn cứ

    # Câu dẫn vào nội dung hướng dẫn
    intro_line_found = False
    for i, line in enumerate(body_lines):
         if i in processed_indices: continue
         stripped_line = line.strip()
         if ("hướng dẫn như sau" in stripped_line.lower() or "hướng dẫn thực hiện" in stripped_line.lower()) and ":" in stripped_line:
             p_intro = document.add_paragraph(stripped_line)
             set_paragraph_format(p_intro, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_before=Pt(6), space_after=Pt(12), first_line_indent=FIRST_LINE_INDENT, line_spacing=1.5)
             set_run_format(p_intro.runs[0], size=FONT_SIZE_DEFAULT)
             processed_indices.add(i)
             intro_line_found = True
             break

    # 4. Nội dung Hướng dẫn
    for i, line in enumerate(body_lines):
        if i in processed_indices: continue # Bỏ qua căn cứ và câu dẫn đã xử lý
        stripped_line = line.strip()
        if not stripped_line: continue
        p = document.add_paragraph()

        # Logic định dạng cơ bản cho các đề mục hướng dẫn
        is_part_roman = re.match(r'^([IVXLCDM]+)\.\s+', stripped_line.upper()) # Mục La Mã I, II
        is_section_digit = re.match(r'^(\d+\.)\s+', stripped_line) # Mục 1, 2, 3
        is_subsection_alpha = re.match(r'^[a-z]\)\s+', stripped_line) # Mục a, b, c
        is_subsubsection_digit = re.match(r'^-?\s?(\d+\.\d+\.\d+)\.?\s+', stripped_line) # Mục 1.1.1
        is_bullet = stripped_line.startswith("-") or stripped_line.startswith("+") or stripped_line.startswith("*") or stripped_line.startswith("•")

        align = WD_ALIGN_PARAGRAPH.JUSTIFY
        left_indent = Cm(0)
        first_indent = FIRST_LINE_INDENT if not (is_part_roman or is_section_digit or is_subsection_alpha or is_subsubsection_digit or is_bullet) else Cm(0)
        is_bold = bool(is_part_roman or is_section_digit) # Mục La Mã, mục số đậm
        is_italic = False
        size = FONT_SIZE_DEFAULT
        space_before = Pt(0)
        space_after = Pt(6)
        line_spacing = 1.5

        if is_part_roman:
            align = WD_ALIGN_PARAGRAPH.LEFT
            space_before = Pt(12)
            size = Pt(13)
        elif is_section_digit:
            align = WD_ALIGN_PARAGRAPH.LEFT
            left_indent = Cm(0.5)
            space_before = Pt(6)
        elif is_subsection_alpha:
             align = WD_ALIGN_PARAGRAPH.LEFT
             left_indent = Cm(1.0)
             is_bold = False # Mục chữ thường không đậm
        elif is_subsubsection_digit:
             align = WD_ALIGN_PARAGRAPH.LEFT
             left_indent = Cm(1.5)
             is_bold = False
        elif is_bullet:
             align = WD_ALIGN_PARAGRAPH.LEFT
             left_indent = Cm(2.0) # Tăng thụt lề cho bullet trong mục con
             first_indent = Cm(-0.5) # Hanging indent
             is_bold = False


        set_paragraph_format(p, alignment=align, left_indent=left_indent, first_line_indent=first_indent, line_spacing=line_spacing, space_before=space_before, space_after=space_after)
        add_run_with_format(p, stripped_line, size=size, bold=is_bold, italic=is_italic)


    # 5. Chữ ký (Sử dụng khối chữ ký cơ bản)
    if 'signer_title' not in data: data['signer_title'] = "THỦ TRƯỞNG CƠ QUAN"
    if 'signer_name' not in data: data['signer_name'] = "[Họ và tên]"
    document.add_paragraph()
    # --- Sửa lỗi thiếu dấu ) ở dòng dưới ---
    format_signature_block(document, data) # Thêm dấu )

    # 6. Nơi nhận (Sử dụng nơi nhận cơ bản)
    if 'recipients' not in data: data['recipients'] = ["- Như trên;", "- Lưu: VT, ...;"]
    format_recipient_list(document, data)

    print("Định dạng Hướng dẫn hoàn tất.")