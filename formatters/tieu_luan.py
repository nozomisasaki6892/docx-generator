# formatters/tieu_luan.py
import re
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from utils import set_paragraph_format, set_run_format, add_run_with_format, add_centered_text
from config import FONT_SIZE_DEFAULT, FONT_SIZE_TITLE, FIRST_LINE_INDENT

# Lưu ý: Định dạng tiểu luận rất đa dạng, phụ thuộc yêu cầu của từng trường/giáo viên.
# Code này tạo cấu trúc cơ bản, cần tùy chỉnh (trang bìa, lề, mục lục...).

def format(document, data):
    print("Bắt đầu định dạng Tiểu luận (cơ bản)...")
    title = data.get("title", "Tiểu luận Môn học ABC")
    body = data.get("body", "Lời mở đầu...\nChương 1: Cơ sở lý luận...\n   1.1...\nChương 2: Thực trạng...\nChương 3: Giải pháp...\nKết luận...\nTài liệu tham khảo...")
    student_name = data.get("student_name", "[Họ và tên Sinh viên]")
    student_id = data.get("student_id", "[Mã số Sinh viên]")
    student_class = data.get("student_class", "[Lớp]")
    instructor_name = data.get("instructor_name", "[Giảng viên hướng dẫn]")
    university_name = data.get("university_name", "TRƯỜNG ĐẠI HỌC XYZ").upper()
    faculty_name = data.get("faculty_name", "KHOA [TÊN KHOA]").upper()
    submission_place = data.get("submission_place", "Hà Nội")
    submission_year = data.get("submission_year", "2025")


    # --- Trang bìa (Ví dụ đơn giản) ---
    # Cần thiết kế chi tiết hơn
    add_centered_text(document, university_name, size=Pt(13), bold=True, space_after=0)
    add_centered_text(document, faculty_name, size=Pt(13), bold=True, space_after=36)

    add_centered_text(document, "TIỂU LUẬN", size=Pt(16), bold=True, space_before=18, space_after=12)
    # Tên đề tài
    essay_topic = title.replace("Tiểu luận", "").strip()
    add_centered_text(document, f"Đề tài: {essay_topic}", size=Pt(14), bold=True, space_after=36)


    p_info = document.add_paragraph()
    set_paragraph_format(p_info, alignment=WD_ALIGN_PARAGRAPH.RIGHT, space_after=Pt(6))
    add_run_with_format(p_info, f"Sinh viên thực hiện: {student_name}\n", size=Pt(13))
    add_run_with_format(p_info, f"MSSV: {student_id}\n", size=Pt(13))
    add_run_with_format(p_info, f"Lớp: {student_class}\n", size=Pt(13))
    add_run_with_format(p_info, f"Giảng viên hướng dẫn: {instructor_name}", size=Pt(13))

    # Địa điểm, năm
    add_centered_text(document, f"{submission_place}, {submission_year}", size=Pt(13), bold=False, space_before=72, space_after=0)

    document.add_page_break()

    # --- Nội dung Tiểu luận ---
    # Cần thêm Mục lục, Lời cảm ơn, Danh mục bảng biểu... nếu cần
    # Code dưới đây chỉ định dạng các đề mục chính

    body_lines = body.split('\n')
    for line in body_lines:
        stripped_line = line.strip()
        if not stripped_line: continue
        p = document.add_paragraph()

        # Logic định dạng cơ bản
        is_chapter = stripped_line.upper().startswith("CHƯƠNG")
        is_section_digit = re.match(r'^(\d+)\.\s+', stripped_line) # Mục 1, 2, 3
        is_subsection_digit = re.match(r'^(\d+\.\d+\.?)\s+', stripped_line) # Mục 1.1, 1.2
        is_subsubsection_digit = re.match(r'^(\d+\.\d+\.\d+\.?)\s+', stripped_line) # Mục 1.1.1
        is_heading = stripped_line.upper() in ["LỜI MỞ ĐẦU", "KẾT LUẬN", "TÀI LIỆU THAM KHẢO", "MỤC LỤC", "LỜI CẢM ƠN"]
        is_bullet = stripped_line.startswith("-") or stripped_line.startswith("+") or stripped_line.startswith("*") or stripped_line.startswith("•")

        align = WD_ALIGN_PARAGRAPH.JUSTIFY
        left_indent = Cm(0)
        # Tiểu luận thường thụt lề dòng đầu tiên của đoạn văn
        first_indent = FIRST_LINE_INDENT if not (is_chapter or is_section_digit or is_subsection_digit or is_subsubsection_digit or is_heading or is_bullet) else Cm(0)
        is_bold = bool(is_chapter or is_section_digit or is_subsection_digit or is_heading) # Mục lớn đậm
        is_italic = False
        size = Pt(13) # Cỡ chữ nội dung tiểu luận thường là 13pt
        space_before = Pt(0)
        space_after = Pt(6)
        line_spacing = 1.5 # Giãn dòng 1.5 lines

        if is_heading or is_chapter:
            align = WD_ALIGN_PARAGRAPH.CENTER
            is_bold = True
            space_before = Pt(18)
            space_after = Pt(12)
            size = Pt(14) # Tiêu đề mục lớn
        elif is_section_digit:
            align = WD_ALIGN_PARAGRAPH.LEFT
            space_before = Pt(12)
        elif is_subsection_digit:
            align = WD_ALIGN_PARAGRAPH.LEFT
            left_indent = Cm(0.5)
            space_before = Pt(6)
        elif is_subsubsection_digit:
             align = WD_ALIGN_PARAGRAPH.LEFT
             left_indent = Cm(1.0)
             is_bold = False # Mục con nhỏ thường không đậm
             is_italic = True # Có thể in nghiêng
        elif is_bullet:
             align = WD_ALIGN_PARAGRAPH.LEFT
             left_indent = Cm(1.5)
             first_indent = Cm(-0.5) # Hanging indent

        set_paragraph_format(p, alignment=align, left_indent=left_indent, first_line_indent=first_indent, line_spacing=line_spacing, space_before=space_before, space_after=space_after)
        # Xử lý riêng phần Tài liệu tham khảo nếu cần định dạng treo
        if "TÀI LIỆU THAM KHẢO" in data.get("current_section", "").upper() and not is_heading:
             set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, left_indent=Cm(1.0), first_line_indent=Cm(-1.0), line_spacing=line_spacing, space_after=Pt(6))

        add_run_with_format(p, stripped_line, size=size, bold=is_bold, italic=is_italic)


    print("Định dạng Tiểu luận (cơ bản) hoàn tất.")