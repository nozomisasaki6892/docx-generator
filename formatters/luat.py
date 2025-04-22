# formatters/luat.py
import re
import time
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from utils import set_paragraph_format, set_run_format, add_run_with_format
from config import FONT_SIZE_DEFAULT, FONT_SIZE_TITLE, FIRST_LINE_INDENT, FONT_SIZE_HEADER

def format_qppl_header(document, issuer_name):
    # Header VBQPPL chỉ có QH/TN và Tên CQ ban hành căn giữa
    p_qh = document.add_paragraph("CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM")
    set_paragraph_format(p_qh, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(0))
    add_run_with_format(p_qh, p_qh.text, size=FONT_SIZE_HEADER, bold=True)

    p_tn = document.add_paragraph("Độc lập - Tự do - Hạnh phúc")
    set_paragraph_format(p_tn, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(0))
    add_run_with_format(p_tn, p_tn.text, size=Pt(14), bold=True) # Tiêu ngữ to hơn

    p_line_tn = document.add_paragraph("-" * 20)
    set_paragraph_format(p_line_tn, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(12))

    # Tên cơ quan ban hành dưới QH/TN
    p_issuer = document.add_paragraph(issuer_name.upper())
    set_paragraph_format(p_issuer, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(12))
    add_run_with_format(p_issuer, p_issuer.text, size=FONT_SIZE_HEADER, bold=False) # Tên CQ không đậm

def format_qppl_signature(document, signer_title, signer_name):
    # Chữ ký VBQPPL (Luật, NQ QH...) căn giữa
    sig_paragraph = document.add_paragraph()
    set_paragraph_format(sig_paragraph, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=Pt(12), space_after=Pt(0), line_spacing=1.15)
    add_run_with_format(sig_paragraph, signer_title.upper() + "\n\n\n\n\n", size=Pt(14), bold=True) # Chức vụ đậm, to
    add_run_with_format(sig_paragraph, signer_name, size=Pt(14), bold=True) # Tên đậm, to

def format(document, data):
    print("Bắt đầu định dạng Luật/Bộ luật...")
    title = data.get("title", "LUẬT ABC").upper() # VD: LUẬT AN NINH MẠNG
    body = data.get("body", "Căn cứ Hiến pháp...\nQuốc hội ban hành Luật...\nChương I...\nĐiều 1...")
    law_number = data.get("law_number", "Luật số: .../.../QH...") # VD: Luật số: 24/2018/QH14
    adoption_date_str = data.get("adoption_date", time.strftime("ngày %d tháng %m năm %Y")) # Ngày QH thông qua

    # 1. Header (Tên cơ quan ban hành là QUỐC HỘI)
    format_qppl_header(document, "QUỐC HỘI")

    # 2. Số hiệu Luật (Căn giữa)
    p_num = document.add_paragraph(law_number)
    set_paragraph_format(p_num, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(6))
    add_run_with_format(p_num, p_num.text, size=FONT_SIZE_DEFAULT) # Cỡ chữ 13-14

    # 3. Tên Luật (Căn giữa, IN HOA, đậm)
    p_tenluat = document.add_paragraph(title)
    set_paragraph_format(p_tenluat, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=Pt(6), space_after=Pt(12))
    add_run_with_format(p_tenluat, p_tenluat.text, size=FONT_SIZE_TITLE, bold=True, uppercase=True) # Cỡ 14

    # 4. Lời nói đầu / Căn cứ ban hành
    body_lines = body.split('\n')
    processed_indices = set()

    for i, line in enumerate(body_lines):
        stripped_line = line.strip()
        if not stripped_line: continue
        # Tìm các dòng căn cứ hoặc lời dẫn đầu
        if stripped_line.lower().startswith("căn cứ") or "quốc hội ban hành" in stripped_line.lower():
            p = document.add_paragraph(stripped_line)
            set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line_indent=FIRST_LINE_INDENT, line_spacing=1.5, space_after=Pt(6))
            set_run_format(p.runs[0], size=FONT_SIZE_DEFAULT, italic=True) # Căn cứ in nghiêng
            processed_indices.add(i)
        # Dừng khi gặp cấu trúc Chương, Điều
        elif stripped_line.upper().startswith("CHƯƠNG") or stripped_line.upper().startswith("ĐIỀU"):
            break
        # Xử lý các dòng khác trong lời nói đầu (nếu có)
        elif i < 5 and len(processed_indices) > 0: # Heuristic cho các dòng đầu sau căn cứ
             p = document.add_paragraph(stripped_line)
             set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line_indent=FIRST_LINE_INDENT, line_spacing=1.5, space_after=Pt(6))
             set_run_format(p.runs[0], size=FONT_SIZE_DEFAULT)
             processed_indices.add(i)

    # 5. Nội dung (Chương, Mục, Điều, Khoản, Điểm)
    for i, line in enumerate(body_lines):
        if i in processed_indices: continue
        stripped_line = line.strip()
        if not stripped_line: continue
        p = document.add_paragraph()

        is_chuong = stripped_line.upper().startswith("CHƯƠNG")
        is_muc = re.match(r'^(MỤC\s+\d+)\.?\s+', stripped_line.upper())
        is_dieu = stripped_line.upper().startswith("ĐIỀU")
        is_khoan = re.match(r'^\d+\.\s+', stripped_line)
        is_diem = re.match(r'^[a-z]\)\s+', stripped_line)

        align = WD_ALIGN_PARAGRAPH.JUSTIFY
        left_indent = Cm(0)
        first_indent = FIRST_LINE_INDENT if not (is_chuong or is_muc or is_dieu or is_khoan or is_diem) else Cm(0)
        is_bold = False
        size = FONT_SIZE_DEFAULT
        space_before = Pt(0)
        space_after = Pt(6)
        line_spacing = 1.5

        if is_chuong:
            align = WD_ALIGN_PARAGRAPH.CENTER
            is_bold = True
            space_before = Pt(12)
            size = Pt(13) # Cỡ chữ như nội dung
        elif is_muc:
            align = WD_ALIGN_PARAGRAPH.CENTER
            is_bold = True
            space_before = Pt(6)
            size = Pt(13)
        elif is_dieu:
            align = WD_ALIGN_PARAGRAPH.LEFT # Điều căn trái
            is_bold = True
            space_before = Pt(6)
            size = Pt(13)
        elif is_khoan:
            left_indent = Cm(0.5)
        elif is_diem:
            left_indent = Cm(1.0)

        set_paragraph_format(p, alignment=align, left_indent=left_indent, first_line_indent=first_indent, line_spacing=line_spacing, space_before=space_before, space_after=space_after)
        add_run_with_format(p, stripped_line, size=size, bold=is_bold)

    # 6. Thông tin thông qua Luật
    adoption_info = data.get("adoption_info", f"{title} này đã được Quốc hội nước Cộng hòa xã hội chủ nghĩa Việt Nam khóa ... kỳ họp thứ ... thông qua ngày {adoption_date_str}.")
    p_adoption = document.add_paragraph(adoption_info)
    set_paragraph_format(p_adoption, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=Pt(12), space_after=Pt(12))
    set_run_format(p_adoption.runs[0], size=FONT_SIZE_DEFAULT, italic=True) # Thông tin thông qua nghiêng

    # 7. Chữ ký Chủ tịch Quốc hội
    signer_title = data.get("signer_title", "CHỦ TỊCH QUỐC HỘI")
    signer_name = data.get("signer_name", "[Tên Chủ tịch QH]")
    format_qppl_signature(document, signer_title, signer_name)

    # Luật thường không có Nơi nhận

    print("Định dạng Luật hoàn tất.")