# formatters/huong_dan_hs.py
import re
import time
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from utils import set_paragraph_format, set_run_format, add_run_with_format
try:
    # Dùng header, signature, recipient của văn bản HC thông thường
    from .common_elements import format_basic_header, format_signature_block, format_recipient_list
except ImportError:
    def format_basic_header(document, data, doc_type): pass
    def format_signature_block(document, data): pass
    def format_recipient_list(document, data): pass

from config import FONT_SIZE_DEFAULT, FONT_SIZE_TITLE, FIRST_LINE_INDENT

def format(document, data):
    # Định dạng này tương tự Thông báo hoặc Công văn
    print("Bắt đầu định dạng Hướng dẫn Hồ sơ...")
    title = data.get("title", "Hướng dẫn Chuẩn bị hồ sơ nhập học/Tuyển sinh...")
    body = data.get("body", "Căn cứ Kế hoạch tuyển sinh..., Nhà trường hướng dẫn hồ sơ nhập học gồm các giấy tờ sau:\n1. Giấy báo trúng tuyển (bản chính).\n2. Học bạ THPT (bản chính và bản sao công chứng).\n3. Bằng tốt nghiệp THPT hoặc Giấy chứng nhận tốt nghiệp tạm thời (bản chính và bản sao công chứng).\n...")
    doc_type_label = "HƯỚNG DẪN" # Hoặc THÔNG BÁO tùy cách gọi

    # 1. Header (Sử dụng header cơ bản)
    format_basic_header(document, data, "HuongDanHS") # Có thể dùng type ThongBaoNT nếu phù hợp hơn

    # 2. Tiêu đề
    p_title = document.add_paragraph(doc_type_label)
    set_paragraph_format(p_title, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=Pt(12), space_after=Pt(6))
    set_run_format(p_title.runs[0], size=FONT_SIZE_TITLE, bold=True)

    # Trích yếu nội dung hướng dẫn
    subject = title.replace("Hướng dẫn", "").strip()
    p_subject = document.add_paragraph(f"V/v: {subject}")
    set_paragraph_format(p_subject, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(12))
    set_run_format(p_subject.runs[0], size=Pt(14), bold=True) # Trích yếu đậm


    # 3. Căn cứ / Lời dẫn (nếu có)
    body_lines = body.split('\n')
    processed_indices = set()

    for i, line in enumerate(body_lines):
        stripped_line = line.strip()
        if not stripped_line: continue
        if stripped_line.lower().startswith("căn cứ"):
            p = document.add_paragraph(stripped_line)
            set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line_indent=FIRST_LINE_INDENT, line_spacing=1.5, space_after=Pt(0))
            set_run_format(p.runs[0], size=FONT_SIZE_DEFAULT, italic=True)
            processed_indices.add(i)
        elif any(l.strip().lower().startswith("căn cứ") for l in body_lines[:i]):
             break

    # Câu dẫn vào nội dung
    intro_line_found = False
    for i, line in enumerate(body_lines):
         if i in processed_indices: continue
         stripped_line = line.strip()
         # Tìm câu dẫn chung chung
         if "hướng dẫn" in stripped_line.lower() and "hồ sơ" in stripped_line.lower() and "sau:" in stripped_line.lower():
             p_intro = document.add_paragraph(stripped_line)
             set_paragraph_format(p_intro, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_before=Pt(6), space_after=Pt(12), first_line_indent=FIRST_LINE_INDENT, line_spacing=1.5)
             set_run_format(p_intro.runs[0], size=FONT_SIZE_DEFAULT)
             processed_indices.add(i)
             intro_line_found = True
             break


    # 4. Nội dung Hướng dẫn (Thường là danh sách)
    for i, line in enumerate(body_lines):
        if i in processed_indices: continue
        stripped_line = line.strip()
        if not stripped_line: continue
        p = document.add_paragraph()

        is_numbered_item = re.match(r'^(\d+)\.\s+', stripped_line) # 1., 2.
        is_bullet = stripped_line.startswith("-") or stripped_line.startswith("+") or stripped_line.startswith("*") or stripped_line.startswith("•")

        align = WD_ALIGN_PARAGRAPH.JUSTIFY
        left_indent = Cm(0.5) # Thụt lề cho danh sách
        first_indent = Cm(0)
        is_bold = False
        is_italic = False
        size = FONT_SIZE_DEFAULT
        space_before = Pt(0)
        space_after = Pt(6)
        line_spacing = 1.5

        if is_numbered_item:
            align = WD_ALIGN_PARAGRAPH.LEFT
            first_indent = Cm(-0.5) # Hanging indent
        elif is_bullet:
             align = WD_ALIGN_PARAGRAPH.LEFT
             left_indent = Cm(1.0) # Thụt lề sâu hơn cho bullet
             first_indent = Cm(-0.5) # Hanging indent

        set_paragraph_format(p, alignment=align, left_indent=left_indent, first_line_indent=first_indent, line_spacing=line_spacing, space_before=space_before, space_after=space_after)
        add_run_with_format(p, stripped_line, size=size, bold=is_bold, italic=is_italic)