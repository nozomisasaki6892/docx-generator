# formatters/quy_che_ts.py
import re
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from utils import set_paragraph_format, set_run_format, add_run_with_format
try:
    # Quy chế TS thường ban hành kèm Quyết định của Hiệu trưởng
    from .common_elements import format_basic_header # Header của trường
except ImportError:
    from common_elements import format_basic_header
from config import FONT_SIZE_DEFAULT, FONT_SIZE_TITLE, FIRST_LINE_INDENT

def format(document, data):
    print("Bắt đầu định dạng Quy chế tuyển sinh...")
    title = data.get("title", "Quy chế tuyển sinh trình độ [Trình độ] năm [Năm]")
    body = data.get("body", "Nội dung quy chế...")
    issuing_org = data.get("issuing_org", "TÊN TRƯỜNG").upper()
    attached_decision = data.get("attached_decision", None) # VD: Quyết định số .../QĐ-ĐH...

    # 1. Header của trường (Nếu là văn bản độc lập)
    # format_basic_header(document, data, "QuyCheTS") # Bỏ qua nếu kèm QĐ

    # 2. Tên loại
    p_tenloai = document.add_paragraph("QUY CHẾ")
    set_paragraph_format(p_tenloai, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=Pt(12), space_after=Pt(6))
    add_run_with_format(p_tenloai, "QUY CHẾ", size=FONT_SIZE_TITLE, bold=True, uppercase=True)

    # 3. Tiêu đề của Quy chế
    qc_title = title.replace("Quy chế", "").strip()
    p_title = document.add_paragraph(qc_title)
    set_paragraph_format(p_title, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(6))
    add_run_with_format(p_title, qc_title, size=Pt(14), bold=True)

    # 4. Dòng "Ban hành kèm theo Quyết định số..."
    if attached_decision:
        p_attach = document.add_paragraph(f"(Ban hành kèm theo {attached_decision})")
        set_paragraph_format(p_attach, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(12))
        add_run_with_format(p_attach, p_attach.text, size=FONT_SIZE_DEFAULT, italic=True)

    # 5. Nội dung (Chương, Điều, Khoản, Điểm)
    body_lines = body.split('\n')
    for line in body_lines:
        stripped_line = line.strip()
        if not stripped_line: continue

        p = document.add_paragraph()
        # Copy logic nhận diện và định dạng Chương, Điều, Khoản, Điểm từ quy_dinh.py
        is_chuong = stripped_line.upper().startswith("CHƯƠNG")
        is_dieu = stripped_line.upper().startswith("ĐIỀU")
        is_khoan = re.match(r'^\d+\.\s+', stripped_line)
        is_diem = re.match(r'^[a-z]\)\s+', stripped_line)

        align = WD_ALIGN_PARAGRAPH.JUSTIFY
        left_indent = Cm(0)
        first_indent = FIRST_LINE_INDENT
        is_bold = False
        size = FONT_SIZE_DEFAULT
        space_before = Pt(0)

        if is_chuong:
            align = WD_ALIGN_PARAGRAPH.CENTER
            first_indent = Cm(0)
            is_bold = True
            space_before = Pt(12)
        elif is_dieu:
            align = WD_ALIGN_PARAGRAPH.LEFT
            first_indent = Cm(0)
            is_bold = True
            space_before = Pt(6)
        elif is_khoan:
            left_indent = Cm(0.5)
            first_indent = Cm(0)
        elif is_diem:
            left_indent = Cm(1.0)
            first_indent = Cm(0)

        set_paragraph_format(p, alignment=align, left_indent=left_indent, first_line_indent=first_indent, line_spacing=1.5, space_before=space_before, space_after=Pt(6))
        add_run_with_format(p, stripped_line, size=size, bold=is_bold)

    # Quy chế ban hành kèm QĐ thường không có chữ ký/nơi nhận riêng

    print("Định dạng Quy chế tuyển sinh hoàn tất.")