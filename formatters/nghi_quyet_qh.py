# formatters/nghi_quyet_qh.py
import re
import time
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
# Giả sử dùng header và signature của Luật
from .luat import format_qppl_header, format_qppl_signature
from utils import set_paragraph_format, set_run_format, add_run_with_format
from config import FONT_SIZE_DEFAULT, FONT_SIZE_TITLE, FIRST_LINE_INDENT

def format(document, data):
    print("Bắt đầu định dạng Nghị quyết QPPL của Quốc hội...")
    title = data.get("title", "Nghị quyết của Quốc hội về ABC")
    body = data.get("body", "Nội dung nghị quyết...")
    resolution_number = data.get("resolution_number", "Nghị quyết số: .../.../QH...")
    adoption_date_str = data.get("adoption_date", time.strftime("ngày %d tháng %m năm %Y"))

    # 1. Header (QUỐC HỘI)
    format_qppl_header(document, "QUỐC HỘI")

    # 2. Số hiệu Nghị quyết
    p_num = document.add_paragraph(resolution_number)
    set_paragraph_format(p_num, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(6))
    add_run_with_format(p_num, p_num.text, size=FONT_SIZE_DEFAULT)

    # 3. Tên Nghị quyết
    nq_title = title.replace("của Quốc hội", "").strip().upper() # Bỏ "của Quốc hội"
    p_tenloai = document.add_paragraph(nq_title)
    set_paragraph_format(p_tenloai, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=Pt(6), space_after=Pt(12))
    add_run_with_format(p_tenloai, p_tenloai.text, size=FONT_SIZE_TITLE, bold=True, uppercase=True)

    # 4. Cơ quan ban hành (QUỐC HỘI)
    p_issuer = document.add_paragraph("QUỐC HỘI")
    set_paragraph_format(p_issuer, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(6))
    add_run_with_format(p_issuer, p_issuer.text, size=FONT_SIZE_DEFAULT, bold=True, uppercase=True)

    # 5. Căn cứ ban hành
    preamble = data.get("preamble", "Căn cứ Hiến pháp nước Cộng hòa xã hội chủ nghĩa Việt Nam;\n[Căn cứ khác nếu có];\nXét đề nghị của ...;") # Mẫu preamble
    preamble_lines = preamble.split('\n')
    for line in preamble_lines:
         p_pre = document.add_paragraph(line)
         set_paragraph_format(p_pre, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line_indent=FIRST_LINE_INDENT, space_after=Pt(0), line_spacing=1.5)
         add_run_with_format(p_pre, line, size=FONT_SIZE_DEFAULT, italic=True)

    # 6. QUYẾT NGHỊ:
    p_qn_label = document.add_paragraph("QUYẾT NGHỊ:")
    set_paragraph_format(p_qn_label, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=Pt(6), space_after=Pt(6))
    add_run_with_format(p_qn_label, "QUYẾT NGHỊ:", size=FONT_SIZE_DEFAULT, bold=True, uppercase=True)

    # 7. Nội dung (Điều, Khoản, Điểm)
    body_lines = body.split('\n')
    for line in body_lines:
        stripped_line = line.strip()
        if not stripped_line: continue
        p = document.add_paragraph()

        is_dieu = stripped_line.upper().startswith("ĐIỀU")
        is_khoan = re.match(r'^\d+\.\s+', stripped_line)
        is_diem = re.match(r'^[a-z]\)\s+', stripped_line)

        align = WD_ALIGN_PARAGRAPH.JUSTIFY
        left_indent = Cm(0)
        first_indent = FIRST_LINE_INDENT
        is_bold = False

        if is_dieu:
            align = WD_ALIGN_PARAGRAPH.LEFT
            first_indent = Cm(0)
            is_bold = True
        elif is_khoan:
            left_indent = Cm(0.5)
            first_indent = Cm(0)
        elif is_diem:
            left_indent = Cm(1.0)
            first_indent = Cm(0)

        set_paragraph_format(p, alignment=align, left_indent=left_indent, first_line_indent=first_indent, line_spacing=1.5, space_after=Pt(6))
        add_run_with_format(p, stripped_line, size=FONT_SIZE_DEFAULT, bold=is_bold)

    # 8. Thông tin thông qua
    p_adoption = document.add_paragraph(f"Nghị quyết này đã được Quốc hội nước Cộng hòa xã hội chủ nghĩa Việt Nam khóa ... kỳ họp thứ ... thông qua ngày {adoption_date_str}.")
    set_paragraph_format(p_adoption, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=Pt(12), space_after=Pt(12))
    add_run_with_format(p_adoption, p_adoption.text, size=FONT_SIZE_DEFAULT, italic=True)

    # 9. Chữ ký Chủ tịch Quốc hội
    format_qppl_signature(document, "CHỦ TỊCH QUỐC HỘI", data.get("signer_name", "[Tên Chủ tịch QH]"))

    print("Định dạng Nghị quyết QH hoàn tất.")