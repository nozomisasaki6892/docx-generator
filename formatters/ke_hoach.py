# formatters/ke_hoach.py
import re
import time
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from utils import set_paragraph_format, set_run_format, add_run_with_format
try:
    from .common_elements import format_basic_header, format_signature_block, format_recipient_list
except ImportError:
    # Fallback nếu chạy độc lập
    def format_basic_header(document, data, doc_type): pass
    def format_signature_block(document, data): pass
    def format_recipient_list(document, data): pass

from config import FONT_SIZE_DEFAULT, FONT_SIZE_TITLE, FIRST_LINE_INDENT

def format(document, data):
    print("Bắt đầu định dạng Kế hoạch...")
    title = data.get("title", "Kế hoạch Tổ chức Hội nghị/Triển khai Công tác...")
    body = data.get("body", "I. MỤC ĐÍCH, YÊU CẦU\nII. NỘI DUNG KẾ HOẠCH\n   1. Thời gian, địa điểm\n   2. Thành phần tham dự\n   3. Nội dung chi tiết\nIII. KINH PHÍ THỰC HIỆN\nIV. TỔ CHỨC THỰC HIỆN...")
    doc_type_label = "KẾ HOẠCH"

    # 1. Header (Sử dụng header cơ bản)
    format_basic_header(document, data, "KeHoach")

    # 2. Tiêu đề
    p_title = document.add_paragraph(doc_type_label)
    set_paragraph_format(p_title, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=Pt(12), space_after=Pt(6))
    set_run_format(p_title.runs[0], size=FONT_SIZE_TITLE, bold=True)

    # Trích yếu nội dung kế hoạch
    subject = title.replace("Kế hoạch", "").strip()
    p_subject = document.add_paragraph(subject)
    set_paragraph_format(p_subject, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(12))
    set_run_format(p_subject.runs[0], size=Pt(14), bold=True) # Trích yếu đậm


    # 3. Căn cứ lập kế hoạch (nếu có)
    body_lines = body.split('\n')
    processed_indices = set()

    for i, line in enumerate(body_lines):
        stripped_line = line.strip()
        if not stripped_line: continue
        if stripped_line.lower().startswith("căn cứ"):
            p = document.add_paragraph(stripped_line)
            set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line_indent=FIRST_LINE_INDENT, line_spacing=1.5, space_after=Pt(0))
            set_run_format(p.runs[0], size=FONT_SIZE_DEFAULT, italic=True)
            processed_indices.add(i)
        elif any(l.strip().lower().startswith("căn cứ") for l in body_lines[:i]):
             break
    if processed_indices: document.add_paragraph() # Thêm khoảng trống sau căn cứ

    # 4. Nội dung Kế hoạch
    for i, line in enumerate(body_lines):
        if i in processed_indices: continue
        stripped_line = line.strip()
        if not stripped_line: continue
        p = document.add_paragraph()

        # Logic định dạng cơ bản cho các đề mục kế hoạch
        is_part_roman = re.match(r'^([IVXLCDM]+)\.\s+', stripped_line.upper()) # I, II, III
        is_section_digit = re.match(r'^(\d+\.)\s+', stripped_line) # 1, 2, 3
        is_subsection_alpha = re.match(r'^[a-z]\)\s+', stripped_line) # a, b, c
        is_bullet = stripped_line.startswith("-") or stripped_line.startswith("+") or stripped_line.startswith("*") or stripped_line.startswith("•")

        align = WD_ALIGN_PARAGRAPH.JUSTIFY
        left_indent = Cm(0)
        first_indent = FIRST_LINE_INDENT if not (is_part_roman or is_section_digit or is_subsection_alpha or is_bullet) else Cm(0)
        is_bold = bool(is_part_roman or is_section_digit) # Mục lớn đậm
        is_italic = False
        size = FONT_SIZE_DEFAULT
        space_before = Pt(0)
        space_after = Pt(6)
        line_spacing = 1.5

        if is_part_roman:
            align = WD_ALIGN_PARAGRAPH.LEFT
            space_before = Pt(12)
            size = Pt(13)
        elif is_section_digit:
            align = WD_ALIGN_PARAGRAPH.LEFT
            left_indent = Cm(0.5)
            space_before = Pt(6)
        elif is_subsection_alpha:
             align = WD_ALIGN_PARAGRAPH.LEFT
             left_indent = Cm(1.0)
        elif is_bullet:
             align = WD_ALIGN_PARAGRAPH.LEFT
             left_indent = Cm(1.5)
             first_indent = Cm(-0.5) # Hanging indent


        set_paragraph_format(p, alignment=align, left_indent=left_indent, first_line_indent=first_indent, line_spacing=line_spacing, space_before=space_before, space_after=space_after)
        add_run_with_format(p, stripped_line, size=size, bold=is_bold, italic=is_italic)


    # 5. Chữ ký (Sử dụng khối chữ ký cơ bản)
    if 'signer_title' not in data: data['signer_title'] = "THỦ TRƯỞNG CƠ QUAN"
    if 'signer_name' not in data: data['signer_name'] = "[Họ và tên]"
    document.add_paragraph()
    format_signature_block(document, data)

    # 6. Nơi nhận (Sử dụng nơi nhận cơ bản)
    if 'recipients' not in data: data['recipients'] = ["- Như trên;", "- Lưu: VT, ...;"]
    format_recipient_list(document, data)

    print("Định dạng Kế hoạch hoàn tất.")