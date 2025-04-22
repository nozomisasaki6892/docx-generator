# formatters/cong_van.py
import time
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from utils import set_paragraph_format, set_run_format, add_run_with_format
try:
    from .common_elements import format_basic_header, format_signature_block, format_recipient_list
except ImportError:
    # Fallback nếu chạy độc lập
    def format_basic_header(document, data, doc_type): pass
    def format_signature_block(document, data): pass
    def format_recipient_list(document, data): pass

from config import FONT_SIZE_DEFAULT, FIRST_LINE_INDENT, FONT_SIZE_SMALL

def format(document, data):
    print("Bắt đầu định dạng Công văn (mặc định)...")
    title = data.get("title", "Công văn về việc ABC") # Thường dùng làm trích yếu
    body = data.get("body", "Nội dung công văn...")
    subject_prefix = "V/v:" # Tiền tố trích yếu của Công văn

    # 1. Header (Sử dụng header cơ bản)
    # Cập nhật doc_number trong data nếu chưa có tiền tố CV
    if 'doc_number' in data and not data['doc_number'].endswith('/CV-'):
         # Giả định mẫu số hiệu, cần chuẩn hóa tốt hơn
         match = re.match(r'Số:\s*(\d+)\s*/(.*)', data['doc_number'])
         if match:
             data['doc_number'] = f"Số: {match.group(1)}/CV-{match.group(2)}"
         else:
             data['doc_number'] = data.get("doc_number", "Số:       /CV-...")
    elif 'doc_number' not in data:
        data['doc_number'] = "Số:       /CV-..."

    format_basic_header(document, data, "CongVan")

    # 2. Trích yếu nội dung công văn (dưới số hiệu, căn trái)
    subject_content = title
    if title.lower().startswith("công văn"):
        subject_content = title.split(" ", 1)[1] # Bỏ chữ "Công văn"
    if subject_content.lower().startswith("về việc"):
         subject_content = subject_content.split(" ", 2)[-1] # Bỏ "Về việc"

    # Thêm đoạn trắng để căn chỉnh thẳng hàng với "Số:"
    num_spaces = data['doc_number'].find('/') + 1 if '/' in data['doc_number'] else 4
    trich_yeu_line = f"{subject_prefix}{' ' * num_spaces}{subject_content}"

    p_subject = document.add_paragraph(trich_yeu_line)
    # Căn trái, không thụt lề đầu dòng
    set_paragraph_format(p_subject, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=Pt(0), space_after=Pt(12), first_line_indent = Cm(0))
    # Có thể định dạng phần V/v đậm nếu muốn
    run_prefix = add_run_with_format(p_subject, f"{subject_prefix}{' ' * num_spaces}", size=FONT_SIZE_SMALL, bold=False)
    run_content = add_run_with_format(p_subject, subject_content, size=FONT_SIZE_SMALL, bold=False)


    # 3. Kính gửi
    recipient_main = data.get("recipient_main", "Kính gửi: [Tên cơ quan/cá nhân nhận]")
    p_kg = document.add_paragraph(recipient_main)
    set_paragraph_format(p_kg, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=Pt(6))
    set_run_format(p_kg.runs[0], size=FONT_SIZE_DEFAULT, bold=True)


    # 4. Nội dung công văn
    body_lines = body.split('\n')
    for line in body_lines:
        stripped_line = line.strip()
        # Bỏ qua các dòng trống hoặc chỉ chứa khoảng trắng
        if not stripped_line:
             # Có thể thêm dòng trống nếu muốn giữ khoảng cách
             # document.add_paragraph()
             continue

        p = document.add_paragraph()
        # Nội dung công văn thường căn đều 2 bên, thụt lề dòng đầu
        set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=Pt(6), first_line_indent=FIRST_LINE_INDENT, line_spacing=1.5)
        add_run_with_format(p, stripped_line, size=FONT_SIZE_DEFAULT)


    # 5. Lời kết (Nếu có)
    # Thường là đoạn cuối cùng của body


    # 6. Chữ ký (Sử dụng khối chữ ký cơ bản)
    # Cần xác định rõ người ký và chức vụ từ data
    if 'signer_title' not in data: data['signer_title'] = "CHỨC VỤ NGƯỜI KÝ"
    if 'signer_name' not in data: data['signer_name'] = "[Họ và tên]"
    document.add_paragraph() # Khoảng cách trước chữ ký
    format_signature_block(document, data)

    # 7. Nơi nhận (Sử dụng nơi nhận cơ bản)
    if 'recipients' not in data:
        recipients_default = []
        if recipient_main and not recipient_main.startswith("Kính gửi:"):
             recipients_default.append(f"- {recipient_main};") # Thêm nơi nhận chính nếu có
        else:
             recipients_default.append("- Như trên;")
        recipients_default.append("- Lưu: VT, ...;")
        data['recipients'] = recipients_default

    format_recipient_list(document, data)

    print("Định dạng Công văn hoàn tất.")