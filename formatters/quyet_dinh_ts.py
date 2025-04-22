# formatters/quyet_dinh_ts.py
import re
import time
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from utils import set_paragraph_format, set_run_format, add_run_with_format
try:
    # Dùng header, signature như Quyết định hành chính thông thường
    from .common_elements import format_basic_header, format_signature_block, format_recipient_list
except ImportError:
    def format_basic_header(document, data, doc_type): pass
    def format_signature_block(document, data): pass
    def format_recipient_list(document, data): pass

from config import FONT_SIZE_DEFAULT, FONT_SIZE_TITLE, FIRST_LINE_INDENT

def format(document, data):
    print("Bắt đầu định dạng Quyết định Tuyển sinh...")
    title = data.get("title", "Quyết định về việc Công nhận thí sinh trúng tuyển...")
    body = data.get("body", "Căn cứ Quy chế tuyển sinh...\nCăn cứ kết quả xét tuyển...\nXét đề nghị của Trưởng phòng Đào tạo...\nQUYẾT ĐỊNH:\nĐiều 1. Công nhận các thí sinh có tên trong danh sách kèm theo...\nĐiều 2. Các Ông/Bà Trưởng phòng... chịu trách nhiệm thi hành Quyết định này.\nĐiều 3. Quyết định này có hiệu lực kể từ ngày ký.")
    issuing_authority = data.get("issuing_org", "HIỆU TRƯỞNG TRƯỜNG ĐH XYZ").upper() # Người ký thường là Hiệu trưởng
    doc_type_label = "QUYẾT ĐỊNH"

    # 1. Header (Tên trường)
    # Cần đảm bảo data['issuing_org'] là tên trường
    data['issuing_org'] = data.get("university_name", issuing_authority)
    format_basic_header(document, data, "QuyetDinhTS")

    # 2. Tên loại và Trích yếu
    p_tenloai = document.add_paragraph(doc_type_label)
    set_paragraph_format(p_tenloai, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=Pt(12), space_after=Pt(6))
    add_run_with_format(p_tenloai, doc_type_label, size=FONT_SIZE_TITLE, bold=True, uppercase=True)

    trich_yeu_text = title.replace("Quyết định", "").strip()
    if trich_yeu_text.lower().startswith("về việc"):
        trich_yeu_text = trich_yeu_text.split(" ", 2)[-1]

    p_trichyeu = document.add_paragraph(f"Về việc {trich_yeu_text}")
    set_paragraph_format(p_trichyeu, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(12))
    add_run_with_format(p_trichyeu, f"Về việc {trich_yeu_text}", size=Pt(14), bold=True)


    # 3. Thẩm quyền ban hành (HIỆU TRƯỞNG)
    signer_title_upper = data.get("signer_title", "HIỆU TRƯỞNG").upper()
    p_authority = document.add_paragraph(signer_title_upper)
    set_paragraph_format(p_authority, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(6))
    add_run_with_format(p_authority, signer_title_upper, size=FONT_SIZE_DEFAULT, bold=True, uppercase=True)


    # 4. Căn cứ ban hành
    body_lines = body.split('\n')
    processed_indices = set()

    for i, line in enumerate(body_lines):
        stripped_line = line.strip()
        if not stripped_line: continue
        if stripped_line.lower().startswith("căn cứ") or stripped_line.lower().startswith("xét đề nghị"):
            p = document.add_paragraph(stripped_line)
            set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line_indent=FIRST_LINE_INDENT, line_spacing=1.5, space_after=Pt(0))
            add_run_with_format(p.runs[0], size=FONT_SIZE_DEFAULT, italic=True)
            processed_indices.add(i)
        elif any(l.strip().lower().startswith("căn cứ") or l.strip().lower().startswith("xét đề nghị") for l in body_lines[:i]):
            break # Dừng khi hết căn cứ

    # 5. Phần QUYẾT ĐỊNH:
    added_qd_label = False
    quyet_dinh_intro_index = -1
    for i, line in enumerate(body_lines):
         if i in processed_indices: continue
         stripped_line = line.strip()
         if "QUYẾT ĐỊNH:" in stripped_line.upper():
             if stripped_line.upper() == "QUYẾT ĐỊNH:":
                 p_qd_label = document.add_paragraph("QUYẾT ĐỊNH:")
                 set_paragraph_format(p_qd_label, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=Pt(12), space_after=Pt(12))
                 add_run_with_format(p_qd_label, "QUYẾT ĐỊNH:", size=FONT_SIZE_DEFAULT, bold=True, uppercase=True)
                 processed_indices.add(i)
                 added_qd_label = True
             else:
                 p_qd_intro = document.add_paragraph(stripped_line)
                 set_paragraph_format(p_qd_intro, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=Pt(12), space_after=Pt(12))
                 add_run_with_format(p_qd_intro, stripped_line, size=FONT_SIZE_DEFAULT, bold=True)
                 processed_indices.add(i)
                 added_qd_label = True
             quyet_dinh_intro_index = i
             break

    if processed_indices and quyet_dinh_intro_index == -1:
         document.add_paragraph()

    if not added_qd_label:
        p_qd_label = document.add_paragraph("QUYẾT ĐỊNH:")
        set_paragraph_format(p_qd_label, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=Pt(12), space_after=Pt(12))
        add_run_with_format(p_qd_label, "QUYẾT ĐỊNH:", size=FONT_SIZE_DEFAULT, bold=True, uppercase=True)


    # 6. Nội dung Quyết định (Điều, Khoản, Điểm)
    # Thường có các Điều về công nhận danh sách, trách nhiệm thi hành, hiệu lực
    for i, line in enumerate(body_lines):
        if i in processed_indices: continue
        stripped_line = line.strip()
        if not stripped_line: continue

        if "QUYẾT ĐỊNH:" in stripped_line.upper() and not added_qd_label:
            continue

        p = document.add_paragraph()
        is_dieu = stripped_line.upper().startswith("ĐIỀU")
        is_khoan = re.match(r'^\d+\.\s+', stripped_line)
        is_diem = re.match(r'^[a-z]\)\s+', stripped_line)

        align = WD_ALIGN_PARAGRAPH.JUSTIFY
        left_indent = Cm(0)
        first_indent = FIRST_LINE_INDENT if not (is_dieu or is_khoan or is_diem) else Cm(0)
        is_bold = False
        size = FONT_SIZE_DEFAULT

        if is_dieu:
            align = WD_ALIGN_PARAGRAPH.LEFT
            is_bold = True
            space_before = Pt(6)
        elif is_khoan:
            left_indent = Cm(0.5)
        elif is_diem:
            left_indent = Cm(1.0)

        set_paragraph_format(p, alignment=align, left_indent=left_indent, first_line_indent=first_indent, line_spacing=1.5, space_after=Pt(6))
        add_run_with_format(p, stripped_line, size=size, bold=is_bold)

    # 7. Chữ ký (Hiệu trưởng)
    if 'signer_title' not in data: data['signer_title'] = signer_title_upper
    if 'signer_name' not in data: data['signer_name'] = "[Họ và tên Hiệu trưởng]"
    format_signature_block(document, data)

    # 8. Nơi nhận
    if 'recipients' not in data: data['recipients'] = ["- Như Điều 2;", "- Ban Giám hiệu;", "- Lưu: VT, P.ĐT."]
    format_recipient_list(document, data)

    print("Định dạng Quyết định Tuyển sinh hoàn tất.")