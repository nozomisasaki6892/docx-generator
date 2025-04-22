# formatters/thong_cao.py
import time
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from utils import set_paragraph_format, set_run_format, add_run_with_format, add_centered_text
try:
    # Thông cáo báo chí có thể có header, signature đơn giản hơn
    from .common_elements import format_signature_block, format_recipient_list # Có thể dùng signature, recipients
except ImportError:
     def format_signature_block(document, data): pass
     def format_recipient_list(document, data): pass

from config import FONT_SIZE_DEFAULT, FONT_SIZE_TITLE, FIRST_LINE_INDENT, FONT_SIZE_HEADER

def format(document, data):
    print("Bắt đầu định dạng Thông cáo báo chí...")
    title = data.get("title", "THÔNG CÁO BÁO CHÍ").upper()
    subject = data.get("subject", "Về việc Tổ chức sự kiện/Ra mắt sản phẩm...")
    body = data.get("body", "[Tên cơ quan/tổ chức] trân trọng thông báo...\nNội dung chính...\nThông tin chi tiết...\nLiên hệ...")
    issuing_org = data.get("issuing_org", "TÊN CƠ QUAN/TỔ CHỨC").upper()
    issuing_location = data.get("issuing_location", "Hà Nội")
    issuing_date_str = data.get("issuing_date", time.strftime(f"ngày %d tháng %m năm %Y"))


    # 1. Tên cơ quan (Căn giữa hoặc trái)
    add_centered_text(document, issuing_org, size=FONT_SIZE_HEADER, bold=True, space_after=12)


    # 2. Tiêu đề Thông cáo báo chí
    add_centered_text(document, title, size=FONT_SIZE_TITLE, bold=True, space_before=12, space_after=6)
    # Ngày tháng ban hành dưới tiêu đề
    add_centered_text(document, f"{issuing_location}, {issuing_date_str}", size=FONT_SIZE_DEFAULT, italic=True, space_after=18)


    # 3. Trích yếu/Tiêu đề phụ (Căn giữa, đậm)
    add_centered_text(document, subject, size=Pt(14), bold=True, space_after=12)


    # 4. Nội dung Thông cáo
    body_lines = body.split('\n')
    for line in body_lines:
        stripped_line = line.strip()
        if not stripped_line: continue
        p = document.add_paragraph()

        # Nội dung TCBC thường căn đều
        is_contact_info = "liên hệ" in stripped_line.lower() or "email" in stripped_line.lower() or "website" in stripped_line.lower() or "điện thoại" in stripped_line.lower()

        align = WD_ALIGN_PARAGRAPH.JUSTIFY
        first_indent = FIRST_LINE_INDENT
        is_bold = False
        is_italic = False
        size = FONT_SIZE_DEFAULT
        space_after = Pt(6)
        line_spacing = 1.5

        if is_contact_info:
            align = WD_ALIGN_PARAGRAPH.LEFT # Thông tin liên hệ căn trái
            first_indent = Cm(0)
            space_before = Pt(12) # Cách đoạn trước liên hệ
        elif stripped_line.endswith("###") or stripped_line.endswith("***"): # Dấu hiệu kết thúc
            align = WD_ALIGN_PARAGRAPH.CENTER
            first_indent = Cm(0)
            is_bold = True
            space_before = Pt(12)

        set_paragraph_format(p, alignment=align, space_after=space_after, first_line_indent=first_indent, line_spacing=line_spacing)
        add_run_with_format(p, stripped_line, size=size, bold=is_bold, italic=is_italic)


    # 5. Chữ ký (Nếu có - thường không có chữ ký trang trọng)
    # Có thể chỉ ghi tên đơn vị phát hành ở cuối
    if data.get('signer_title') or data.get('signer_name'):
         document.add_paragraph()
         # Dùng signature block đơn giản hoặc chỉ ghi tên
         # format_signature_block(document, data)
         p_issuer_end = document.add_paragraph(issuing_org)
         set_paragraph_format(p_issuer_end, alignment=WD_ALIGN_PARAGRAPH.RIGHT, space_before=Pt(12))
         set_run_format(p_issuer_end.runs[0], size=FONT_SIZE_DEFAULT, bold=True)


    print("Định dạng Thông cáo báo chí hoàn tất.")