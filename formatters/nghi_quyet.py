# formatters/nghi_quyet.py
import re
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from utils import set_paragraph_format, set_run_format, add_run_with_format
try:
    from .common_elements import format_basic_header, format_signature_block, format_recipient_list
except ImportError:
    from common_elements import format_basic_header, format_signature_block, format_recipient_list

from config import FONT_SIZE_DEFAULT, FONT_SIZE_TITLE, FIRST_LINE_INDENT

def format(document, data):
    print("Bắt đầu định dạng Nghị quyết (cá biệt)...")
    title = data.get("title", "Nghị quyết về việc ABC")
    body = data.get("body", "Căn cứ...\nXét Tờ trình...\nHội đồng nhân dân... QUYẾT NGHỊ:\nĐiều 1...\nĐiều 2...")
    issuing_authority = data.get("issuing_org", "HỘI ĐỒNG NHÂN DÂN TỈNH/THÀNH PHỐ").upper()

    # 1. Header (Của cơ quan ban hành NQ cá biệt)
    data['issuing_org'] = issuing_authority # Cập nhật data để header dùng đúng tên
    format_basic_header(document, data, "NghiQuyet")

    # 2. Tên loại
    p_tenloai = document.add_paragraph("NGHỊ QUYẾT")
    set_paragraph_format(p_tenloai, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=Pt(12), space_after=Pt(6))
    add_run_with_format(p_tenloai, "NGHỊ QUYẾT", size=FONT_SIZE_TITLE, bold=True, uppercase=True)

    # 3. Trích yếu
    trich_yeu_text = title.replace("Nghị quyết", "").strip()
    # Bỏ các tiền tố thường gặp
    if trich_yeu_text.lower().startswith("về việc"):
        trich_yeu_text = trich_yeu_text.split(" ", 2)[-1]
    elif trich_yeu_text.lower().startswith("ban hành"):
         trich_yeu_text = trich_yeu_text.split(" ", 1)[-1]

    p_trichyeu = document.add_paragraph(trich_yeu_text)
    set_paragraph_format(p_trichyeu, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(6))
    add_run_with_format(p_trichyeu, trich_yeu_text, size=Pt(14), bold=True)
    p_line_ty = document.add_paragraph("-" * 15)
    set_paragraph_format(p_line_ty, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(12))


    # 4. Cơ quan ban hành (Không lặp lại vì đã có trong header)

    # 5. Nội dung (Căn cứ, Phần QUYẾT NGHỊ:, các Điều)
    body_lines = body.split('\n')
    processed_indices = set()
    added_quyet_nghi_label = False

    # Xử lý Căn cứ trước
    for i, line in enumerate(body_lines):
        stripped_line = line.strip()
        if not stripped_line: continue

        if stripped_line.lower().startswith("căn cứ") or stripped_line.lower().startswith("xét tờ trình"):
            p = document.add_paragraph(stripped_line)
            set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line_indent=FIRST_LINE_INDENT, line_spacing=1.5, space_after=Pt(0))
            add_run_with_format(p, stripped_line, size=FONT_SIZE_DEFAULT, italic=True)
            processed_indices.add(i)
        elif any(l.strip().lower().startswith("căn cứ") for l in body_lines[:i]):
            break # Dừng khi hết căn cứ

    # Tìm câu dẫn quyết nghị nếu có
    quyet_nghi_intro_index = -1
    for i, line in enumerate(body_lines):
         if i in processed_indices: continue
         stripped_line = line.strip()
         # Tìm dòng chứa tên cơ quan + QUYẾT NGHỊ:
         if issuing_authority in stripped_line.upper() and "QUYẾT NGHỊ:" in stripped_line.upper():
             p_intro_qn = document.add_paragraph(stripped_line)
             set_paragraph_format(p_intro_qn, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=Pt(6), space_after=Pt(12))
             add_run_with_format(p_intro_qn, stripped_line, size=FONT_SIZE_DEFAULT, bold=True)
             processed_indices.add(i)
             added_quyet_nghi_label = True # Đã có dòng này, không cần thêm label riêng
             quyet_nghi_intro_index = i
             break

    if processed_indices and quyet_nghi_intro_index == -1: # Thêm khoảng trắng sau căn cứ nếu không có câu dẫn
        document.add_paragraph()

    # Thêm nhãn QUYẾT NGHỊ: nếu chưa có trong câu dẫn
    if not added_quyet_nghi_label:
        p_qn_label = document.add_paragraph("QUYẾT NGHỊ:")
        set_paragraph_format(p_qn_label, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=Pt(6), space_after=Pt(12))
        add_run_with_format(p_qn_label, "QUYẾT NGHỊ:", size=FONT_SIZE_DEFAULT, bold=True, uppercase=True)

    # Xử lý phần Điều khoản và nội dung còn lại
    for i, line in enumerate(body_lines):
        if i in processed_indices: continue # Bỏ qua các dòng đã xử lý
        stripped_line = line.strip()
        if not stripped_line: continue

        # Bỏ qua nếu dòng này chính là QUYẾT NGHỊ: trong body gốc và ta đã thêm label
        if "QUYẾT NGHỊ:" in stripped_line.upper() and not added_quyet_nghi_label:
            continue

        p = document.add_paragraph()
        is_dieu = stripped_line.upper().startswith("ĐIỀU")
        is_khoan = re.match(r'^\d+\.\s+', stripped_line)
        is_diem = re.match(r'^[a-z]\)\s+', stripped_line)

        align = WD_ALIGN_PARAGRAPH.JUSTIFY
        left_indent = Cm(0)
        first_indent = FIRST_LINE_INDENT if not (is_dieu or is_khoan or is_diem) else Cm(0)
        is_bold = False
        size = FONT_SIZE_DEFAULT

        if is_dieu:
            align = WD_ALIGN_PARAGRAPH.LEFT
            is_bold = True
            space_before = Pt(6)
        elif is_khoan:
            left_indent = Cm(0.5)
        elif is_diem:
            left_indent = Cm(1.0)

        set_paragraph_format(p, alignment=align, left_indent=left_indent, first_line_indent=first_indent, line_spacing=1.5, space_after=Pt(6))
        add_run_with_format(p, stripped_line, size=size, bold=is_bold)

    # 6. Chữ ký
    # Chữ ký NQ cá biệt thường có Thẩm quyền ký (TM. HĐND...)
    authority_signer = data.get("authority_signer")
    if not authority_signer:
        authority_signer = f"TM. {issuing_authority}"
    data['authority_signer'] = authority_signer

    if 'signer_title' not in data: data['signer_title'] = "CHỦ TỊCH" # Ví dụ Chủ tịch HĐND
    if 'signer_name' not in data: data['signer_name'] = "[Họ và tên]"
    format_signature_block(document, data)

    # 7. Nơi nhận
    format_recipient_list(document, data)

    print("Định dạng Nghị quyết (cá biệt) hoàn tất.")