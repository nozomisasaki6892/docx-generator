# formatters/giay_uy_quyen.py
import re
import time
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from utils import set_paragraph_format, set_run_format, add_run_with_format, add_centered_text
from config import FONT_SIZE_DEFAULT, FONT_SIZE_TITLE, FONT_SIZE_SIGNATURE, FONT_SIZE_SIGNER_NAME, FONT_SIZE_PLACE_DATE, FONT_SIZE_HEADER, FIRST_LINE_INDENT

def format_party_info(document, party_label, party_data):
    # Hàm helper để định dạng thông tin một bên
    p_label = document.add_paragraph(party_label)
    set_paragraph_format(p_label, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=Pt(6))
    set_run_format(p_label.runs[0], size=FONT_SIZE_DEFAULT, bold=True)

    default_info = {
        "name": "[Họ và tên]",
        "dob": "[Ngày sinh]",
        "id_card": "[Số CMND/CCCD]",
        "id_date": "[Ngày cấp]",
        "id_place": "[Nơi cấp]",
        "address": "[Địa chỉ thường trú]",
        "phone": "[Số điện thoại]"
    }

    for key, label in [("name", "Họ và tên"), ("dob", "Sinh ngày"), ("id_card", "Số CMND/CCCD"),
                       ("id_date", "Ngày cấp"), ("id_place", "Nơi cấp"), ("address", "Địa chỉ thường trú"),
                       ("phone", "Điện thoại liên hệ")]:
        p_info = document.add_paragraph()
        set_paragraph_format(p_info, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=Pt(3), left_indent=Cm(0.5))
        add_run_with_format(p_info, f"- {label}: ", size=FONT_SIZE_DEFAULT)
        add_run_with_format(p_info, party_data.get(key, default_info[key]), size=FONT_SIZE_DEFAULT)


def format(document, data):
    print("Bắt đầu định dạng Giấy ủy quyền...")
    title = data.get("title", "GIẤY ỦY QUYỀN").upper()
    issuing_location = data.get("issuing_location", "Hà Nội")
    current_date_str = time.strftime(f"ngày %d tháng %m năm %Y")

    # Lấy thông tin các bên từ data (dạng dict)
    authorizer_data = data.get("authorizer", {}) # Bên ủy quyền
    authorized_data = data.get("authorized", {}) # Bên được ủy quyền
    authorization_content = data.get("authorization_content", "Bằng Giấy này, Bên A ủy quyền cho Bên B thực hiện các công việc sau:\n- ...\n- ...")
    authorization_scope = data.get("authorization_scope", "Phạm vi ủy quyền:\n- ...")
    authorization_duration = data.get("authorization_duration", "Thời hạn ủy quyền: Kể từ ngày ... đến ngày ... (hoặc cho đến khi công việc hoàn thành).")


    # 1. Quốc hiệu, Tiêu ngữ
    add_centered_text(document, "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM", size=FONT_SIZE_HEADER, bold=True, space_after=0)
    add_centered_text(document, "Độc lập - Tự do - Hạnh phúc", size=Pt(13), bold=True, space_after=18)

    # 2. Tên Giấy ủy quyền
    add_centered_text(document, title, size=FONT_SIZE_TITLE, bold=True, space_before=12, space_after=12)

    # 3. Địa điểm, Ngày tháng
    p_place_date = document.add_paragraph(f"{issuing_location}, {current_date_str}")
    set_paragraph_format(p_place_date, alignment=WD_ALIGN_PARAGRAPH.RIGHT, space_after=Pt(12))
    set_run_format(p_place_date.runs[0], size=FONT_SIZE_DEFAULT, italic=True)

    # 4. Thông tin Bên ủy quyền (Bên A)
    format_party_info(document, "BÊN ỦY QUYỀN (Gọi tắt là Bên A):", authorizer_data)
    document.add_paragraph()

    # 5. Thông tin Bên được ủy quyền (Bên B)
    format_party_info(document, "BÊN ĐƯỢC ỦY QUYỀN (Gọi tắt là Bên B):", authorized_data)
    document.add_paragraph()

    # 6. Nội dung ủy quyền
    p_content_label = document.add_paragraph("NỘI DUNG ỦY QUYỀN:")
    set_paragraph_format(p_content_label, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=Pt(6))
    set_run_format(p_content_label.runs[0], size=FONT_SIZE_DEFAULT, bold=True)
    content_lines = authorization_content.split('\n')
    for line in content_lines:
        stripped_line = line.strip()
        if not stripped_line: continue
        p = document.add_paragraph()
        set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=Pt(6), first_line_indent=FIRST_LINE_INDENT, line_spacing=1.5)
        add_run_with_format(p, stripped_line, size=FONT_SIZE_DEFAULT)

    # 7. Phạm vi ủy quyền
    scope_lines = authorization_scope.split('\n')
    for line in scope_lines:
        stripped_line = line.strip()
        if not stripped_line: continue
        p = document.add_paragraph()
        set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=Pt(6), first_line_indent=FIRST_LINE_INDENT, line_spacing=1.5)
        add_run_with_format(p, stripped_line, size=FONT_SIZE_DEFAULT)


    # 8. Thời hạn ủy quyền
    duration_lines = authorization_duration.split('\n')
    for line in duration_lines:
        stripped_line = line.strip()
        if not stripped_line: continue
        p = document.add_paragraph()
        set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=Pt(6), first_line_indent=FIRST_LINE_INDENT, line_spacing=1.5)
        add_run_with_format(p, stripped_line, size=FONT_SIZE_DEFAULT)

    # 9. Cam kết
    p_commit = document.add_paragraph("Hai bên cam kết thực hiện đúng các nội dung đã nêu trong Giấy ủy quyền này. Mọi tranh chấp phát sinh (nếu có) sẽ được giải quyết trên tinh thần hợp tác, thương lượng. Nếu không tự giải quyết được, vụ việc sẽ được đưa ra Tòa án có thẩm quyền để giải quyết theo quy định của pháp luật.")
    set_paragraph_format(p_commit, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_before=Pt(12), space_after=Pt(12), first_line_indent=FIRST_LINE_INDENT, line_spacing=1.5)
    set_run_format(p_commit.runs[0], size=FONT_SIZE_DEFAULT)

    p_validity = document.add_paragraph("Giấy ủy quyền này được lập thành 02 (hai) bản, mỗi bên giữ 01 (một) bản và có giá trị pháp lý như nhau.")
    set_paragraph_format(p_validity, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=Pt(18), first_line_indent=FIRST_LINE_INDENT, line_spacing=1.5)
    set_run_format(p_validity.runs[0], size=FONT_SIZE_DEFAULT)


    # 10. Chữ ký hai bên
    table = document.add_table(rows=1, cols=2)
    table.autofit = False
    table.allow_autofit = False
    table.columns[0].width = Inches(3.0)
    table.columns[1].width = Inches(3.0)

    cell_a = table.cell(0, 0)
    cell_b = table.cell(0, 1)
    cell_a._element.clear_content()
    cell_b._element.clear_content()

    p_a_title = cell_a.add_paragraph("BÊN ỦY QUYỀN")
    set_paragraph_format(p_a_title, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(0))
    set_run_format(p_a_title.runs[0], size=FONT_SIZE_SIGNATURE, bold=True)
    p_a_note = cell_a.add_paragraph("(Ký và ghi rõ họ tên)")
    set_paragraph_format(p_a_note, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(6))
    set_run_format(p_a_note.runs[0], size=Pt(11), italic=True)
    cell_a.add_paragraph("\n\n\n\n")
    p_a_name = cell_a.add_paragraph(authorizer_data.get("name", "[Họ tên Bên A]"))
    set_paragraph_format(p_a_name, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_run_format(p_a_name.runs[0], size=FONT_SIZE_SIGNER_NAME, bold=True)

    p_b_title = cell_b.add_paragraph("BÊN ĐƯỢC ỦY QUYỀN")
    set_paragraph_format(p_b_title, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(0))
    set_run_format(p_b_title.runs[0], size=FONT_SIZE_SIGNATURE, bold=True)
    p_b_note = cell_b.add_paragraph("(Ký và ghi rõ họ tên)")
    set_paragraph_format(p_b_note, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(6))
    set_run_format(p_b_note.runs[0], size=Pt(11), italic=True)
    cell_b.add_paragraph("\n\n\n\n")
    p_b_name = cell_b.add_paragraph(authorized_data.get("name", "[Họ tên Bên B]"))
    set_paragraph_format(p_b_name, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_run_format(p_b_name.runs[0], size=FONT_SIZE_SIGNER_NAME, bold=True)


    print("Định dạng Giấy ủy quyền hoàn tất.")