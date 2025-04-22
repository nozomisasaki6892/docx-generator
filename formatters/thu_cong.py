# formatters/thu_cong.py
import time
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from utils import set_paragraph_format, set_run_format, add_run_with_format
try:
    # Dùng header, signature của văn bản hành chính
    from .common_elements import format_basic_header, format_signature_block
except ImportError:
    def format_basic_header(document, data, doc_type): pass
    def format_signature_block(document, data): pass

from config import FONT_SIZE_DEFAULT, FONT_SIZE_TITLE, FIRST_LINE_INDENT

def format(document, data):
    print("Bắt đầu định dạng Thư công...")
    title = data.get("title", "THƯ CÔNG") # Hoặc không có tiêu đề loại
    subject = data.get("subject", "V/v: Chúc mừng/Cảm ơn/Trao đổi công việc...")
    body = data.get("body", "Kính gửi: [Ông/Bà/Ngài]...\n[Nội dung thư]...\nTrân trọng.")
    recipient = data.get("recipient_main", "Kính gửi: [Tên đầy đủ và Chức vụ người nhận]")


    # 1. Header (Cơ quan/Cá nhân gửi thư)
    format_basic_header(document, data, "ThuCong")


    # 2. Tiêu đề (nếu có) và Trích yếu
    if title.strip(): # Nếu title không rỗng
         p_title = document.add_paragraph(title.upper())
         set_paragraph_format(p_title, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=Pt(12), space_after=Pt(6))
         set_run_format(p_title.runs[0], size=FONT_SIZE_TITLE, bold=True)

    p_subject = document.add_paragraph(subject)
    set_paragraph_format(p_subject, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(12))
    set_run_format(p_subject.runs[0], size=Pt(14), bold=True)


    # 3. Kính gửi
    p_kg = document.add_paragraph(recipient)
    set_paragraph_format(p_kg, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=Pt(12))
    set_run_format(p_kg.runs[0], size=FONT_SIZE_DEFAULT, bold=True)


    # 4. Nội dung Thư
    body_lines = body.split('\n')
    for line in body_lines:
        stripped_line = line.strip()
        if not stripped_line:
            document.add_paragraph() # Giữ dòng trống
            continue

        p = document.add_paragraph()
        # Nội dung thư thường căn đều, thụt lề dòng đầu
        is_ending = stripped_line.lower().startswith("trân trọng") or stripped_line.lower().startswith("kính thư")

        align = WD_ALIGN_PARAGRAPH.JUSTIFY
        first_indent = FIRST_LINE_INDENT
        is_bold = False
        is_italic = False
        size = FONT_SIZE_DEFAULT
        line_spacing = 1.5

        if is_ending:
            align = WD_ALIGN_PARAGRAPH.RIGHT # Lời kết cuối thư căn phải
            first_indent = Cm(0)
            is_italic = True # Nghiêng
            space_before = Pt(12)
        else:
             space_after = Pt(6)

        set_paragraph_format(p, alignment=align, space_after=space_after if not is_ending else Pt(0), first_line_indent=first_indent, line_spacing=line_spacing)
        add_run_with_format(p, stripped_line, size=size, bold=is_bold, italic=is_italic)


    # 5. Chữ ký (Người gửi thư)
    if 'signer_title' not in data: data['signer_title'] = "CHỨC VỤ NGƯỜI GỬI"
    if 'signer_name' not in data: data['signer_name'] = "[Họ và tên]"
    # Giảm khoảng cách trước chữ ký vì lời kết đã có space_before
    sig_paragraph = document.add_paragraph()
    set_paragraph_format(sig_paragraph, alignment=WD_ALIGN_PARAGRAPH.RIGHT, space_before=Pt(0), space_after=Pt(0), line_spacing=1.0)
    run_title = add_run_with_format(sig_paragraph, data['signer_title'].upper() + "\n", size=FONT_SIZE_DEFAULT, bold=True)
    sig_paragraph.add_run("\n\n\n\n")
    run_name = add_run_with_format(sig_paragraph, data['signer_name'], size=FONT_SIZE_DEFAULT, bold=True)


    # Thư công thường không có Nơi nhận kiểu NĐ30

    print("Định dạng Thư công hoàn tất.")