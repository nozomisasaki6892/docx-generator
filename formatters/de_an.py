# formatters/de_an.py
import re
import time
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from utils import set_paragraph_format, set_run_format, add_run_with_format
try:
    from .common_elements import format_basic_header, format_signature_block, format_recipient_list
except ImportError:
    # Fallback nếu chạy độc lập
    def format_basic_header(document, data, doc_type): pass
    def format_signature_block(document, data): pass
    def format_recipient_list(document, data): pass

from config import FONT_SIZE_DEFAULT, FONT_SIZE_TITLE, FIRST_LINE_INDENT

def format(document, data):
    print("Bắt đầu định dạng Đề án...")
    title = data.get("title", "Đề án Phát triển ứng dụng XYZ")
    body = data.get("body", "Phần I: SỰ CẦN THIẾT VÀ CĂN CỨ PHÁP LÝ\nPhần II: MỤC TIÊU, NHIỆM VỤ VÀ GIẢI PHÁP\nPhần III: KINH PHÍ THỰC HIỆN\nPhần IV: TỔ CHỨC THỰC HIỆN...")
    doc_type_label = "ĐỀ ÁN"
    issuing_authority = data.get("issuing_org", "CƠ QUAN CHỦ TRÌ").upper()

    # 1. Header (Sử dụng header cơ bản)
    # Có thể dùng header hoặc không tùy yêu cầu trình bày đề án
    format_basic_header(document, data, "DeAn")

    # 2. Tên Đề án
    p_title = document.add_paragraph(doc_type_label)
    set_paragraph_format(p_title, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=Pt(12), space_after=Pt(6))
    set_run_format(p_title.runs[0], size=FONT_SIZE_TITLE, bold=True)

    # Tên đề án cụ thể
    project_name = title.replace("Đề án", "").strip()
    p_name = document.add_paragraph(project_name.upper()) # Tên đề án thường IN HOA
    set_paragraph_format(p_name, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(12))
    set_run_format(p_name.runs[0], size=Pt(14), bold=True)


    # 3. Lời mở đầu / Tóm tắt (Nếu có)
    summary = data.get("summary", None)
    if summary:
        p_sum = document.add_paragraph(summary)
        set_paragraph_format(p_sum, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=Pt(12), first_line_indent=FIRST_LINE_INDENT, line_spacing=1.5)
        set_run_format(p_sum.runs[0], size=FONT_SIZE_DEFAULT, italic=True)


    # 4. Nội dung Đề án
    body_lines = body.split('\n')
    for line in body_lines:
        stripped_line = line.strip()
        if not stripped_line: continue
        p = document.add_paragraph()

        # Logic định dạng cơ bản cho các đề mục đề án
        is_part_roman = re.match(r'^(PHẦN\s+[IVXLCDM]+)\.?\s+', stripped_line.upper())
        is_section_digit = re.match(r'^([IVXLCDM]+)\.\s+', stripped_line.upper()) # Mục La Mã trong Phần
        is_subsection_digit = re.match(r'^(\d+\.)\s+', stripped_line) # Mục 1, 2, 3
        is_subsubsection_alpha = re.match(r'^[a-z]\)\s+', stripped_line) # Mục a, b, c
        is_bullet = stripped_line.startswith("-") or stripped_line.startswith("+") or stripped_line.startswith("*")

        align = WD_ALIGN_PARAGRAPH.JUSTIFY
        left_indent = Cm(0)
        first_indent = FIRST_LINE_INDENT if not (is_part_roman or is_section_digit or is_subsection_digit or is_subsubsection_alpha or is_bullet) else Cm(0)
        is_bold = bool(is_part_roman or is_section_digit or is_subsection_digit) # Mục lớn đậm
        is_italic = False
        size = FONT_SIZE_DEFAULT
        space_before = Pt(0)
        space_after = Pt(6)
        line_spacing = 1.5

        if is_part_roman:
            align = WD_ALIGN_PARAGRAPH.CENTER
            space_before = Pt(18)
            size = Pt(14) # Phần lớn, đậm, giữa
        elif is_section_digit:
            align = WD_ALIGN_PARAGRAPH.LEFT
            space_before = Pt(12)
            size = Pt(13) # Mục La Mã, đậm, trái
        elif is_subsection_digit:
            align = WD_ALIGN_PARAGRAPH.LEFT
            left_indent = Cm(0.5)
            space_before = Pt(6)
        elif is_subsubsection_alpha:
             align = WD_ALIGN_PARAGRAPH.LEFT
             left_indent = Cm(1.0)
        elif is_bullet:
             align = WD_ALIGN_PARAGRAPH.LEFT
             left_indent = Cm(1.5)
             first_indent = Cm(-0.5) # Hanging indent

        set_paragraph_format(p, alignment=align, left_indent=left_indent, first_line_indent=first_indent, line_spacing=line_spacing, space_before=space_before, space_after=space_after)
        add_run_with_format(p, stripped_line, size=size, bold=is_bold, italic=is_italic)


    # 5. Chữ ký (Thường là Thủ trưởng cơ quan chủ trì)
    if 'signer_title' not in data: data['signer_title'] = "THỦ TRƯỞNG CƠ QUAN CHỦ TRÌ"
    if 'signer_name' not in data: data['signer_name'] = "[Họ và tên]"
    document.add_paragraph()
    format_signature_block(document, data)

    # 6. Nơi nhận (Nếu đề án được ban hành như văn bản)
    if data.get('recipients'):
         format_recipient_list(document, data)

    print("Định dạng Đề án hoàn tất.")