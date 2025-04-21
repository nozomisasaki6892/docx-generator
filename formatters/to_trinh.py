# formatters/to_trinh.py
import re
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from utils import set_paragraph_format, set_run_format, add_run_with_format
try:
    from .common_elements import format_basic_header, format_signature_block, format_recipient_list
except ImportError:
    from common_elements import format_basic_header, format_signature_block, format_recipient_list
from config import FONT_SIZE_DEFAULT, FONT_SIZE_TITLE, FIRST_LINE_INDENT

def format(document, data):
    print("Bắt đầu định dạng Tờ trình...")
    title = data.get("title", "Tờ trình về việc ABC")
    body = data.get("body", "Nội dung tờ trình...")
    recipients_to = data.get("recipients_to", "Kính gửi: [Lãnh đạo/Cấp trên]") # Người nhận Tờ trình
    issuing_org = data.get("issuing_org", "TÊN ĐƠN VỊ TRÌNH").upper()

    # 1. Header (Tên đơn vị trình bên trái)
    data['issuing_org'] = issuing_org
    format_basic_header(document, data, "ToTrinh") # Header căn trái CQBH

    # 2. Tên loại TỜ TRÌNH
    p_tenloai = document.add_paragraph("TỜ TRÌNH")
    set_paragraph_format(p_tenloai, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=Pt(12), space_after=Pt(6))
    add_run_with_format(p_tenloai, "TỜ TRÌNH", size=FONT_SIZE_TITLE, bold=True, uppercase=True)

    # 3. Trích yếu
    tt_title = title.replace("Tờ trình", "").strip()
    p_title = document.add_paragraph(f"Về việc {tt_title}")
    set_paragraph_format(p_title, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(12))
    add_run_with_format(p_title, f"Về việc {tt_title}", size=Pt(14), bold=True)

    # 4. Kính gửi
    p_kg = document.add_paragraph(recipients_to)
    set_paragraph_format(p_kg, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(12)) # Kính gửi căn giữa
    add_run_with_format(p_kg, recipients_to, size=FONT_SIZE_DEFAULT, bold=True)

    # 5. Nội dung Tờ trình
    body_lines = body.split('\n')
    for line in body_lines:
        stripped_line = line.strip()
        if stripped_line:
            p = document.add_paragraph()
            # Nhận diện mục đánh số
            is_numbered_item = re.match(r'^\d+\.\s+', stripped_line)
            align = WD_ALIGN_PARAGRAPH.JUSTIFY
            left_indent = Cm(0.5) if is_numbered_item else Cm(0)
            first_indent = Cm(0) if is_numbered_item else FIRST_LINE_INDENT

            set_paragraph_format(p, alignment=align, left_indent=left_indent, first_line_indent=first_indent, line_spacing=1.5, space_after=Pt(6))
            add_run_with_format(p, stripped_line, size=FONT_SIZE_DEFAULT, bold=bool(is_numbered_item))

    # 6. Lời kết (Kính trình...)
    p_closing = document.add_paragraph(f"{issuing_org} kính trình [Lãnh đạo] xem xét, phê duyệt.")
    set_paragraph_format(p_closing, alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=FIRST_LINE_INDENT, space_before=Pt(6), space_after=Pt(6))
    add_run_with_format(p_closing, p_closing.text, size=FONT_SIZE_DEFAULT, italic=True)


    # 7. Chữ ký (Người đứng đầu đơn vị trình)
    format_signature_block(document, data)

    # 8. Nơi nhận
    if not data.get('recipients'):
        data['recipients'] = ["- Như trên;", "- Lưu: VT, [Đơn vị soạn]."]
    format_recipient_list(document, data)

    print("Định dạng Tờ trình hoàn tất.")