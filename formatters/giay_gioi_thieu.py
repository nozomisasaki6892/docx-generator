# formatters/giay_gioi_thieu.py
import time
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from utils import set_paragraph_format, set_run_format, add_run_with_format
try:
    from .common_elements import format_basic_header, format_signature_block # Chỉ dùng header và signature
except ImportError:
    def format_basic_header(document, data, doc_type): pass
    def format_signature_block(document, data): pass

from config import FONT_SIZE_DEFAULT, FONT_SIZE_TITLE, FIRST_LINE_INDENT, FONT_SIZE_SMALL

def format(document, data):
    print("Bắt đầu định dạng Giấy giới thiệu...")
    title = "GIẤY GIỚI THIỆU"
    body = data.get("body", "Trân trọng giới thiệu Ông/Bà:...\nChức vụ:...\nĐược cử đến:...\nĐể liên hệ giải quyết công việc về:...\nĐề nghị Quý cơ quan tạo điều kiện giúp đỡ Ông/Bà ... hoàn thành nhiệm vụ.\nGiấy giới thiệu có giá trị đến hết ngày .../.../......")
    recipient_org = data.get("recipient_org", "Kính gửi: [Tên cơ quan/đơn vị nơi đến]")


    # 1. Header (Sử dụng header cơ bản)
    format_basic_header(document, data, "GiayGioiThieu")


    # 2. Tên loại văn bản
    p_title = document.add_paragraph(title)
    set_paragraph_format(p_title, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=Pt(12), space_after=Pt(18))
    set_run_format(p_title.runs[0], size=FONT_SIZE_TITLE, bold=True)

    # 3. Kính gửi
    p_kg = document.add_paragraph(recipient_org)
    set_paragraph_format(p_kg, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(12))
    set_run_format(p_kg.runs[0], size=FONT_SIZE_DEFAULT, bold=True)

    # 4. Nội dung giới thiệu
    body_lines = body.split('\n')
    intro_org = data.get("issuing_org", "[Tên cơ quan giới thiệu]") # Lấy tên CQ từ data header
    p_intro = document.add_paragraph(f"{intro_org} trân trọng giới thiệu:")
    set_paragraph_format(p_intro, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=Pt(6), first_line_indent=FIRST_LINE_INDENT, line_spacing=1.5)
    set_run_format(p_intro.runs[0], size=FONT_SIZE_DEFAULT)

    for line in body_lines:
        stripped_line = line.strip()
        if not stripped_line or "trân trọng giới thiệu" in stripped_line.lower(): continue # Bỏ dòng đầu nếu có trong body

        p = document.add_paragraph()
        is_info_line = ":" in stripped_line and len(stripped_line.split(":")[0]) < 30

        align = WD_ALIGN_PARAGRAPH.LEFT
        first_indent = FIRST_LINE_INDENT
        if is_info_line:
            first_indent = Cm(0) # Không thụt lề dòng thông tin

        set_paragraph_format(p, alignment=align, space_after=Pt(6), first_line_indent=first_indent, left_indent=Cm(1.0), line_spacing=1.5) # Thụt lề nội dung chính

        if is_info_line:
            parts = stripped_line.split(":", 1)
            add_run_with_format(p, parts[0] + ":", size=FONT_SIZE_DEFAULT)
            add_run_with_format(p, parts[1], size=FONT_SIZE_DEFAULT)
        elif "đề nghị quý cơ quan" in stripped_line.lower():
             # Đoạn đề nghị không thụt lề trái
             set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=Pt(6), first_line_indent=FIRST_LINE_INDENT, left_indent=Cm(0), line_spacing=1.5)
             add_run_with_format(p, stripped_line, size=FONT_SIZE_DEFAULT)
        elif "có giá trị đến" in stripped_line.lower():
             # Đoạn giá trị hiệu lực không thụt lề, nghiêng
              set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=Pt(12), space_after=Pt(6), first_line_indent=Cm(0), left_indent=Cm(0), line_spacing=1.5)
              add_run_with_format(p, stripped_line, size=FONT_SIZE_SMALL, italic=True)
        else:
             add_run_with_format(p, stripped_line, size=FONT_SIZE_DEFAULT)

    # 5. Chữ ký (Sử dụng khối chữ ký cơ bản)
    if 'signer_title' not in data: data['signer_title'] = "THỦ TRƯỞNG CƠ QUAN" # Hoặc chức vụ ký GGT
    if 'signer_name' not in data: data['signer_name'] = "[Họ và tên]"
    document.add_paragraph()
    format_signature_block(document, data)

    # Giấy giới thiệu thường không có Nơi nhận ở cuối

    print("Định dạng Giấy giới thiệu hoàn tất.")