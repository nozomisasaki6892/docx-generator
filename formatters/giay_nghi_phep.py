# formatters/giay_nghi_phep.py
import time
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from utils import set_paragraph_format, set_run_format, add_run_with_format, add_centered_text
from config import FONT_SIZE_DEFAULT, FONT_SIZE_TITLE, FIRST_LINE_INDENT, FONT_SIZE_HEADER

def format(document, data):
    print("Bắt đầu định dạng Giấy/Đơn xin nghỉ phép...")
    title = data.get("title", "ĐƠN XIN NGHỈ PHÉP").upper()
    body = data.get("body", "Kính gửi:...\nTên tôi là:...\nChức vụ/Bộ phận:...\nNay tôi làm đơn này xin phép được nghỉ... ngày, từ ngày... đến ngày...\nLý do:...\nTrong thời gian nghỉ, công việc của tôi sẽ do Ông/Bà... phụ trách.\nKính mong Ban Lãnh đạo/Phòng ban xem xét, chấp thuận.\nTôi xin chân thành cảm ơn.")
    applicant_name = data.get("applicant_name", "[Họ và tên người làm đơn]")
    issuing_location = data.get("issuing_location", "Hà Nội") # Nơi viết đơn
    submission_date_str = data.get("submission_date", time.strftime(f"ngày %d tháng %m năm %Y")) # Ngày viết đơn


    # 1. Quốc hiệu, Tiêu ngữ (Căn giữa)
    add_centered_text(document, "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM", size=FONT_SIZE_HEADER, bold=True, space_after=0)
    add_centered_text(document, "Độc lập - Tự do - Hạnh phúc", size=Pt(13), bold=True, space_after=18)

    # 2. Tên Đơn
    add_centered_text(document, title, size=FONT_SIZE_TITLE, bold=True, space_before=12, space_after=12)

    # 3. Kính gửi
    recipient = data.get("recipient", "Ban Giám đốc Công ty ABC\nvà Trưởng phòng [Tên phòng ban]") # VD
    p_kg = document.add_paragraph(f"Kính gửi: {recipient}")
    set_paragraph_format(p_kg, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=Pt(12))
    set_run_format(p_kg.runs[0], size=FONT_SIZE_DEFAULT, bold=True)

    # 4. Nội dung Đơn
    body_lines = body.split('\n')
    for line in body_lines:
        stripped_line = line.strip()
        if not stripped_line: continue

        p = document.add_paragraph()
        # Nội dung đơn thường căn trái hoặc đều, thụt lề dòng đầu
        is_info_line = ":" in stripped_line and len(stripped_line.split(":")[0]) < 30 # Heuristic

        align = WD_ALIGN_PARAGRAPH.JUSTIFY
        first_indent = FIRST_LINE_INDENT
        is_bold = False
        if is_info_line:
            align = WD_ALIGN_PARAGRAPH.LEFT
            first_indent = Cm(0)

        set_paragraph_format(p, alignment=align, space_after=Pt(6), first_line_indent=first_indent, line_spacing=1.5)

        if is_info_line:
            parts = stripped_line.split(":", 1)
            add_run_with_format(p, parts[0] + ":", size=FONT_SIZE_DEFAULT)
            add_run_with_format(p, parts[1], size=FONT_SIZE_DEFAULT)
        elif "kính mong" in stripped_line.lower():
            add_run_with_format(p, stripped_line, size=FONT_SIZE_DEFAULT, bold=True) # Có thể đậm câu kính mong
        else:
             add_run_with_format(p, stripped_line, size=FONT_SIZE_DEFAULT)


    # 5. Lời cảm ơn
    p_thanks = document.add_paragraph("Tôi xin chân thành cảm ơn!")
    set_paragraph_format(p_thanks, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=Pt(12), space_after=Pt(12), first_line_indent=FIRST_LINE_INDENT, line_spacing=1.5)
    set_run_format(p_thanks.runs[0], size=FONT_SIZE_DEFAULT)

    # 6. Bảng chữ ký (Người làm đơn, Người duyệt - nếu cần)
    table = document.add_table(rows=1, cols=2) # Có thể 3 cột nếu có ý kiến bộ phận
    table.autofit = False
    table.allow_autofit = False
    table.columns[0].width = Cm(8.0) # Cột ý kiến duyệt (nếu có)
    table.columns[1].width = Cm(8.5) # Cột người làm đơn

    # --- Cột Phải: Người làm đơn ---
    cell_applicant = table.cell(0, 1)
    cell_applicant._element.clear_content()

    p_app_date = cell_applicant.add_paragraph(f"{issuing_location}, {submission_date_str}")
    set_paragraph_format(p_app_date, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(6))
    set_run_format(p_app_date.runs[0], size=FONT_SIZE_DEFAULT, italic=True)

    p_app_label = cell_applicant.add_paragraph("NGƯỜI LÀM ĐƠN")
    set_paragraph_format(p_app_label, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(0))
    set_run_format(p_app_label.runs[0], size=FONT_SIZE_DEFAULT, bold=True)

    p_app_note = cell_applicant.add_paragraph("(Ký và ghi rõ họ tên)")
    set_paragraph_format(p_app_note, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(12))
    set_run_format(p_app_note.runs[0], size=Pt(11), italic=True)

    cell_applicant.add_paragraph("\n\n\n") # Khoảng trống ký

    p_app_name = cell_applicant.add_paragraph(applicant_name)
    set_paragraph_format(p_app_name, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_run_format(p_app_name.runs[0], size=FONT_SIZE_DEFAULT, bold=True)

    # --- Cột Trái: Ý kiến phê duyệt (Tùy chọn) ---
    approver_title = data.get("approver_title", "Ý KIẾN CỦA TRƯỞNG BỘ PHẬN").upper()
    approver_name = data.get("approver_name", "")
    if data.get("needs_approval", False): # Thêm cờ để biết có cần cột duyệt không
        cell_approver = table.cell(0, 0)
        cell_approver._element.clear_content()

        p_appr_label = cell_approver.add_paragraph(approver_title)
        set_paragraph_format(p_appr_label, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(0))
        set_run_format(p_appr_label.runs[0], size=FONT_SIZE_DEFAULT, bold=True)

        p_appr_note = cell_approver.add_paragraph("(Duyệt và ký)")
        set_paragraph_format(p_appr_note, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(12))
        set_run_format(p_appr_note.runs[0], size=Pt(11), italic=True)

        cell_approver.add_paragraph("\n\n\n") # Khoảng trống ký

        if approver_name:
             p_appr_name = cell_approver.add_paragraph(approver_name)
             set_paragraph_format(p_appr_name, alignment=WD_ALIGN_PARAGRAPH.CENTER)
             set_run_format(p_appr_name.runs[0], size=FONT_SIZE_DEFAULT, bold=True)


    print("Định dạng Giấy/Đơn xin nghỉ phép hoàn tất.")