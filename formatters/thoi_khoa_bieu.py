# formatters/thoi_khoa_bieu.py
import re
import time
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_ALIGN_VERTICAL
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from utils import set_paragraph_format, set_run_format, add_run_with_format
from config import FONT_SIZE_DEFAULT, FONT_SIZE_TITLE, FONT_NAME, FONT_SIZE_HEADER

def set_cell_properties(cell, text, bold=False, italic=False, size=Pt(10), align=WD_ALIGN_PARAGRAPH.CENTER, valign=WD_ALIGN_VERTICAL.CENTER):
    """Helper to format cell content."""
    # Xóa paragraph cũ nếu có
    if len(cell.paragraphs) > 0:
        p = cell.paragraphs[0]
        p.clear()
    else:
         p = cell.add_paragraph()

    run = p.add_run(text)
    set_run_format(run, size=size, bold=bold, italic=italic)
    p.alignment = align
    cell.vertical_alignment = valign

def format(document, data):
    print("Bắt đầu định dạng Thời khóa biểu...")
    # Dữ liệu mẫu (cần cấu trúc chuẩn từ data)
    schedule_info = data.get("schedule_info", {
        "term": "Học kỳ I", "year": "2024 - 2025", "class": "Lớp XYZ", "faculty": "Khoa ABC",
        "issuer": "Phòng Đào tạo"
    })
    # Dữ liệu thời khóa biểu: list of dictionaries hoặc list of lists
    schedule_data = data.get("schedule_data", [
        # [Tiết, Thứ 2, Thứ 3, Thứ 4, Thứ 5, Thứ 6, Thứ 7]
        ["Sáng\n(7:00-11:30)", "Toán cao cấp A1\nPhòng: A101\nGV: Nguyễn Văn A", "Vật lý đại cương\nPhòng: B203\nGV: Trần Thị B", "", "Toán cao cấp A1\nPhòng: A101\nGV: Nguyễn Văn A", "Vật lý đại cương\nPhòng: B203\nGV: Trần Thị B", ""],
        ["Chiều\n(13:00-17:30)", "", "Tiếng Anh C1\nPhòng: C305\nGV: Lê Văn C", "Lập trình C++\nPhòng: D407\nGV: Phạm Thị D", "", "Lập trình C++\nPhòng: D407\nGV: Phạm Thị D", ""]
        # Thêm các buổi khác nếu cần
    ])
    issuing_org = data.get("issuing_org", "TÊN TRƯỜNG").upper()

    # 1. Header (Tên trường, đơn vị lập TKB)
    p_org = document.add_paragraph(issuing_org)
    set_paragraph_format(p_org, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(0))
    add_run_with_format(p_org, issuing_org, size=FONT_SIZE_HEADER, bold=True)
    p_dept = document.add_paragraph(schedule_info['issuer'].upper())
    set_paragraph_format(p_dept, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(12))
    add_run_with_format(p_dept, schedule_info['issuer'].upper(), size=Pt(11), bold=True)

    # 2. Tên Thời khóa biểu
    p_title = document.add_paragraph(f"THỜI KHÓA BIỂU {schedule_info['term'].upper()} NĂM HỌC {schedule_info['year']}")
    set_paragraph_format(p_title, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(6))
    add_run_with_format(p_title, p_title.text, size=FONT_SIZE_TITLE, bold=True, uppercase=True)
    p_class_info = document.add_paragraph(f"(Lớp: {schedule_info['class']} - Khoa: {schedule_info['faculty']})")
    set_paragraph_format(p_class_info, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(12))
    add_run_with_format(p_class_info, p_class_info.text, size=FONT_SIZE_DEFAULT)


    # 3. Tạo bảng thời khóa biểu
    # Xác định số hàng, số cột
    num_rows = len(schedule_data) + 1 # +1 cho header thứ
    num_cols = len(schedule_data[0]) if schedule_data else 7 # Mặc định 7 cột (Buổi/Tiết + 6 ngày)

    if num_cols < 2: # Cần ít nhất cột Buổi/Tiết và 1 ngày
         document.add_paragraph("Lỗi: Dữ liệu thời khóa biểu không hợp lệ.")
         return

    table = document.add_table(rows=num_rows, cols=num_cols)
    table.style = 'Table Grid' # Kiểu bảng có đường kẻ
    table.autofit = False # Tắt tự động điều chỉnh

    # Thiết lập độ rộng cột (ước lượng)
    # Cột đầu (Buổi/Tiết) nhỏ hơn các cột ngày
    table.columns[0].width = Inches(0.8)
    day_col_width = Inches(5.5 / (num_cols - 1)) # Chia đều phần còn lại
    for i in range(1, num_cols):
        table.columns[i].width = day_col_width

    # Header bảng (Thứ 2 -> Thứ 7)
    header_cells = table.rows[0].cells
    days = ["Buổi/Tiết", "Thứ 2", "Thứ 3", "Thứ 4", "Thứ 5", "Thứ 6", "Thứ 7"] # Có thể tùy chỉnh
    for i in range(num_cols):
        col_header = days[i] if i < len(days) else f"Cột {i+1}"
        set_cell_properties(header_cells[i], col_header, bold=True, size=Pt(11))

    # Đổ dữ liệu vào bảng
    for r, row_data in enumerate(schedule_data):
        row_cells = table.rows[r + 1].cells
        for c, cell_data in enumerate(row_data):
             if c < num_cols: # Đảm bảo không ghi ra ngoài số cột
                 # Cột đầu tiên (Buổi/Tiết) có thể in đậm
                 is_bold = (c == 0)
                 align = WD_ALIGN_PARAGRAPH.LEFT if c > 0 else WD_ALIGN_PARAGRAPH.CENTER # Căn trái nội dung môn học
                 set_cell_properties(row_cells[c], str(cell_data), bold=is_bold, size=Pt(10), align=align, valign=WD_ALIGN_VERTICAL.TOP)


    # 4. Ghi chú (nếu có)
    notes = data.get("notes", None)
    if notes:
        p_notes_label = document.add_paragraph("Ghi chú:", space_before=Pt(12))
        add_run_with_format(p_notes_label, "Ghi chú:", bold=True)
        if isinstance(notes, list):
            for note in notes:
                 p_note = document.add_paragraph(f"- {note}")
                 set_paragraph_format(p_note, left_indent=Cm(0.5), space_before=Pt(0), space_after=Pt(0), line_spacing=1.0)
        else:
             p_note = document.add_paragraph(f"- {notes}")
             set_paragraph_format(p_note, left_indent=Cm(0.5), space_before=Pt(0), space_after=Pt(0), line_spacing=1.0)

    # 5. Ngày tháng lập biểu và Chữ ký người lập
    p_date_create = document.add_paragraph(f"{data.get('issuing_location', '........')}, ngày {time.strftime('%d')} tháng {time.strftime('%m')} năm {time.strftime('%Y')}")
    set_paragraph_format(p_date_create, alignment=WD_ALIGN_PARAGRAPH.RIGHT, space_before=Pt(12), space_after=Pt(0))
    add_run_with_format(p_date_create, p_date_create.text, size=FONT_SIZE_DEFAULT, italic=True)

    p_signer_title = document.add_paragraph()
    set_paragraph_format(p_signer_title, alignment=WD_ALIGN_PARAGRAPH.RIGHT, space_before=Pt(0), space_after=Pt(60)) # Chừa chỗ ký
    add_run_with_format(p_signer_title, data.get("signer_title", "NGƯỜI LẬP BIỂU").upper(), bold=True)

    p_signer_name = document.add_paragraph()
    set_paragraph_format(p_signer_name, alignment=WD_ALIGN_PARAGRAPH.RIGHT, space_before=Pt(0))
    add_run_with_format(p_signer_name, data.get("signer_name", "[Ký, ghi rõ họ tên]"), bold=True)


    print("Định dạng Thời khóa biểu hoàn tất.")