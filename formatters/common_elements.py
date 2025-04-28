# common_elements.py
import time
import traceback
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
# Import utils mới và config mới
import utils
from config import (
    FONT_NAME,
    FONT_SIZE_HEADER, FONT_SIZE_TIEUNGU_DIADANH, FONT_SIZE_SOKYHIEU,
    FONT_SIZE_SIGNATURE_AUTH, FONT_SIZE_SIGNATURE_NAME,
    FONT_SIZE_RECIPIENT_LABEL, FONT_SIZE_RECIPIENT_LIST,
    FONT_SIZE_SMALL # FONT_SIZE_11
)

def add_header_elements(document, data):
    """Thêm các thành phần Header chuẩn NĐ30 vào đầu tài liệu dùng table."""
    print("  COMMON: Adding Header Elements (New Arch)...", flush=True)
    try:
        header_table = document.add_table(rows=1, cols=2)
        header_table.autofit = False
        header_table.allow_autofit = False
        header_table.columns[0].width = Inches(2.9)
        header_table.columns[1].width = Inches(3.3)
        cell_left = header_table.cell(0, 0)
        cell_right = header_table.cell(0, 1)
        cell_left._element.clear_content()
        cell_right._element.clear_content()

        # --- Ô Trái: Tên CQ, Số/KH (Ô 2, 3) ---
        print("    COMMON: Adding Ten Co Quan...", flush=True)
        ten_cqcq = data.get("issuing_org_parent") # Lấy từ data do app.py chuẩn bị
        ten_cqbh = data.get("issuing_org", "[TÊN CƠ QUAN]").upper()
        if ten_cqcq:
            p_cqcq = cell_left.add_paragraph()
            utils.set_paragraph_format(p_cqcq, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(0))
            utils.add_run_with_format(p_cqcq, ten_cqcq.upper(), size=FONT_SIZE_HEADER, bold=False) # NĐ30: CQ chủ quản ko đậm
        p_cqbh = cell_left.add_paragraph()
        utils.set_paragraph_format(p_cqbh, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(0))
        utils.add_run_with_format(p_cqbh, ten_cqbh, size=FONT_SIZE_HEADER, bold=True) # NĐ30: CQ ban hành đậm
        p_line_cq = cell_left.add_paragraph()
        utils.set_paragraph_format(p_line_cq, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(0))
        utils.add_run_with_format(p_line_cq, "________", size=FONT_SIZE_HEADER, bold=True) # Dòng kẻ đậm

        print("    COMMON: Adding So Ky Hieu...", flush=True)
        so_vb = data.get("doc_number_only", "...")
        ky_hieu = data.get("doc_symbol", "...")
        p_skh = cell_left.add_paragraph()
        # Căn giữa trong ô trái, cách dòng kẻ trên 6pt
        utils.set_paragraph_format(p_skh, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=Pt(6), space_after=Pt(0))
        utils.add_run_with_format(p_skh, f"Số: {so_vb}/", size=FONT_SIZE_SOKYHIEU, bold=False)
        utils.add_run_with_format(p_skh, ky_hieu.upper(), size=FONT_SIZE_SOKYHIEU, bold=False)

        # --- Ô Phải: QH/TN, Địa danh/TG (Ô 1, 4) ---
        print("    COMMON: Adding Quoc Hieu Tieu Ngu...", flush=True)
        p_qh = cell_right.add_paragraph()
        utils.set_paragraph_format(p_qh, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(0))
        utils.add_run_with_format(p_qh, "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM", size=FONT_SIZE_HEADER, bold=True) # NĐ30: QH đậm
        p_tn = cell_right.add_paragraph()
        utils.set_paragraph_format(p_tn, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(0))
        utils.add_run_with_format(p_tn, "Độc lập - Tự do - Hạnh phúc", size=FONT_SIZE_TIEUNGU_DIADANH, bold=True) # NĐ30: TN đậm
        p_line_tn = cell_right.add_paragraph()
        utils.set_paragraph_format(p_line_tn, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(0))
        utils.add_run_with_format(p_line_tn, "_______________", size=FONT_SIZE_TIEUNGU_DIADANH, bold=True) # Dòng kẻ đậm

        print("    COMMON: Adding Dia Danh Thoi Gian...", flush=True)
        dia_danh = data.get("issuing_location", "Hà Nội")
        try:
            ngay = int(data.get("issuing_day", time.strftime("%d")))
            thang = int(data.get("issuing_month", time.strftime("%m")))
            nam = int(data.get("issuing_year", time.strftime("%Y")))
        except ValueError:
             ngay, thang, nam = int(time.strftime("%d")), int(time.strftime("%m")), int(time.strftime("%Y"))
             print("    WARNING: Invalid date data in common_elements, using current date.", flush=True)
        thoi_gian_str = f"ngày {ngay:02d} tháng {thang:02d} năm {nam}"
        p_ddtg = cell_right.add_paragraph()
        # Căn giữa trong ô phải, cách dòng kẻ trên 6pt
        utils.set_paragraph_format(p_ddtg, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=Pt(6), space_after=Pt(0))
        utils.add_run_with_format(p_ddtg, f"{dia_danh}, {thoi_gian_str}", size=FONT_SIZE_TIEUNGU_DIADANH, italic=True) # NĐ30: Nghiêng

        print("  COMMON: Header Elements Added OK.", flush=True)
        # Thêm khoảng trắng sau header để cách với tên loại VB
        document.add_paragraph()

    except Exception as e:
        print(f"!!!!!!!! ERROR Adding Header Elements !!!!!!!!!!", flush=True)
        print(traceback.format_exc(), flush=True)
        document.add_paragraph(f"[Lỗi tạo Header: {e}]")


def add_signature_block(document, data):
    """Thêm khối chữ ký chuẩn NĐ30."""
    print("  COMMON: Adding Signature Block (New Arch)...", flush=True)
    try:
        authority_signer = data.get("authority_signer") # VD: "TM. CHÍNH PHỦ"
        signer_title = data.get("signer_title", "[CHỨC VỤ]").upper() # VD: "THỦ TƯỚNG"
        signer_name = data.get("signer_name", "[Họ tên]")
        signer_note = data.get("signer_note") # VD: "(Đã ký)"

        sig_table = document.add_table(rows=1, cols=2)
        sig_table.autofit = False
        sig_table.allow_autofit = False
        sig_table.columns[0].width = Inches(3.0) # Cột trái trống
        sig_table.columns[1].width = Inches(3.3) # Cột phải chứa chữ ký
        cell_sig = sig_table.cell(0, 1)
        cell_sig._element.clear_content()
        sig_align = WD_ALIGN_PARAGRAPH.CENTER # Căn giữa nội dung trong ô phải

        # Ô 7a: Quyền hạn (nếu có)
        if authority_signer:
            p_auth = cell_sig.add_paragraph()
            utils.set_paragraph_format(p_auth, alignment=sig_align, space_after=Pt(0))
            utils.add_run_with_format(p_auth, authority_signer.upper(), size=FONT_SIZE_SIGNATURE_AUTH, bold=True)

        # Ô 7a: Chức vụ
        p_title = cell_sig.add_paragraph()
        utils.set_paragraph_format(p_title, alignment=sig_align, space_after=Pt(0))
        utils.add_run_with_format(p_title, signer_title, size=FONT_SIZE_SIGNATURE_AUTH, bold=True)

        # Ghi chú (nếu có)
        if signer_note:
            p_note = cell_sig.add_paragraph()
            utils.set_paragraph_format(p_note, alignment=sig_align, space_after=Pt(0))
            utils.add_run_with_format(p_note, signer_note, size=FONT_SIZE_SMALL, italic=True)

        # Ô 7c: Khoảng trống ký
        p_space = cell_sig.add_paragraph("\n\n\n\n")
        utils.set_paragraph_format(p_space, alignment=sig_align, space_after=Pt(0))

        # Ô 7b: Tên người ký
        p_name = cell_sig.add_paragraph()
        utils.set_paragraph_format(p_name, alignment=sig_align, space_after=Pt(0))
        utils.add_run_with_format(p_name, signer_name, size=FONT_SIZE_SIGNATURE_NAME, bold=True)

        print("  COMMON: Signature Block Added OK.", flush=True)

    except Exception as e:
        print(f"!!!!!!!! ERROR Adding Signature Block !!!!!!!!!!", flush=True)
        print(traceback.format_exc(), flush=True)
        try: document.add_paragraph(f"[Lỗi tạo Chữ ký: {e}]")
        except: pass


def add_recipient_list(document, data):
    """Thêm khối Nơi nhận chuẩn NĐ30."""
    print("  COMMON: Adding Recipient List (New Arch)...", flush=True)
    try:
        recipients = data.get("recipients") # Lấy list từ data do app.py chuẩn bị
        if not recipients:
             print("  WARNING: No recipients list found in data. Adding default.", flush=True)
             recipients = ["- Như Điều ...;", "- Lưu: VT, ...;"] # Mặc định nếu thiếu

        # Ô 9b: Chữ "Nơi nhận:"
        p_label = document.add_paragraph()
        # Vị trí: Ngang hàng quyền hạn ký, sát lề trái
        # (Do khối ký dùng table nên ta chỉ cần căn trái paragraph này)
        utils.set_paragraph_format(p_label, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=Pt(0), space_after=Pt(0))
        utils.add_run_with_format(p_label, "Nơi nhận:", size=FONT_SIZE_RECIPIENT_LABEL, bold=True, italic=True)

        # Danh sách nơi nhận
        for recipient in recipients:
            p_rec = document.add_paragraph()
            utils.set_paragraph_format(
                p_rec, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                left_indent=Cm(0), first_line_indent=Cm(0), # Nơi nhận sát lề trái
                space_before=Pt(0), space_after=Pt(0), line_spacing=1.0
            )
            # Thêm gạch đầu dòng thủ công nếu cần (hoặc dùng style list của Word)
            text = recipient if recipient.startswith("-") else f"- {recipient}"
            utils.add_run_with_format(p_rec, text, size=FONT_SIZE_RECIPIENT_LIST, bold=False)

        print("  COMMON: Recipient List Added OK.", flush=True)
    except Exception as e:
        print(f"!!!!!!!! ERROR Adding Recipient List !!!!!!!!!!", flush=True)
        print(traceback.format_exc(), flush=True)
        try: document.add_paragraph(f"[Lỗi tạo Nơi nhận: {e}]")
        except: pass