# formatters/giay_bao_trung_tuyen.py
import time
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from utils import set_paragraph_format, set_run_format, add_run_with_format
try:
    # Giấy báo thường có header, signature riêng
    from .common_elements import format_recipient_list # Có thể dùng nơi nhận nếu cần gửi nhiều
except ImportError:
    def format_recipient_list(document, data): pass

from config import FONT_SIZE_DEFAULT, FONT_SIZE_TITLE, FIRST_LINE_INDENT, FONT_SIZE_SIGNATURE, FONT_SIZE_SIGNER_NAME, FONT_SIZE_HEADER

def format_gbtt_header(document, data):
    # Header Giấy báo tương tự header cơ bản NĐ30
    issuing_org_parent = data.get("issuing_org_parent", "BỘ GIÁO DỤC VÀ ĐÀO TẠO").upper() # Ví dụ
    issuing_org = data.get("issuing_org", "TRƯỜNG ĐẠI HỌC XYZ").upper()
    doc_number = data.get("doc_number", "Số:       /GBTT-...")
    issuing_location = data.get("issuing_location", "Hà Nội")
    current_date_str = time.strftime(f"ngày %d tháng %m năm %Y")

    header_table = document.add_table(rows=1, cols=2)
    header_table.autofit = False
    header_table.allow_autofit = False
    header_table.columns[0].width = Cm(8.0)
    header_table.columns[1].width = Cm(8.5)

    # Cột trái: CQ chủ quản, CQ ban hành, Số hiệu
    cell_left = header_table.cell(0, 0)
    cell_left._element.clear_content()
    p_parent = cell_left.add_paragraph(issuing_org_parent)
    set_paragraph_format(p_parent, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(0))
    set_run_format(p_parent.runs[0], size=FONT_SIZE_HEADER, bold=False)
    p_org = cell_left.add_paragraph(issuing_org)
    set_paragraph_format(p_org, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(0))
    set_run_format(p_org.runs[0], size=FONT_SIZE_HEADER, bold=True)
    p_line_org = cell_left.add_paragraph("_______")
    set_paragraph_format(p_line_org, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(6))
    set_run_format(p_line_org.runs[0], size=FONT_SIZE_HEADER, bold=True)
    p_num = cell_left.add_paragraph(doc_number)
    set_paragraph_format(p_num, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=Pt(6))
    set_run_format(p_num.runs[0], size=Pt(13))

    # Cột phải: QH, TN, Ngày tháng
    cell_right = header_table.cell(0, 1)
    cell_right._element.clear_content()
    p_qh = cell_right.add_paragraph("CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM")
    set_paragraph_format(p_qh, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(0))
    set_run_format(p_qh.runs[0], size=FONT_SIZE_HEADER, bold=True)
    p_tn = cell_right.add_paragraph("Độc lập - Tự do - Hạnh phúc")
    set_paragraph_format(p_tn, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(0))
    set_run_format(p_tn.runs[0], size=Pt(13), bold=True)
    p_line_tn = cell_right.add_paragraph("-" * 20)
    set_paragraph_format(p_line_tn, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(6))
    set_run_format(p_line_tn.runs[0], size=Pt(13), bold=True)
    p_date = cell_right.add_paragraph(f"{issuing_location}, {current_date_str}")
    set_paragraph_format(p_date, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=Pt(6))
    set_run_format(p_date.runs[0], size=Pt(13), italic=True)

    document.add_paragraph()


def format(document, data):
    print("Bắt đầu định dạng Giấy báo trúng tuyển...")
    title = data.get("title", "GIẤY BÁO TRÚNG TUYỂN VÀ NHẬP HỌC").upper()
    body = data.get("body", "Hội đồng tuyển sinh... trân trọng thông báo...\nThí sinh:...\nĐã trúng tuyển vào ngành:...\nThời gian nhập học:...\nĐịa điểm:...\nHồ sơ cần chuẩn bị:...")
    student_name = data.get("student_name", "[Họ và tên thí sinh]")

    # 1. Header
    format_gbtt_header(document, data)

    # 2. Tên Giấy báo
    p_title = document.add_paragraph(title)
    set_paragraph_format(p_title, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=Pt(12), space_after=Pt(18))
    set_run_format(p_title.runs[0], size=FONT_SIZE_TITLE, bold=True)


    # 3. Nội dung thông báo
    body_lines = body.split('\n')
    for line in body_lines:
        stripped_line = line.strip()
        if not stripped_line: continue

        p = document.add_paragraph()
        # Nội dung giấy báo thường căn trái hoặc đều
        is_info_line = ":" in stripped_line and len(stripped_line.split(":")[0]) < 40 # Heuristic
        is_list_item = stripped_line.startswith("-") or stripped_line.startswith("+") or stripped_line.startswith("*")

        align = WD_ALIGN_PARAGRAPH.JUSTIFY
        first_indent = FIRST_LINE_INDENT
        left_indent = Cm(0)
        is_bold = False

        if is_info_line:
            align = WD_ALIGN_PARAGRAPH.LEFT
            first_indent = Cm(0)
            parts = stripped_line.split(":", 1)
            # In đậm phần thông tin thí sinh trúng tuyển
            if "thí sinh" in parts[0].lower() or "trúng tuyển" in parts[0].lower() or "ngành" in parts[0].lower():
                 add_run_with_format(p, parts[0] + ":", size=FONT_SIZE_DEFAULT, bold=True)
                 add_run_with_format(p, parts[1], size=FONT_SIZE_DEFAULT, bold=True)
                 is_bold = True # Flag để không add run lần nữa
            else:
                 add_run_with_format(p, parts[0] + ":", size=FONT_SIZE_DEFAULT)
                 add_run_with_format(p, parts[1], size=FONT_SIZE_DEFAULT)

        elif is_list_item:
            align = WD_ALIGN_PARAGRAPH.LEFT
            left_indent = Cm(1.0)
            first_indent = Cm(-0.5) # Hanging indent
            add_run_with_format(p, stripped_line, size=FONT_SIZE_DEFAULT)
        else:
             add_run_with_format(p, stripped_line, size=FONT_SIZE_DEFAULT)


        set_paragraph_format(p, alignment=align, space_after=Pt(6), first_line_indent=first_indent, left_indent=left_indent, line_spacing=1.5)
        # Chỉ add_run nếu chưa add ở is_info_line
        # if not is_info_line and not is_list_item:
        #      run = p.runs[0] if p.runs else p.add_run()
        #      run.text = stripped_line
        #      set_run_format(run, size=FONT_SIZE_DEFAULT)


    # 4. Chữ ký (Thường là Chủ tịch Hội đồng tuyển sinh / Hiệu trưởng)
    signer_title = data.get("signer_title", "CHỦ TỊCH HỘI ĐỒNG TUYỂN SINH").upper()
    signer_name = data.get("signer_name", "[Họ và tên]")

    sig_paragraph = document.add_paragraph()
    set_paragraph_format(sig_paragraph, alignment=WD_ALIGN_PARAGRAPH.RIGHT, space_before=Pt(18), space_after=Pt(0), line_spacing=1.0)
    add_run_with_format(sig_paragraph, signer_title + "\n\n\n\n\n", size=FONT_SIZE_SIGNATURE, bold=True)
    add_run_with_format(sig_paragraph, signer_name, size=FONT_SIZE_SIGNER_NAME, bold=True)


    print("Định dạng Giấy báo trúng tuyển hoàn tất.")