# formatters/thong_tu.py
import re
import time
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
# Header, signature giống văn bản hành chính thông thường
try:
    from .common_elements import format_basic_header, format_signature_block, format_recipient_list
except ImportError:
    def format_basic_header(document, data, doc_type): pass
    def format_signature_block(document, data): pass
    def format_recipient_list(document, data): pass

from config import FONT_SIZE_DEFAULT, FONT_SIZE_TITLE, FIRST_LINE_INDENT

def format(document, data):
    print("Bắt đầu định dạng Thông tư...")
    title = data.get("title", "Thông tư Quy định chi tiết/Hướng dẫn thực hiện...")
    body = data.get("body", "Căn cứ Nghị định...\nTheo đề nghị của Vụ trưởng Vụ...\nBộ trưởng Bộ... ban hành Thông tư:\nChương I...\nĐiều 1...")
    doc_type_label = "THÔNG TƯ"
    minister_name = data.get("minister_name", "BỘ TRƯỞNG BỘ [TÊN BỘ]").upper() # Thẩm quyền ban hành

    # 1. Header (Bộ ban hành)
    # Cần đảm bảo data['issuing_org'] là tên Bộ
    if 'issuing_org' not in data: data['issuing_org'] = f"BỘ {data.get('ministry_name', '[TÊN BỘ]')}".upper()
    format_basic_header(document, data, "ThongTu")


    # 2. Tên loại và Trích yếu
    p_tenloai = document.add_paragraph(doc_type_label)
    set_paragraph_format(p_tenloai, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=Pt(12), space_after=Pt(6))
    add_run_with_format(p_tenloai, doc_type_label, size=FONT_SIZE_TITLE, bold=True, uppercase=True)

    subject = title.replace("Thông tư", "").strip()
    if subject.lower().startswith("quy định"):
         subject = subject.split(" ", 1)[-1]
    elif subject.lower().startswith("hướng dẫn"):
         subject = subject.split(" ", 1)[-1]

    p_subject = document.add_paragraph(subject)
    set_paragraph_format(p_subject, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(12))
    add_run_with_format(p_subject.runs[0], size=Pt(14), bold=True) # Trích yếu đậm


    # 3. Thẩm quyền ban hành (BỘ TRƯỞNG)
    p_authority = document.add_paragraph(minister_name)
    set_paragraph_format(p_authority, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(6))
    add_run_with_format(p_authority, minister_name, size=FONT_SIZE_DEFAULT, bold=True, uppercase=True)


    # 4. Căn cứ ban hành
    body_lines = body.split('\n')
    processed_indices = set()

    for i, line in enumerate(body_lines):
        stripped_line = line.strip()
        if not stripped_line: continue
        if stripped_line.lower().startswith("căn cứ") or "theo đề nghị của" in stripped_line.lower() or "ban hành thông tư" in stripped_line.lower():
            p = document.add_paragraph(stripped_line)
            set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line_indent=FIRST_LINE_INDENT, line_spacing=1.5, space_after=Pt(0))
            set_run_format(p.runs[0], size=FONT_SIZE_DEFAULT, italic=True)
            processed_indices.add(i)
            if "ban hành thông tư" in stripped_line.lower():
                 break # Dừng sau câu này
        # Dừng khi hết căn cứ hoặc gặp Chương/Điều
        elif any(l.strip().lower().startswith("căn cứ") or "theo đề nghị của" in l.strip().lower() for l in body_lines[:i]):
             if stripped_line.upper().startswith("CHƯƠNG") or stripped_line.upper().startswith("ĐIỀU"):
                 break

    if processed_indices: document.add_paragraph()


    # 5. Nội dung Thông tư (Chương, Mục, Điều, Khoản, Điểm)
    for i, line in enumerate(body_lines):
        if i in processed_indices: continue
        stripped_line = line.strip()
        if not stripped_line: continue
        # Bỏ qua câu ban hành nếu đã xử lý ở căn cứ
        if "ban hành thông tư" in stripped_line.lower(): continue

        p = document.add_paragraph()

        is_chuong = stripped_line.upper().startswith("CHƯƠNG")
        is_muc = re.match(r'^(MỤC\s+\d+)\.?\s+', stripped_line.upper())
        is_dieu = stripped_line.upper().startswith("ĐIỀU")
        is_khoan = re.match(r'^\d+\.\s+', stripped_line)
        is_diem = re.match(r'^[a-z]\)\s+', stripped_line)

        align = WD_ALIGN_PARAGRAPH.JUSTIFY
        left_indent = Cm(0)
        first_indent = FIRST_LINE_INDENT if not (is_chuong or is_muc or is_dieu or is_khoan or is_diem) else Cm(0)
        is_bold = False
        size = FONT_SIZE_DEFAULT
        space_before = Pt(0)
        space_after = Pt(6)
        line_spacing = 1.5

        if is_chuong:
            align = WD_ALIGN_PARAGRAPH.CENTER
            is_bold = True
            space_before = Pt(12)
            size = Pt(13)
        elif is_muc:
            align = WD_ALIGN_PARAGRAPH.CENTER
            is_bold = True
            space_before = Pt(6)
            size = Pt(13)
        elif is_dieu:
            align = WD_ALIGN_PARAGRAPH.LEFT
            is_bold = True
            space_before = Pt(6)
            size = Pt(13)
        elif is_khoan:
            left_indent = Cm(0.5)
        elif is_diem:
            left_indent = Cm(1.0)

        set_paragraph_format(p, alignment=align, left_indent=left_indent, first_line_indent=first_indent, line_spacing=line_spacing, space_before=space_before, space_after=space_after)
        add_run_with_format(p, stripped_line, size=size, bold=is_bold)


    # 6. Chữ ký (Bộ trưởng/Thứ trưởng)
    if 'signer_title' not in data: data['signer_title'] = "BỘ TRƯỞNG" # Hoặc KT. BỘ TRƯỞNG \n THỨ TRƯỞNG
    if 'signer_name' not in data: data['signer_name'] = "[Họ và tên Bộ trưởng]"
    format_signature_block(document, data) # Dùng signature block chuẩn

    # 7. Nơi nhận
    format_recipient_list(document, data)

    print("Định dạng Thông tư hoàn tất.")