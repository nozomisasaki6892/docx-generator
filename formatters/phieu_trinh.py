# formatters/phieu_trinh.py
import time
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from utils import set_paragraph_format, set_run_format, add_run_with_format
try:
    # Phiếu trình thường có cấu trúc chữ ký riêng (Người trình, Người duyệt)
    from .common_elements import format_basic_header # Có thể dùng header đơn vị
except ImportError:
    def format_basic_header(document, data, doc_type): pass

from config import FONT_SIZE_DEFAULT, FONT_SIZE_TITLE, FIRST_LINE_INDENT, FONT_SIZE_HEADER

def format_presentation_signatures(document, presenter_data, approver_data):
    table = document.add_table(rows=1, cols=2)
    table.autofit = False
    table.allow_autofit = False
    table.columns[0].width = Cm(8.0)
    table.columns[1].width = Cm(8.5)

    # Cột Trái: Người trình
    cell_presenter = table.cell(0, 0)
    cell_presenter._element.clear_content()
    p_pres_label = cell_presenter.add_paragraph("NGƯỜI TRÌNH")
    set_paragraph_format(p_pres_label, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(6))
    set_run_format(p_pres_label.runs[0], size=FONT_SIZE_DEFAULT, bold=True)
    cell_presenter.add_paragraph("\n\n\n")
    p_pres_name = cell_presenter.add_paragraph(presenter_data.get("name", "[Tên người trình]"))
    set_paragraph_format(p_pres_name, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_run_format(p_pres_name.runs[0], size=FONT_SIZE_DEFAULT, bold=True)

    # Cột Phải: Người duyệt
    cell_approver = table.cell(0, 1)
    cell_approver._element.clear_content()
    p_appr_label = cell_approver.add_paragraph(approver_data.get("title", "NGƯỜI DUYỆT").upper())
    set_paragraph_format(p_appr_label, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(6))
    set_run_format(p_appr_label.runs[0], size=FONT_SIZE_DEFAULT, bold=True)
    cell_approver.add_paragraph("\n\n\n")
    p_appr_name = cell_approver.add_paragraph(approver_data.get("name", "[Tên người duyệt]"))
    set_paragraph_format(p_appr_name, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_run_format(p_appr_name.runs[0], size=FONT_SIZE_DEFAULT, bold=True)


def format(document, data):
    print("Bắt đầu định dạng Phiếu trình...")
    title = data.get("title", "PHIẾU TRÌNH").upper()
    subject = data.get("subject", "V/v: Xin ý kiến chỉ đạo về việc ABC") # Nội dung trình
    body = data.get("body", "Kính gửi: [Lãnh đạo duyệt]\n1. Nội dung sự việc:\n...\n2. Đề xuất/Kiến nghị:\n...\nKính trình Lãnh đạo xem xét, cho ý kiến chỉ đạo.")
    presenter_data = data.get("presenter", {}) # {'name': '...', 'title': '...'}
    approver_data = data.get("approver", {}) # {'name': '...', 'title': '...'}
    issuing_date_str = data.get("issuing_date", time.strftime(f"ngày %d tháng %m năm %Y"))


    # 1. Header (Tên đơn vị trình - nếu cần)
    if data.get("issuing_org"):
        format_basic_header(document, data, "PhieuTrinh")

    # 2. Tên Phiếu trình
    p_title = document.add_paragraph(title)
    set_paragraph_format(p_title, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=Pt(12), space_after=Pt(6))
    set_run_format(p_title.runs[0], size=FONT_SIZE_TITLE, bold=True)

    # Nội dung trình (Trích yếu)
    p_subject = document.add_paragraph(subject)
    set_paragraph_format(p_subject, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(12))
    set_run_format(p_subject.runs[0], size=Pt(14), bold=True)


    # 3. Kính gửi
    recipient = data.get("recipient", "Kính gửi: [Tên Lãnh đạo duyệt]")
    p_kg = document.add_paragraph(recipient)
    set_paragraph_format(p_kg, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=Pt(12))
    set_run_format(p_kg.runs[0], size=FONT_SIZE_DEFAULT, bold=True)


    # 4. Nội dung trình bày
    body_lines = body.split('\n')
    for line in body_lines:
        stripped_line = line.strip()
        if not stripped_line: continue

        p = document.add_paragraph()
        # Nội dung phiếu trình thường căn trái hoặc đều
        is_section_digit = re.match(r'^(\d+\.)\s+', stripped_line) # Mục 1, 2
        is_bullet = stripped_line.startswith("-") or stripped_line.startswith("+") or stripped_line.startswith("*")

        align = WD_ALIGN_PARAGRAPH.JUSTIFY
        left_indent = Cm(0)
        first_indent = FIRST_LINE_INDENT if not (is_section_digit or is_bullet) else Cm(0)
        is_bold = bool(is_section_digit) # Mục lớn có thể đậm
        is_italic = False
        size = FONT_SIZE_DEFAULT
        line_spacing = 1.5

        if is_section_digit:
            align = WD_ALIGN_PARAGRAPH.LEFT
            space_before = Pt(6)
        elif is_bullet:
            align = WD_ALIGN_PARAGRAPH.LEFT
            left_indent = Cm(0.5)
            first_indent = Cm(-0.5) # Hanging indent

        set_paragraph_format(p, alignment=align, space_after=Pt(6), first_line_indent=first_indent, left_indent=left_indent, line_spacing=line_spacing)

        # Xử lý câu kết "Kính trình..."
        if "kính trình" in stripped_line.lower() and "xem xét" in stripped_line.lower():
             add_run_with_format(p, stripped_line, size=FONT_SIZE_DEFAULT, bold=True) # Đậm câu kết
        else:
             add_run_with_format(p, stripped_line, size=size, bold=is_bold, italic=is_italic)


    # 5. Chữ ký (Người trình, Người duyệt)
    document.add_paragraph() # Khoảng cách
    format_presentation_signatures(document, presenter_data, approver_data)

    # 6. Ý kiến chỉ đạo (Phần để trống cho lãnh đạo ghi) - Có thể thêm nếu cần
    document.add_paragraph("\n")
    p_ykien_label = document.add_paragraph("Ý KIẾN CHỈ ĐẠO:")
    set_paragraph_format(p_ykien_label, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=Pt(12))
    set_run_format(p_ykien_label.runs[0], size=FONT_SIZE_DEFAULT, bold=True)
    document.add_paragraph("\n\n\n\n") # Để trống nhiều dòng

    print("Định dạng Phiếu trình hoàn tất.")