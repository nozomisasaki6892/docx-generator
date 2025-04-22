# formatters/de_cuong_mh.py
import re
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from utils import set_paragraph_format, set_run_format, add_run_with_format
from config import FONT_SIZE_DEFAULT, FONT_SIZE_TITLE, FIRST_LINE_INDENT

def format(document, data):
    print("Bắt đầu định dạng Đề cương môn học...")
    title = data.get("title", "Đề cương chi tiết môn học ABC")
    body = data.get("body", "1. Thông tin chung về môn học\n2. Mục tiêu môn học\n3. Nội dung chi tiết...\n4. Tài liệu tham khảo...")
    university_name = data.get("university_name", "TRƯỜNG ĐẠI HỌC XYZ").upper()
    faculty_name = data.get("faculty_name", "KHOA CÔNG NGHỆ THÔNG TIN").upper()

    # 1. Thông tin Trường, Khoa (Căn giữa)
    p_uni = document.add_paragraph(university_name)
    set_paragraph_format(p_uni, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(0))
    set_run_format(p_uni.runs[0], size=Pt(13), bold=True)

    p_fac = document.add_paragraph(faculty_name)
    set_paragraph_format(p_fac, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(18))
    set_run_format(p_fac.runs[0], size=Pt(13), bold=True)

    # 2. Tên Đề cương
    p_title = document.add_paragraph("ĐỀ CƯƠNG CHI TIẾT MÔN HỌC")
    set_paragraph_format(p_title, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=Pt(12), space_after=Pt(6))
    set_run_format(p_title.runs[0], size=FONT_SIZE_TITLE, bold=True)

    # Tên môn học cụ thể
    course_name = title.replace("Đề cương chi tiết môn học", "").replace("Đề cương môn học", "").strip()
    p_name = document.add_paragraph(course_name.upper()) # Tên môn học IN HOA
    set_paragraph_format(p_name, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(18))
    set_run_format(p_name.runs[0], size=Pt(14), bold=True)


    # 3. Nội dung Đề cương
    body_lines = body.split('\n')
    for line in body_lines:
        stripped_line = line.strip()
        if not stripped_line: continue
        p = document.add_paragraph()

        # Logic định dạng cơ bản cho các đề mục đề cương
        is_section_digit = re.match(r'^(\d+)\.\s+', stripped_line) # Mục 1, 2, 3
        is_subsection_digit = re.match(r'^(\d+\.\d+\.?)\s+', stripped_line) # Mục 1.1, 1.2
        is_subsubsection_alpha = re.match(r'^[a-z]\)\s+', stripped_line) # Mục a, b, c
        is_bullet = stripped_line.startswith("-") or stripped_line.startswith("+") or stripped_line.startswith("*") or stripped_line.startswith("•")

        align = WD_ALIGN_PARAGRAPH.JUSTIFY
        left_indent = Cm(0)
        first_indent = FIRST_LINE_INDENT if not (is_section_digit or is_subsection_digit or is_subsubsection_alpha or is_bullet) else Cm(0)
        is_bold = bool(is_section_digit or is_subsection_digit) # Mục lớn đậm
        is_italic = False
        size = FONT_SIZE_DEFAULT
        space_before = Pt(0)
        space_after = Pt(6)
        line_spacing = 1.5

        if is_section_digit:
            align = WD_ALIGN_PARAGRAPH.LEFT
            space_before = Pt(12)
            size = Pt(13)
        elif is_subsection_digit:
            align = WD_ALIGN_PARAGRAPH.LEFT
            left_indent = Cm(0.5) # Thụt lề mục con
            space_before = Pt(6)
        elif is_subsubsection_alpha:
             align = WD_ALIGN_PARAGRAPH.LEFT
             left_indent = Cm(1.0)
        elif is_bullet:
             align = WD_ALIGN_PARAGRAPH.LEFT
             left_indent = Cm(1.5)
             first_indent = Cm(-0.5) # Hanging indent

        set_paragraph_format(p, alignment=align, left_indent=left_indent, first_line_indent=first_indent, line_spacing=line_spacing, space_before=space_before, space_after=space_after)
        add_run_with_format(p, stripped_line, size=size, bold=is_bold, italic=is_italic)


    # 4. Thông tin giảng viên, ngày phê duyệt (Nếu có, thường cuối file)
    # Cần lấy thông tin này từ data

    # Ví dụ phần cuối
    document.add_paragraph() # Thêm khoảng trống
    p_approval_loc_date = document.add_paragraph(f"Hà Nội, ngày ...... tháng ...... năm ......")
    set_paragraph_format(p_approval_loc_date, alignment=WD_ALIGN_PARAGRAPH.RIGHT, space_after=Pt(6))
    set_run_format(p_approval_loc_date.runs[0], size=FONT_SIZE_DEFAULT, italic=True)

    # Bảng chữ ký nếu cần (Trưởng Khoa, Giảng viên biên soạn) - Tương tự bảng chữ ký Biên bản
    # signature_data = data.get('signatures', [{'title':'TRƯỞNG KHOA', 'name':'...'}, {'title':'GIẢNG VIÊN BIÊN SOẠN', 'name':'...'}])
    # format_signatures_in_table(document, signature_data) # Cần định nghĩa hàm này hoặc import từ biên bản nếu dùng chung


    print("Định dạng Đề cương môn học hoàn tất.")