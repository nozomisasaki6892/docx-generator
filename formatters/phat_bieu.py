# formatters/phat_bieu.py
import time
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from utils import set_paragraph_format, set_run_format, add_run_with_format, add_centered_text
from config import FONT_SIZE_DEFAULT, FONT_SIZE_TITLE, FIRST_LINE_INDENT

# Lưu ý: Bài phát biểu không có cấu trúc chuẩn như văn bản hành chính.
# Định dạng chủ yếu là tiêu đề và nội dung diễn văn.

def format(document, data):
    print("Bắt đầu định dạng Bài phát biểu...")
    title = data.get("title", "BÀI PHÁT BIỂU KHAI MẠC/CHÀO MỪNG/...") # Tiêu đề bài phát biểu
    subtitle = data.get("subtitle", None) # Phụ đề (VD: Tại Hội nghị ABC)
    body = data.get("body", "Kính thưa quý vị đại biểu...\nThưa toàn thể hội nghị...\n...")
    speaker_info = data.get("speaker_info", "Người phát biểu: [Tên], Chức vụ: [Chức vụ]")
    location_date = data.get("location_date", f"Hà Nội, ngày {time.strftime('%d tháng %m năm %Y')}")


    # 1. Tiêu đề Bài phát biểu (Căn giữa, đậm, IN HOA)
    add_centered_text(document, title.upper(), size=FONT_SIZE_TITLE, bold=True, space_before=Pt(18), space_after=Pt(6))
    if subtitle:
        add_centered_text(document, f"({subtitle})", size=Pt(14), bold=True, italic=True, space_after=Pt(18))


    # 2. Nội dung Bài phát biểu
    body_lines = body.split('\n')
    for line in body_lines:
        stripped_line = line.strip()
        if not stripped_line:
             document.add_paragraph() # Giữ lại dòng trống để tạo khoảng cách đoạn
             continue

        p = document.add_paragraph()
        # Lời chào, kính thưa thường căn trái, đậm hoặc nghiêng tùy ngữ cảnh
        is_greeting = stripped_line.lower().startswith("kính thưa") or stripped_line.lower().startswith("thưa")

        align = WD_ALIGN_PARAGRAPH.JUSTIFY
        first_indent = FIRST_LINE_INDENT
        is_bold = False
        is_italic = False
        size = Pt(14) # Cỡ chữ bài phát biểu thường lớn hơn (14pt)
        line_spacing = 1.5

        if is_greeting:
            align = WD_ALIGN_PARAGRAPH.LEFT
            first_indent = Cm(0)
            is_bold = True # Lời chào đậm
            is_italic = False
            space_after = Pt(12) # Giãn cách sau lời chào
        else:
             space_after = Pt(6) # Giãn cách thông thường


        set_paragraph_format(p, alignment=align, space_after=space_after, first_line_indent=first_indent, line_spacing=line_spacing)
        add_run_with_format(p, stripped_line, size=size, bold=is_bold, italic=is_italic)


    # 3. Lời kết và Cảm ơn
    p_thanks = document.add_paragraph("Xin trân trọng cảm ơn!")
    set_paragraph_format(p_thanks, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=Pt(12), space_after=Pt(12))
    set_run_format(p_thanks.runs[0], size=Pt(14), bold=True, italic=True)


    # 4. Thông tin người phát biểu và ngày tháng (Căn phải)
    p_loc_date = document.add_paragraph(location_date)
    set_paragraph_format(p_loc_date, alignment=WD_ALIGN_PARAGRAPH.RIGHT, space_after=Pt(6))
    set_run_format(p_loc_date.runs[0], size=Pt(13), italic=True)

    p_speaker = document.add_paragraph(speaker_info)
    set_paragraph_format(p_speaker, alignment=WD_ALIGN_PARAGRAPH.RIGHT, space_after=Pt(0))
    set_run_format(p_speaker.runs[0], size=Pt(13), bold=True)


    print("Định dạng Bài phát biểu hoàn tất.")