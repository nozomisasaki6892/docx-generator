# formatters/giay_xac_nhan_sv.py
import time
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from utils import set_paragraph_format, set_run_format, add_run_with_format, add_centered_text
try:
    # Dùng header, signature của văn bản hành chính thông thường
    from .common_elements import format_basic_header, format_signature_block
except ImportError:
    def format_basic_header(document, data, doc_type): pass
    def format_signature_block(document, data): pass

from config import FONT_SIZE_DEFAULT, FONT_SIZE_TITLE, FIRST_LINE_INDENT, FONT_SIZE_HEADER

def format(document, data):
    print("Bắt đầu định dạng Giấy xác nhận sinh viên...")
    title = data.get("title", "GIẤY XÁC NHẬN").upper()
    confirmation_subject = data.get("confirmation_subject", "(V/v: Xác nhận sinh viên)") # Nội dung xác nhận
    body = data.get("body", "Trường [Tên trường] xác nhận:\nAnh/Chị:...\nNgày sinh:...\nLà sinh viên năm thứ... Lớp... Khoa... Hệ đào tạo...\nKhóa học:...\nMã số sinh viên:...\nHiện đang học tập tại trường.\nLý do xin xác nhận:...\nGiấy xác nhận này có giá trị trong vòng ... tháng kể từ ngày ký.")
    student_name = data.get("student_name", "[Họ và tên sinh viên]")


    # 1. Header (Trường cấp xác nhận)
    # Đảm bảo data có 'issuing_org' là tên trường
    if 'issuing_org' not in data: data['issuing_org'] = "TRƯỜNG ĐẠI HỌC XYZ"
    format_basic_header(document, data, "GiayXacNhanSV")


    # 2. Tên Giấy xác nhận
    add_centered_text(document, title, size=FONT_SIZE_TITLE, bold=True, space_before=12, space_after=6)
    # Chủ đề xác nhận (nếu có)
    add_centered_text(document, confirmation_subject, size=Pt(14), bold=True, space_after=18)


    # 3. Nội dung xác nhận
    body_lines = body.split('\n')
    issuing_org_name = data.get("issuing_org", "[Tên trường]")
    p_intro = document.add_paragraph(f"{issuing_org_name} xác nhận:")
    set_paragraph_format(p_intro, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=Pt(6), first_line_indent=FIRST_LINE_INDENT, line_spacing=1.5)
    set_run_format(p_intro.runs[0], size=FONT_SIZE_DEFAULT)

    for line in body_lines:
        stripped_line = line.strip()
        if not stripped_line or "xác nhận:" in stripped_line.lower(): continue

        p = document.add_paragraph()
        is_info_line = ":" in stripped_line and len(stripped_line.split(":")[0]) < 30

        align = WD_ALIGN_PARAGRAPH.LEFT
        first_indent = Cm(0) # Thông tin xác nhận căn trái, không thụt lề

        set_paragraph_format(p, alignment=align, space_after=Pt(6), first_line_indent=first_indent, left_indent=Cm(1.0), line_spacing=1.5) # Thụt lề thông tin

        if is_info_line:
            parts = stripped_line.split(":", 1)
            # In đậm thông tin sinh viên
            if "anh/chị" in parts[0].lower() or "ngày sinh" in parts[0].lower() or "sinh viên năm" in parts[0].lower() or "lớp" in parts[0].lower() or "khoa" in parts[0].lower() or "mã số" in parts[0].lower():
                add_run_with_format(p, parts[0] + ":", size=FONT_SIZE_DEFAULT, bold=True)
                add_run_with_format(p, parts[1], size=FONT_SIZE_DEFAULT, bold=True)
            else:
                add_run_with_format(p, parts[0] + ":", size=FONT_SIZE_DEFAULT)
                add_run_with_format(p, parts[1], size=FONT_SIZE_DEFAULT)
        elif "có giá trị" in stripped_line.lower():
             # Giá trị hiệu lực, nghiêng
             set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=Pt(12), space_after=Pt(6), first_line_indent=Cm(0), left_indent=Cm(0), line_spacing=1.5)
             add_run_with_format(p, stripped_line, size=FONT_SIZE_DEFAULT, italic=True)
        else:
             add_run_with_format(p, stripped_line, size=FONT_SIZE_DEFAULT)


    # 4. Chữ ký (Thường là Thủ trưởng đơn vị cấp xác nhận: Hiệu trưởng, Trưởng phòng CTSV...)
    if 'signer_title' not in data: data['signer_title'] = "KT. HIỆU TRƯỞNG\nTRƯỞNG PHÒNG CÔNG TÁC SINH VIÊN" # Ví dụ
    if 'signer_name' not in data: data['signer_name'] = "[Họ và tên]"
    document.add_paragraph()
    format_signature_block(document, data) # Dùng signature block chuẩn


    print("Định dạng Giấy xác nhận sinh viên hoàn tất.")