# formatters/nghi_dinh.py
import re
import time
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from utils import (
    set_paragraph_format,
    set_run_format,
    add_run_with_format,
    add_paragraph_with_text
)
# Import các hàm đã xây dựng lại từ common_elements
from common_elements import (
    add_quoc_hieu_tieu_ngu,
    add_ten_co_quan_ban_hanh,
    add_so_ky_hieu,
    add_dia_danh_thoi_gian,
    add_signature_block,
    add_recipient_list
)

# Giả định các hằng số cỡ chữ đã được định nghĩa chuẩn trong config.py hoặc utils.py
FONT_SIZE_13 = Pt(13)
FONT_SIZE_14 = Pt(14)
FIRST_LINE_INDENT = Cm(1.0) # Hoặc Cm(1.27)

def format(document, data):
    """Định dạng văn bản theo thể thức Nghị định của Chính phủ (QPPL)."""
    print("Bắt đầu định dạng Nghị định Chính phủ (QPPL)...")

    # 1. Thiết lập Header theo chuẩn Nghị định 30
    # Tạo bảng 1 hàng 2 cột để bố trí header
    header_table = document.add_table(rows=1, cols=2)
    header_table.autofit = False
    header_table.allow_autofit = False
    # Điều chỉnh độ rộng cột theo Phụ lục I, Mục IV
    # Cột 1 (Tên CQ, Số/KH): Khoảng 2.9 Inches
    # Cột 2 (QH/TN, Ngày tháng): Khoảng 3.3 Inches
    header_table.columns[0].width = Inches(2.9)
    header_table.columns[1].width = Inches(3.3)

    # Lấy ô cho từng cột
    cell_left = header_table.cell(0, 0)
    cell_right = header_table.cell(0, 1)
    cell_left._element.clear_content()
    cell_right._element.clear_content()

    # Điền nội dung vào các ô header
    # Ô Trái (Ô 2, Ô 3)
    add_ten_co_quan_ban_hanh(cell_left, None, "CHÍNH PHỦ") # Không có CQ chủ quản
    so_van_ban = data.get("decree_number_only", "...") # Chỉ lấy phần số
    ky_hieu_van_ban = data.get("decree_symbol", "NĐ-CP") # Chỉ lấy phần ký hiệu
    add_so_ky_hieu(cell_left, so_van_ban, ky_hieu_van_ban)

    # Ô Phải (Ô 1, Ô 4)
    add_quoc_hieu_tieu_ngu(cell_right)
    dia_danh = data.get("issuing_location", "Hà Nội")
    ngay = int(data.get("issuing_day", time.strftime("%d")))
    thang = int(data.get("issuing_month", time.strftime("%m")))
    nam = int(data.get("issuing_year", time.strftime("%Y")))
    add_dia_danh_thoi_gian(cell_right, dia_danh, ngay, thang, nam)

    document.add_paragraph() # Khoảng cách sau header

    # 2. Tên loại văn bản "NGHỊ ĐỊNH" (Ô 5a)
    add_paragraph_with_text(
        document,
        "NGHỊ ĐỊNH",
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_before=Pt(6), # Điều chỉnh space_before/after nếu cần
        space_after=Pt(6),
        size=FONT_SIZE_14, # Cỡ 14
        bold=True,
        uppercase=True
    )

    # 3. Trích yếu nội dung Nghị định (Ô 5a)
    trich_yeu = data.get("title", "Quy định chi tiết thi hành...")
    paragraph_trichyeu = add_paragraph_with_text(
        document,
        trich_yeu,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=Pt(6),
        size=FONT_SIZE_14, # Cỡ 14
        bold=True # Đậm
    )
    # Thêm dòng kẻ dưới trích yếu
    paragraph_line = document.add_paragraph()
    set_paragraph_format(paragraph_line, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(12))
    add_run_with_format(paragraph_line, "________", size=FONT_SIZE_14, bold=True) # Điều chỉnh độ dài

    # 4. Cơ quan ban hành (lặp lại, căn giữa, đậm) - Theo thể thức NĐ QPPL
    add_paragraph_with_text(
        document,
        "CHÍNH PHỦ",
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=Pt(12),
        size=FONT_SIZE_13, # Cỡ 13-14
        bold=True,
        uppercase=True
    )

    # 5. Phần căn cứ ban hành (Ô 6)
    can_cu_list = data.get("can_cu", [])
    if can_cu_list:
        for item in can_cu_list:
            stripped_item = item.strip()
            if stripped_item:
                add_paragraph_with_text(
                    document,
                    stripped_item,
                    alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                    first_line_indent=FIRST_LINE_INDENT,
                    line_spacing=1.5,
                    space_after=Pt(0), # Không cách đoạn giữa các căn cứ
                    size=FONT_SIZE_13, # Cỡ 13-14
                    italic=True # Căn cứ in nghiêng
                )
        # Thêm dòng trống sau khối căn cứ
        document.add_paragraph()

    # 6. Phần nội dung (Ô 6 - Chương, Điều, Khoản, Điểm)
    body_content = data.get("body", "")
    body_lines = body_content.strip().split('\n')
    nghi_dinh_title_upper = "NGHỊ ĐỊNH" # Tiêu đề cần tránh lặp

    for line in body_lines:
        stripped_line = line.strip()
        if not stripped_line:
            continue

        # Bỏ qua nếu dòng là tiêu đề "NGHỊ ĐỊNH" bị lặp lại
        if stripped_line.upper() == nghi_dinh_title_upper:
            print(f"Bỏ qua dòng tiêu đề bị lặp: '{stripped_line}'")
            continue

        paragraph = document.add_paragraph()
        # Mặc định cho nội dung thường
        align = WD_ALIGN_PARAGRAPH.JUSTIFY
        left_indent = Cm(0)
        first_indent = FIRST_LINE_INDENT
        is_bold = False
        size = FONT_SIZE_13 # Cỡ 13-14
        space_before = Pt(0)
        space_after = Pt(6) # Mặc định cách 6pt sau mỗi đoạn
        line_spacing = 1.5 # Mặc định 1.5 lines
        text_to_add = stripped_line

        # --- Logic nhận diện và định dạng Chương, Điều, Khoản, Điểm ---
        # Chương: Căn giữa, đậm, cỡ 13-14
        match_chuong = re.match(r'^(Chương\s+[IVXLCDM]+)\s*(.*?)$', stripped_line, re.IGNORECASE)
        if match_chuong:
            chuong_num = match_chuong.group(1)
            chuong_title = match_chuong.group(2).strip().upper()
            align = WD_ALIGN_PARAGRAPH.CENTER
            first_indent = Cm(0)
            is_bold = True
            size = FONT_SIZE_13
            space_before = Pt(12)
            space_after = Pt(6)
            text_to_add = f"{chuong_num}\n{chuong_title}" # Xuống dòng giữa số và tên chương
            add_run_with_format(paragraph, text_to_add, size=size, bold=is_bold, uppercase=False) # Đã uppercase tên chương

        # Điều: Căn trái, đậm (số và tiêu đề), cỡ 13-14
        match_dieu = re.match(r'^(Điều\s+\d+)\.\s*(.*)', stripped_line, re.IGNORECASE)
        if match_dieu and not match_chuong: # Đảm bảo không phải là dòng Chương
            dieu_num_title = match_dieu.group(1) # Giữ Điều
            dieu_content = match_dieu.group(2).strip()
            align = WD_ALIGN_PARAGRAPH.LEFT
            first_indent = Cm(0)
            is_bold = True
            size = FONT_SIZE_13
            space_before = Pt(6)
            space_after = Pt(3) # Giảm cách sau Điều
            # Chỉ format lại text nếu chưa phải là Chương
            paragraph.clear()
            add_run_with_format(paragraph, f"{dieu_num_title}. ", size=size, bold=is_bold)
            add_run_with_format(paragraph, dieu_content, size=size, bold=is_bold) # Tiêu đề Điều cũng đậm
            text_to_add = None # Đã add run

        # Khoản: Thụt lề 1cm/1.27cm, cỡ 13-14
        match_khoan = re.match(r'^(\d+)\.\s+(.*)', stripped_line)
        if match_khoan and not match_chuong and not match_dieu:
            khoan_num = match_khoan.group(1)
            khoan_content = match_khoan.group(2).strip()
            align = WD_ALIGN_PARAGRAPH.JUSTIFY
            left_indent = FIRST_LINE_INDENT # Thụt lề đúng bằng first indent của đoạn thường
            first_indent = Cm(0)
            is_bold = False
            size = FONT_SIZE_13
            space_before = Pt(3)
            space_after = Pt(3)
            # Chỉ format lại text nếu chưa phải là Chương/Điều
            paragraph.clear()
            add_run_with_format(paragraph, f"{khoan_num}. {khoan_content}", size=size, bold=is_bold)
            text_to_add = None # Đã add run

        # Điểm: Thụt lề thêm 1cm/1.27cm, cỡ 13-14
        match_diem = re.match(r'^([a-z])\)\s+(.*)', stripped_line)
        if match_diem and not match_chuong and not match_dieu and not match_khoan:
            diem_marker = match_diem.group(1)
            diem_content = match_diem.group(2).strip()
            align = WD_ALIGN_PARAGRAPH.JUSTIFY
            left_indent = FIRST_LINE_INDENT + Cm(0.5) # Thụt thêm so với Khoản
            first_indent = Cm(0)
            is_bold = False
            size = FONT_SIZE_13
            space_before = Pt(3)
            space_after = Pt(3)
            # Chỉ format lại text nếu chưa phải là Chương/Điều/Khoản
            paragraph.clear()
            add_run_with_format(paragraph, f"{diem_marker}) {diem_content}", size=size, bold=is_bold)
            text_to_add = None # Đã add run

        # Áp dụng định dạng đoạn
        set_paragraph_format(paragraph, alignment=align, left_indent=left_indent, first_line_indent=first_indent, line_spacing=line_spacing, space_before=space_before, space_after=space_after)
        # Thêm text nếu chưa được add run cụ thể ở trên
        if text_to_add:
            add_run_with_format(paragraph, text_to_add, size=size, bold=is_bold)

    # 7. Khối chữ ký (Ô 7)
    signer_authority = data.get("authority_signer", "TM. CHÍNH PHỦ") # Thẩm quyền ký NĐ CP
    signer_title = data.get("signer_title", "THỦ TƯỚNG") # Chức vụ ký
    signer_name = data.get("signer_name", "[Tên Thủ tướng]") # Tên người ký
    add_signature_block(
        document,
        authority_signer=signer_authority,
        signer_title=signer_title,
        signer_name=signer_name,
        alignment=WD_ALIGN_PARAGRAPH.RIGHT, # Chữ ký NĐ CP căn phải
        use_table=True # Dùng table để căn phải dễ hơn
    )

    # 8. Nơi nhận (Ô 9) - Lấy danh sách chuẩn cho NĐ CP
    default_recipients = [
        "- Ban Bí thư Trung ương Đảng;",
        "- Thủ tướng, các Phó Thủ tướng Chính phủ;",
        "- Các bộ, cơ quan ngang bộ, cơ quan thuộc Chính phủ;",
        "- HĐND, UBND các tỉnh, thành phố trực thuộc trung ương;",
        "- Văn phòng Trung ương và các Ban của Đảng;",
        "- Văn phòng Tổng Bí thư;",
        "- Văn phòng Chủ tịch nước;",
        "- Hội đồng Dân tộc và các Ủy ban của Quốc hội;",
        "- Văn phòng Quốc hội;",
        "- Tòa án nhân dân tối cao;",
        "- Viện kiểm sát nhân dân tối cao;",
        "- Kiểm toán Nhà nước;",
        # Thêm các cơ quan khác nếu cần...
        "- VPCP: BTCN, các PCN, Trợ lý TTg, TGĐ Cổng TTĐT, các Vụ, Cục, đơn vị trực thuộc, Công báo;",
        "- Lưu: VT, [Ký hiệu đơn vị soạn thảo]." # Thay ký hiệu đúng
    ]
    recipients = data.get("recipients", default_recipients)
    add_recipient_list(document, recipients)

    print("Định dạng Nghị định Chính phủ (QPPL) hoàn tất.")