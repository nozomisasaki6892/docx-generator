# formatters/du_an.py
import re
import time
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from utils import set_paragraph_format, set_run_format, add_run_with_format
try:
    # Dự án có thể dùng header, signature, recipient như văn bản hành chính thông thường
    from .common_elements import format_basic_header, format_signature_block, format_recipient_list
except ImportError:
    # Fallback nếu chạy độc lập
    def format_basic_header(document, data, doc_type): pass
    def format_signature_block(document, data): pass
    def format_recipient_list(document, data): pass

from config import FONT_SIZE_DEFAULT, FONT_SIZE_TITLE, FIRST_LINE_INDENT

def format(document, data):
    print("Bắt đầu định dạng Dự án...")
    title = data.get("title", "Dự án Đầu tư Xây dựng Công trình XYZ")
    body = data.get("body", "Phần I: THÔNG TIN CHUNG VỀ DỰ ÁN\nPhần II: MỤC TIÊU VÀ QUY MÔ ĐẦU TƯ\nPhần III: GIẢI PHÁP THỰC HIỆN\nPhần IV: TỔNG MỨC ĐẦU TƯ VÀ NGUỒN VỐN...")
    doc_type_label = "DỰ ÁN"
    investor = data.get("investor", "CHỦ ĐẦU TƯ: [Tên chủ đầu tư]") # Thông tin chủ đầu tư

    # 1. Header (Tùy chọn, có thể không cần nếu trình bày dạng báo cáo)
    # format_basic_header(document, data, "DuAn")

    # 2. Tên Dự án
    p_title = document.add_paragraph(doc_type_label)
    set_paragraph_format(p_title, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=Pt(18), space_after=Pt(6))
    set_run_format(p_title.runs[0], size=FONT_SIZE_TITLE, bold=True)

    # Tên dự án cụ thể
    project_name = title.replace("Dự án", "").strip()
    p_name = document.add_paragraph(project_name.upper()) # Tên dự án IN HOA
    set_paragraph_format(p_name, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(12))
    set_run_format(p_name.runs[0], size=Pt(14), bold=True)

    # Thông tin chủ đầu tư (nếu có)
    if investor:
        p_investor = document.add_paragraph(investor)
        set_paragraph_format(p_investor, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(18))
        set_run_format(p_investor.runs[0], size=FONT_SIZE_DEFAULT, bold=True)


    # 3. Nội dung Dự án
    body_lines = body.split('\n')
    for line in body_lines:
        stripped_line = line.strip()
        if not stripped_line: continue
        p = document.add_paragraph()

        # Logic định dạng cơ bản cho các đề mục dự án (tương tự Đề án)
        is_part_roman = re.match(r'^(PHẦN\s+[IVXLCDM]+)\.?\s+', stripped_line.upper())
        is_section_digit = re.match(r'^([IVXLCDM]+)\.\s+', stripped_line.upper()) # Mục La Mã
        is_subsection_digit = re.match(r'^(\d+\.)\s+', stripped_line) # Mục 1, 2, 3
        is_subsubsection_alpha = re.match(r'^[a-z]\)\s+', stripped_line) # Mục a, b, c
        is_bullet = stripped_line.startswith("-") or stripped_line.startswith("+") or stripped_line.startswith("*")

        align = WD_ALIGN_PARAGRAPH.JUSTIFY
        left_indent = Cm(0)
        first_indent = FIRST_LINE_INDENT if not (is_part_roman or is_section_digit or is_subsection_digit or is_subsubsection_alpha or is_bullet) else Cm(0)
        is_bold = bool(is_part_roman or is_section_digit or is_subsection_digit) # Mục lớn đậm
        is_italic = False
        size = FONT_SIZE_DEFAULT
        space_before = Pt(0)
        space_after = Pt(6)
        line_spacing = 1.5

        if is_part_roman:
            align = WD_ALIGN_PARAGRAPH.CENTER
            space_before = Pt(18)
            size = Pt(14)
        elif is_section_digit:
            align = WD_ALIGN_PARAGRAPH.LEFT
            space_before = Pt(12)
            size = Pt(13)
        elif is_subsection_digit:
            align = WD_ALIGN_PARAGRAPH.LEFT
            left_indent = Cm(0.5)
            space_before = Pt(6)
        elif is_subsubsection_alpha:
             align = WD_ALIGN_PARAGRAPH.LEFT
             left_indent = Cm(1.0)
        elif is_bullet:
             align = WD_ALIGN_PARAGRAPH.LEFT
             left_indent = Cm(1.5)
             first_indent = Cm(-0.5) # Hanging indent

        set_paragraph_format(p, alignment=align, left_indent=left_indent, first_line_indent=first_indent, line_spacing=line_spacing, space_before=space_before, space_after=space_after)
        add_run_with_format(p, stripped_line, size=size, bold=is_bold, italic=is_italic)


    # 4. Chữ ký (Đại diện Chủ đầu tư / Cơ quan phê duyệt)
    if 'signer_title' not in data: data['signer_title'] = "ĐẠI DIỆN CHỦ ĐẦU TƯ" # Hoặc chức vụ phê duyệt
    if 'signer_name' not in data: data['signer_name'] = "[Họ và tên]"
    document.add_paragraph()
    format_signature_block(document, data)

    # 5. Nơi nhận (Nếu dự án được ban hành như văn bản)
    if data.get('recipients'):
         format_recipient_list(document, data)

    print("Định dạng Dự án hoàn tất.")