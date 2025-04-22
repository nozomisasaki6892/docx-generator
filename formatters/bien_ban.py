# formatters/bien_ban.py
import time
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from utils import set_paragraph_format, set_run_format, add_run_with_format
try:
    # Biên bản thường không có chữ ký theo kiểu NĐ30, chữ ký nằm cuối nội dung
    from .common_elements import format_basic_header # Chỉ dùng header
except ImportError:
    def format_basic_header(document, data, doc_type): pass

from config import FONT_SIZE_DEFAULT, FONT_SIZE_TITLE, FIRST_LINE_INDENT

def format_signatures_in_table(document, participants_signatures):
    if not participants_signatures: return

    num_participants = len(participants_signatures)
    # Tạo bảng chữ ký, có thể 2 hoặc 3 cột tùy số lượng
    num_cols = 2 if num_participants <= 4 else 3
    num_rows = (num_participants + num_cols - 1) // num_cols

    table = document.add_table(rows=num_rows, cols=num_cols)
    table.style = 'Table Grid' # Có thể bỏ viền nếu muốn
    table.autofit = False
    table.allow_autofit = False
    col_width = Cm(16.5 / num_cols) # Chia đều chiều rộng
    for col in table.columns:
        col.width = col_width

    idx = 0
    for r in range(num_rows):
        for c in range(num_cols):
            if idx < num_participants:
                participant = participants_signatures[idx]
                cell = table.cell(r, c)
                cell._element.clear_content()
                # Chức vụ/Vai trò
                p_title = cell.add_paragraph(participant.get('title', 'Thành phần tham dự').upper())
                set_paragraph_format(p_title, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(6))
                set_run_format(p_title.runs[0], size=Pt(11), bold=True)
                # Khoảng trống
                cell.add_paragraph("\n\n\n")
                # Tên
                p_name = cell.add_paragraph(participant.get('name', '[Họ và tên]'))
                set_paragraph_format(p_name, alignment=WD_ALIGN_PARAGRAPH.CENTER)
                set_run_format(p_name.runs[0], size=Pt(11), bold=True)
                idx += 1
            else:
                 # Xóa nội dung ô thừa nếu số người không chia hết
                 table.cell(r, c)._element.clear_content()


def format(document, data):
    print("Bắt đầu định dạng Biên bản...")
    title = data.get("title", "Biên bản họp/làm việc/nghiệm thu...")
    body = data.get("body", "I. Thời gian, địa điểm...\nII. Thành phần tham dự...\nIII. Nội dung...\nIV. Kết luận...")
    doc_type_label = "BIÊN BẢN"

    # 1. Header (Sử dụng header cơ bản)
    # Biên bản có thể có hoặc không có header CQBH tùy ngữ cảnh
    format_basic_header(document, data, "BienBan")

    # 2. Tiêu đề
    p_title = document.add_paragraph(doc_type_label)
    set_paragraph_format(p_title, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=Pt(12), space_after=Pt(6))
    set_run_format(p_title.runs[0], size=FONT_SIZE_TITLE, bold=True)

    # Trích yếu nội dung biên bản
    subject = title.replace("Biên bản", "").strip()
    p_subject = document.add_paragraph(f"V/v: {subject}")
    set_paragraph_format(p_subject, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(12))
    set_run_format(p_subject.runs[0], size=Pt(14), bold=True)


    # 3. Nội dung biên bản
    body_lines = body.split('\n')
    participants_data = [] # Tách thông tin người ký nếu có cấu trúc

    for line in body_lines:
        stripped_line = line.strip()
        if not stripped_line: continue
        p = document.add_paragraph()

        # Logic định dạng cơ bản cho các đề mục biên bản
        is_section_roman = re.match(r'^([IVXLCDM]+)\.\s+', stripped_line.upper()) # I, II, III
        is_subsection_digit = re.match(r'^(\d+\.)\s+', stripped_line) # 1, 2, 3
        is_subsubsection_alpha = re.match(r'^[a-z]\)\s+', stripped_line) # a, b, c
        is_bullet = stripped_line.startswith("-") or stripped_line.startswith("+") or stripped_line.startswith("*")
        is_participant_line = "chủ trì" in stripped_line.lower() or "thư ký" in stripped_line.lower() or "thành phần" in stripped_line.lower()

        align = WD_ALIGN_PARAGRAPH.JUSTIFY
        left_indent = Cm(0)
        first_indent = FIRST_LINE_INDENT if not (is_section_roman or is_subsection_digit or is_subsubsection_alpha or is_bullet) else Cm(0)
        is_bold = bool(is_section_roman or is_subsection_digit or is_participant_line)
        is_italic = False
        size = FONT_SIZE_DEFAULT
        space_before = Pt(0)
        space_after = Pt(6)
        line_spacing = 1.5

        if is_section_roman:
            align = WD_ALIGN_PARAGRAPH.LEFT
            space_before = Pt(12)
            size = Pt(13)
        elif is_subsection_digit:
            align = WD_ALIGN_PARAGRAPH.LEFT
            left_indent = Cm(0.5) # Thụt lề mục con
            space_before = Pt(6)
        elif is_subsubsection_alpha:
            align = WD_ALIGN_PARAGRAPH.LEFT
            left_indent = Cm(1.0)
        elif is_bullet:
             align = WD_ALIGN_PARAGRAPH.LEFT
             left_indent = Cm(1.5)
             first_indent = Cm(-0.5) # Hanging indent

        set_paragraph_format(p, alignment=align, left_indent=left_indent, first_line_indent=first_indent, line_spacing=line_spacing, space_before=space_before, space_after=space_after)
        add_run_with_format(p, stripped_line, size=size, bold=is_bold, italic=is_italic)

        # Trích xuất thông tin người ký (cần logic phức tạp hơn)
        # if is_participant_line:
        #    match = re.search(r'(Ông|Bà)\s+(.*?)\s+Chức vụ:\s+(.*)', stripped_line)
        #    if match:
        #        participants_data.append({'name': match.group(2), 'title': match.group(3)})


    # 4. Kết thúc biên bản
    p_end = document.add_paragraph(f"Biên bản này được lập thành ... bản, có giá trị pháp lý như nhau. Các bên đã đọc lại, thống nhất nội dung và cùng ký tên dưới đây.") # Hoặc câu kết thúc phù hợp
    set_paragraph_format(p_end, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_before=Pt(12), space_after=Pt(12), first_line_indent=FIRST_LINE_INDENT, line_spacing=1.5)
    set_run_format(p_end.runs[0], size=FONT_SIZE_DEFAULT, italic=True)


    # 5. Chữ ký các bên tham gia (dạng bảng)
    # Dữ liệu chữ ký cần được gửi trong data['participants_signatures'] dưới dạng list of dicts [{'title': '...', 'name': '...'}, ...]
    participants_signatures = data.get('participants_signatures', [])
    if not participants_signatures:
         # Thêm mẫu nếu không có dữ liệu
         participants_signatures = [{'title': 'Chủ trì', 'name': '[Họ tên]'}, {'title': 'Thư ký', 'name': '[Họ tên]'}]
    format_signatures_in_table(document, participants_signatures)

    # Biên bản thường không có Nơi nhận riêng

    print("Định dạng Biên bản hoàn tất.")