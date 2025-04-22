# formatters/thong_bao.py
import re
import time
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from utils import set_paragraph_format, set_run_format, add_run_with_format
try:
    from .common_elements import format_basic_header, format_signature_block, format_recipient_list
except ImportError:
    # Fallback nếu chạy độc lập
    def format_basic_header(document, data, doc_type): pass
    def format_signature_block(document, data): pass
    def format_recipient_list(document, data): pass

from config import FONT_SIZE_DEFAULT, FONT_SIZE_TITLE, FIRST_LINE_INDENT

def format(document, data):
    print("Bắt đầu định dạng Thông báo...")
    title = data.get("title", "Thông báo về việc Nghỉ lễ/Thay đổi lịch...")
    body = data.get("body", "Căn cứ...\n[Cơ quan] thông báo đến... nội dung như sau:\n1...\n2...\nTrân trọng thông báo!")
    doc_type_label = "THÔNG BÁO"

    # 1. Header (Sử dụng header cơ bản)
    format_basic_header(document, data, "ThongBao")

    # 2. Tiêu đề
    p_title = document.add_paragraph(doc_type_label)
    set_paragraph_format(p_title, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=Pt(12), space_after=Pt(6))
    set_run_format(p_title.runs[0], size=FONT_SIZE_TITLE, bold=True)

    # Trích yếu nội dung thông báo
    subject = title.replace("Thông báo", "").strip()
    if subject.lower().startswith("về việc"):
        subject = subject.split(" ", 2)[-1]
    p_subject = document.add_paragraph(f"V/v: {subject}")
    set_paragraph_format(p_subject, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(12))
    set_run_format(p_subject.runs[0], size=Pt(14), bold=True) # Trích yếu đậm


    # 3. Kính gửi (Nếu có)
    recipient = data.get("recipient_main", None)
    if recipient:
        p_kg = document.add_paragraph(f"Kính gửi: {recipient}")
        set_paragraph_format(p_kg, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=Pt(12))
        set_run_format(p_kg.runs[0], size=FONT_SIZE_DEFAULT, bold=True)


    # 4. Nội dung Thông báo
    body_lines = body.split('\n')
    processed_indices = set()

    # Xử lý căn cứ nếu có
    for i, line in enumerate(body_lines):
        stripped_line = line.strip()
        if not stripped_line: continue
        if stripped_line.lower().startswith("căn cứ"):
            p = document.add_paragraph(stripped_line)
            set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line_indent=FIRST_LINE_INDENT, line_spacing=1.5, space_after=Pt(0))
            set_run_format(p.runs[0], size=FONT_SIZE_DEFAULT, italic=True)
            processed_indices.add(i)
        elif any(l.strip().lower().startswith("căn cứ") for l in body_lines[:i]):
             break
    if processed_indices: document.add_paragraph()

    # Xử lý nội dung chính
    for i, line in enumerate(body_lines):
        if i in processed_indices: continue
        stripped_line = line.strip()
        if not stripped_line: continue
        p = document.add_paragraph()

        is_numbered_item = re.match(r'^(\d+)\.\s+', stripped_line) # 1., 2.
        is_bullet = stripped_line.startswith("-") or stripped_line.startswith("+") or stripped_line.startswith("*") or stripped_line.startswith("•")
        is_ending = "trân trọng thông báo" in stripped_line.lower() or "đề nghị các đơn vị" in stripped_line.lower()

        align = WD_ALIGN_PARAGRAPH.JUSTIFY
        left_indent = Cm(0)
        first_indent = FIRST_LINE_INDENT if not (is_numbered_item or is_bullet) else Cm(0)
        is_bold = False
        is_italic = False
        size = FONT_SIZE_DEFAULT
        space_before = Pt(0)
        space_after = Pt(6)
        line_spacing = 1.5

        if is_numbered_item:
            align = WD_ALIGN_PARAGRAPH.LEFT
            left_indent = Cm(0.5) # Thụt lề mục
        elif is_bullet:
             align = WD_ALIGN_PARAGRAPH.LEFT
             left_indent = Cm(1.0)
             first_indent = Cm(-0.5) # Hanging indent
        elif is_ending:
            align = WD_ALIGN_PARAGRAPH.LEFT
            first_indent = FIRST_LINE_INDENT
            is_italic = True # Câu kết thường nghiêng
            space_before = Pt(12)

        set_paragraph_format(p, alignment=align, left_indent=left_indent, first_line_indent=first_indent, line_spacing=line_spacing, space_before=space_before, space_after=space_after)
        add_run_with_format(p, stripped_line, size=size, bold=is_bold, italic=is_italic)


    # 5. Chữ ký (Sử dụng khối chữ ký cơ bản)
    if 'signer_title' not in data: data['signer_title'] = "THỦ TRƯỞNG CƠ QUAN"
    if 'signer_name' not in data: data['signer_name'] = "[Họ và tên]"
    # Thông báo có thể không cần khoảng cách lớn trước chữ ký nếu câu kết đã có space_before
    format_signature_block(document, data)

    # 6. Nơi nhận (Sử dụng nơi nhận cơ bản)
    if 'recipients' not in data: data['recipients'] = ["- Như trên;", "- Lưu: VT, ...;"]
    format_recipient_list(document, data)

    print("Định dạng Thông báo hoàn tất.")