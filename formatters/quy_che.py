# formatters/quy_che.py
import re
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from utils import set_paragraph_format, set_run_format, add_run_with_format
try:
    from .common_elements import format_basic_header # Chỉ dùng nếu là QC độc lập
except ImportError:
    from common_elements import format_basic_header
from config import FONT_SIZE_DEFAULT, FONT_SIZE_TITLE, FIRST_LINE_INDENT, FONT_SIZE_HEADER

def format(document, data):
    print("Bắt đầu định dạng Quy chế...")
    title = data.get("title", "Quy chế làm việc của [Tên đơn vị]")
    body = data.get("body", "Nội dung quy chế...")
    issuing_org = data.get("issuing_org", "TÊN CƠ QUAN BAN HÀNH").upper()
    attached_decision = data.get("attached_decision", None)

    # 1. Header (Tùy chọn, thường của CQ ban hành QĐ kèm theo)
    # Nếu là QC độc lập, có thể thêm header:
    # data['issuing_org'] = issuing_org
    # format_basic_header(document, data, "QuyChe")

    # 2. Tên loại QUY CHẾ
    p_tenloai = document.add_paragraph("QUY CHẾ")
    set_paragraph_format(p_tenloai, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=Pt(12), space_after=Pt(6))
    add_run_with_format(p_tenloai, "QUY CHẾ", size=FONT_SIZE_TITLE, bold=True, uppercase=True)

    # 3. Tiêu đề Quy chế
    qc_title = title.replace("Quy chế", "").strip()
    p_title = document.add_paragraph(qc_title)
    set_paragraph_format(p_title, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(6))
    add_run_with_format(p_title, qc_title, size=Pt(14), bold=True) # Tiêu đề QC đậm

    # 4. Dòng kèm theo Quyết định
    if attached_decision:
        p_attach = document.add_paragraph(f"(Ban hành kèm theo {attached_decision})")
        set_paragraph_format(p_attach, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(12))
        add_run_with_format(p_attach, p_attach.text, size=FONT_SIZE_DEFAULT, italic=True)

    # 5. Nội dung (Chương, Điều, Khoản, Điểm)
    body_lines = body.split('\n')
    for line in body_lines:
        stripped_line = line.strip()
        if not stripped_line: continue
        p = document.add_paragraph()
        is_chuong = stripped_line.upper().startswith("CHƯƠNG")
        is_dieu = stripped_line.upper().startswith("ĐIỀU")
        is_khoan = re.match(r'^\d+\.\s+', stripped_line)
        is_diem = re.match(r'^[a-z]\)\s+', stripped_line)

        align = WD_ALIGN_PARAGRAPH.JUSTIFY
        left_indent = Cm(0)
        first_indent = FIRST_LINE_INDENT
        is_bold = False
        size = FONT_SIZE_DEFAULT
        space_before = Pt(0)

        if is_chuong:
            align = WD_ALIGN_PARAGRAPH.CENTER
            first_indent = Cm(0)
            is_bold = True
            space_before = Pt(12)
        elif is_dieu:
            align = WD_ALIGN_PARAGRAPH.LEFT
            first_indent = Cm(0)
            is_bold = True
            space_before = Pt(6)
        elif is_khoan:
            left_indent = Cm(0.5)
            first_indent = Cm(0)
        elif is_diem:
            left_indent = Cm(1.0)
            first_indent = Cm(0)

        set_paragraph_format(p, alignment=align, left_indent=left_indent, first_line_indent=first_indent, line_spacing=1.5, space_before=space_before, space_after=Pt(6))
        add_run_with_format(p, stripped_line, size=size, bold=is_bold)

    print("Định dạng Quy chế hoàn tất.")