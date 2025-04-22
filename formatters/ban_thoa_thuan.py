# formatters/ban_thoa_thuan.py
import re
import time
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from utils import set_paragraph_format, set_run_format, add_run_with_format
from config import FONT_SIZE_DEFAULT, FONT_SIZE_TITLE, FONT_SIZE_SIGNATURE, FONT_SIZE_SIGNER_NAME, FONT_SIZE_PLACE_DATE, FONT_SIZE_HEADER, FIRST_LINE_INDENT

# Sử dụng lại hàm format parties và signature từ ban_ghi_nho
try:
    from .ban_ghi_nho import format_parties_block, format_signature_mou
except ImportError:
    # Fallback nếu chạy độc lập (ít xảy ra)
    def format_parties_block(document, parties): pass
    def format_signature_mou(document, signer_a_title, signer_a_name, signer_b_title, signer_b_name): pass


def format(document, data):
    print("Bắt đầu định dạng Bản thỏa thuận...")
    title = data.get("title", "BẢN THỎA THUẬN").upper()
    subtitle = data.get("subtitle", "Về việc Hợp tác XYZ")
    agreement_number = data.get("agreement_number", None) # Số thỏa thuận nếu có
    issuing_location = data.get("issuing_location", "Hà Nội")
    current_date_str = time.strftime(f"ngày %d tháng %m năm %Y")
    preamble = data.get("preamble", ["Căn cứ Bộ luật Dân sự số 91/2015/QH13 ngày 24 tháng 11 năm 2015;", "Căn cứ vào nhu cầu và khả năng của các Bên;"])
    parties_info = data.get("parties", [])
    body = data.get("body", "Điều 1: Nội dung thỏa thuận\n...")
    signer_a_title = data.get("signer_a_title", "ĐẠI DIỆN BÊN A")
    signer_a_name = data.get("signer_a_name", "[Tên người ký Bên A]")
    signer_b_title = data.get("signer_b_title", "ĐẠI DIỆN BÊN B")
    signer_b_name = data.get("signer_b_name", "[Tên người ký Bên B]")


    # 1. Quốc hiệu, Tiêu ngữ (Căn giữa)
    p_qh = document.add_paragraph("CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM")
    set_paragraph_format(p_qh, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(0))
    set_run_format(p_qh.runs[0], size=FONT_SIZE_HEADER, bold=True)
    p_tn = document.add_paragraph("Độc lập - Tự do - Hạnh phúc")
    set_paragraph_format(p_tn, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(0))
    set_run_format(p_tn.runs[0], size=Pt(13), bold=True)
    p_line_tn = document.add_paragraph("-" * 20)
    set_paragraph_format(p_line_tn, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(18))

    # 2. Tên loại văn bản
    p_title = document.add_paragraph(title)
    set_paragraph_format(p_title, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=Pt(12), space_after=Pt(6))
    set_run_format(p_title.runs[0], size=FONT_SIZE_TITLE, bold=True)

    # 3. Tiêu đề phụ / Số thỏa thuận (nếu có)
    sub_line = subtitle
    if agreement_number:
         sub_line += f"\nSố: {agreement_number}"

    p_subtitle = document.add_paragraph(sub_line)
    set_paragraph_format(p_subtitle, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(12))
    set_run_format(p_subtitle.runs[0], size=Pt(14), bold=True) # Có thể tách riêng số để định dạng khác

    # 4. Địa điểm, Ngày tháng (Căn phải, nghiêng)
    p_place_date = document.add_paragraph(f"{issuing_location}, {current_date_str}")
    set_paragraph_format(p_place_date, alignment=WD_ALIGN_PARAGRAPH.RIGHT, space_after=Pt(12))
    set_run_format(p_place_date.runs[0], size=FONT_SIZE_PLACE_DATE, italic=True)

    # 5. Căn cứ ban hành
    if isinstance(preamble, str): preamble = preamble.split('\n')
    if preamble:
        for line in preamble:
            p_pre = document.add_paragraph(line)
            set_paragraph_format(p_pre, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=Pt(0), first_line_indent=FIRST_LINE_INDENT, line_spacing=1.5)
            set_run_format(p_pre.runs[0], size=FONT_SIZE_DEFAULT, italic=True)
        document.add_paragraph() # Khoảng cách sau căn cứ

    p_intro = document.add_paragraph("Hôm nay, ngày... tháng... năm..., tại..., chúng tôi gồm:") # Có thể lấy từ data
    set_paragraph_format(p_intro, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=Pt(6), line_spacing=1.5)
    set_run_format(p_intro.runs[0], size=FONT_SIZE_DEFAULT)

    # 6. Thông tin các bên
    format_parties_block(document, parties_info)
    document.add_paragraph()

    p_agree = document.add_paragraph("Sau khi bàn bạc, hai bên cùng thống nhất ký kết Bản thỏa thuận này với các điều khoản cụ thể như sau:")
    set_paragraph_format(p_agree, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=Pt(6), first_line_indent=FIRST_LINE_INDENT, line_spacing=1.5)
    set_run_format(p_agree.runs[0], size=FONT_SIZE_DEFAULT)

    # 7. Nội dung chính (Điều khoản)
    body_lines = body.split('\n')
    for line in body_lines:
        stripped_line = line.strip()
        if not stripped_line: continue
        p = document.add_paragraph()
        is_dieu = stripped_line.upper().startswith("ĐIỀU")
        is_khoan = re.match(r'^\d+\.\s+', stripped_line)
        is_diem = re.match(r'^[a-z]\)\s+', stripped_line)

        align = WD_ALIGN_PARAGRAPH.JUSTIFY
        left_indent = Cm(0)
        first_indent = FIRST_LINE_INDENT
        is_bold = False
        size = FONT_SIZE_DEFAULT
        space_before = Pt(0)
        space_after = Pt(6)

        if is_dieu:
            align = WD_ALIGN_PARAGRAPH.LEFT
            first_indent = Cm(0)
            is_bold = True
            space_before = Pt(6)
        elif is_khoan:
            left_indent = Cm(0.5)
            first_indent = Cm(0)
        elif is_diem:
            left_indent = Cm(1.0)
            first_indent = Cm(0)

        set_paragraph_format(p, alignment=align, left_indent=left_indent, first_line_indent=first_indent, line_spacing=1.5, space_before=space_before, space_after=space_after)
        add_run_with_format(p, stripped_line, size=size, bold=is_bold)

    # 8. Điều khoản cuối (Hiệu lực, số bản...)
    # Thường nằm trong body, nên xử lý như Điều khoản thông thường

    # 9. Chữ ký (Hai bên)
    document.add_paragraph() # Khoảng cách trước chữ ký
    format_signature_mou(document, signer_a_title, signer_a_name, signer_b_title, signer_b_name)

    print("Định dạng Bản thỏa thuận hoàn tất.")