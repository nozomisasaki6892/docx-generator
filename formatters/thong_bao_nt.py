# formatters/thong_bao_nt.py
import re
import time
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from utils import set_paragraph_format, set_run_format, add_run_with_format
try:
    # Dùng header, signature, recipient như Thông báo hành chính
    from .common_elements import format_basic_header, format_signature_block, format_recipient_list
except ImportError:
    def format_basic_header(document, data, doc_type): pass
    def format_signature_block(document, data): pass
    def format_recipient_list(document, data): pass

from config import FONT_SIZE_DEFAULT, FONT_SIZE_TITLE, FIRST_LINE_INDENT

def format(document, data):
    print("Bắt đầu định dạng Thông báo Nhà trường...")
    title = data.get("title", "Thông báo về việc Nộp học phí/Lịch thi...")
    body = data.get("body", "Kính gửi: Toàn thể sinh viên...\nNhà trường thông báo kế hoạch... như sau:\n1...\n2...\nĐề nghị sinh viên nghiêm túc thực hiện.")
    doc_type_label = "THÔNG BÁO"

    # 1. Header (Tên trường)
    if 'issuing_org' not in data: data['issuing_org'] = "TRƯỜNG ĐẠI HỌC XYZ"
    format_basic_header(document, data, "ThongBaoNT")

    # 2. Tiêu đề
    p_title = document.add_paragraph(doc_type_label)
    set_paragraph_format(p_title, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=Pt(12), space_after=Pt(6))
    set_run_format(p_title.runs[0], size=FONT_SIZE_TITLE, bold=True)

    # Trích yếu nội dung thông báo
    subject = title.replace("Thông báo", "").strip()
    if subject.lower().startswith("về việc"):
        subject = subject.split(" ", 2)[-1]
    p_subject = document.add_paragraph(f"V/v: {subject}")
    set_paragraph_format(p_subject, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(12))
    set_run_format(p_subject.runs[0], size=Pt(14), bold=True)


    # 3. Kính gửi (Toàn thể SV, các Khoa...)
    recipient = data.get("recipient_main", "Kính gửi: Toàn thể Sinh viên")
    p_kg = document.add_paragraph(recipient)
    set_paragraph_format(p_kg, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=Pt(12))
    set_run_format(p_kg.runs[0], size=FONT_SIZE_DEFAULT, bold=True)


    # 4. Nội dung Thông báo
    body_lines = body.split('\n')
    processed_indices = set()

    # Xử lý căn cứ nếu có
    for i, line in enumerate(body_lines):
        stripped_line = line.strip()
        if not stripped_line: continue
        if stripped_line.lower().startswith("căn cứ"):
            p = document.add_paragraph(stripped_line)
            set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line_indent=FIRST_LINE_INDENT, line_spacing=1.5, space_after=Pt(0))
            set_run_format(p.runs[0], size=FONT_SIZE_DEFAULT, italic=True)
            processed_indices.add(i)
        elif any(l.strip().lower().startswith("căn cứ") for l in body_lines[:i]):
             break
    if processed_indices: document.add_paragraph()

    # Xử lý nội dung chính
    for i, line in enumerate(body_lines):
        if i in processed_indices: continue
        stripped_line = line.strip()
        if not stripped_line: continue
        p = document.add_paragraph()

        is_numbered_item = re.match(r'^(\d+)\.\s+', stripped_line) # 1., 2.
        is_bullet = stripped_line.startswith("-") or stripped_line.startswith("+") or stripped_line.startswith("*") or stripped_line.startswith("•")
        is_ending = "trân trọng thông báo" in stripped_line.lower() or "đề nghị sinh viên" in stripped_line.lower() or "đề nghị các đơn vị" in stripped_line.lower()

        align = WD_ALIGN_PARAGRAPH.JUSTIFY
        left_indent = Cm(0)
        first_indent = FIRST_LINE_INDENT if not (is_numbered_item or is_bullet) else Cm(0)
        is_bold = False
        is_italic = False
        size = FONT_SIZE_DEFAULT
        space_before = Pt(0)
        space_after = Pt(6)
        line_spacing = 1.5

        if is_numbered_item:
            align = WD_ALIGN_PARAGRAPH.LEFT
            left_indent = Cm(0.5)
        elif is_bullet:
             align = WD_ALIGN_PARAGRAPH.LEFT
             left_indent = Cm(1.0)
             first_indent = Cm(-0.5)
        elif is_ending:
            align = WD_ALIGN_PARAGRAPH.LEFT
            first_indent = FIRST_LINE_INDENT
            is_italic = True
            space_before = Pt(12)

        set_paragraph_format(p, alignment=align, left_indent=left_indent, first_line_indent=first_indent, line_spacing=line_spacing, space_before=space_before, space_after=space_after)
        add_run_with_format(p, stripped_line, size=size, bold=is_bold, italic=is_italic)


    # 5. Chữ ký (Thường là Hiệu trưởng/Phó hiệu trưởng, Trưởng phòng ban)
    if 'signer_title' not in data: data['signer_title'] = "KT. HIỆU TRƯỞNG\nPHÓ HIỆU TRƯỞNG" # Ví dụ
    if 'signer_name' not in data: data['signer_name'] = "[Họ và tên]"
    format_signature_block(document, data)

    # 6. Nơi nhận
    if 'recipients' not in data: data['recipients'] = ["- Như kính gửi;", "- Các Khoa, Phòng ban;", "- Lưu: VT, CTSV."]
    format_recipient_list(document, data)

    print("Định dạng Thông báo Nhà trường hoàn tất.")